#!/usr/bin/env python3
"""
Test script to verify ProductStrain fix for vertical template
"""

import sys
import os
import logging

# Add src to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

def test_product_strain_fix():
    """Test that ProductStrain values now appear in vertical template"""
    
    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(),
        ]
    )
    logger = logging.getLogger(__name__)
    
    try:
        # Load data
        from src.core.data.excel_processor import ExcelProcessor
        from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
        
        # Initialize processor with database file
        database_path = "/Users/adamcordova/Desktop/labelMaker_ QR copy SAFEST copy 4/uploads/A Greener Today - Bothell_inventory_09-26-2025  4_51 PM.xlsx"
        if not os.path.exists(database_path):
            logger.error(f"Database file not found: {database_path}")
            return
            
        processor = ExcelProcessor()
        if not processor.load_file(database_path):
            logger.error(f"Failed to load database file: {database_path}")
            return
            
        # Convert to records format
        all_records = processor.df.to_dict('records')
        logger.info(f"Loaded {len(all_records)} records from Excel")
        
        # Find a product with ProductStrain value for testing
        test_product = None
        for record in all_records:
            strain = record.get('Product Strain', '').strip()
            if strain and strain not in ['', 'Mixed', 'Paraphernalia']:
                test_product = record
                break
                
        if not test_product:
            logger.error("No suitable test product found with ProductStrain value")
            return
            
        logger.info(f"Testing with product: {test_product.get('Product Name', 'Unknown')}")
        logger.info(f"ProductStrain: '{test_product.get('Product Strain', 'N/A')}'")
        
        # Generate vertical template
        font_scheme = get_font_scheme('vertical')
        template_processor = TemplateProcessor('vertical', font_scheme, 1.0)
        
        # Generate document
        logger.info("=== GENERATING VERTICAL TEMPLATE WITH PRODUCT STRAIN ===")
        final_doc = template_processor.process_records([test_product])
        
        if not final_doc:
            logger.error("❌ Template processor returned None - no document generated")
            return
            
        # Save the document for inspection
        output_path = "./test_vertical_fix_output.docx"
        final_doc.save(output_path)
        logger.info(f"Generated document: {output_path}")
        
        # Read the generated document and check for ProductStrain content
        from docx import Document
        doc = Document(output_path)
        
        product_strain_found = False
        strain_value = test_product.get('Product Strain', '')
        
        logger.info(f"=== CHECKING GENERATED DOCUMENT FOR STRAIN VALUE: '{strain_value}' ===")
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if strain_value in cell.text:
                        logger.info(f"✅ SUCCESS: Found ProductStrain value '{strain_value}' in generated document!")
                        logger.info(f"   Cell text: {cell.text.strip()}")
                        product_strain_found = True
        
        # Check paragraphs
        for para in doc.paragraphs:
            if strain_value in para.text:
                logger.info(f"✅ SUCCESS: Found ProductStrain value '{strain_value}' in paragraph!")
                logger.info(f"   Paragraph text: {para.text.strip()}")
                product_strain_found = True
        
        if product_strain_found:
            logger.info("🎉 FIX SUCCESSFUL: ProductStrain is now appearing in vertical template!")
        else:
            logger.error("❌ FIX FAILED: ProductStrain value not found in generated document")
            # Show what placeholders exist
            logger.info("Checking for unreplaced placeholders...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'ProductStrain' in cell.text:
                            logger.error(f"   Found unreplaced placeholder: {cell.text.strip()}")
        
        # Keep the output file for manual inspection
        logger.info(f"Generated test file saved as: {output_path} (inspect manually)")
            
    except Exception as e:
        logger.error(f"Test failed: {e}", exc_info=True)

if __name__ == "__main__":
    test_product_strain_fix()