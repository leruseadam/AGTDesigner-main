#!/usr/bin/env python3
"""
Test script to verify preroll bold formatting fix
This script tests the formatting specifically for preroll products
"""

import sys
import os

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

def test_preroll_formatting():
    """Test the preroll formatting function"""
    try:
        from core.generation.docx_formatting import enforce_preroll_bold_formatting
        from docx import Document
        
        print("Testing preroll bold formatting...")
        
        # Create a test document with preroll content
        doc = Document()
        
        # Add preroll-related content
        doc.add_paragraph("CONSTELLATION CANNABIS")
        doc.add_paragraph("GMO Infused Pre-Roll")
        doc.add_paragraph("-0.5g x 2 Pack")
        doc.add_paragraph("INDICA")
        doc.add_paragraph("CHAPTER 246-70 WAC GENERAL USE COMPLIANT")
        doc.add_paragraph("$15")
        
        # Add non-preroll content
        doc.add_paragraph("Regular Product Name")
        doc.add_paragraph("Some other content")
        
        print("Created test document with preroll content")
        
        # Apply preroll formatting
        formatted_doc = enforce_preroll_bold_formatting(doc)
        
        print("Applied preroll formatting")
        
        # Check if formatting was applied
        preroll_paragraphs = []
        for paragraph in formatted_doc.paragraphs:
            if any(keyword in paragraph.text.lower() for keyword in ['constellation', 'infused', 'pre-roll', 'indica', 'chapter']):
                preroll_paragraphs.append(paragraph)
        
        print(f"Found {len(preroll_paragraphs)} preroll-related paragraphs")
        
        # Check bold formatting
        bold_count = 0
        for paragraph in preroll_paragraphs:
            for run in paragraph.runs:
                if run.font.bold:
                    bold_count += 1
                    print(f"✅ Bold formatting applied to: '{run.text[:30]}...'")
                else:
                    print(f"❌ No bold formatting for: '{run.text[:30]}...'")
        
        print(f"Total bold runs: {bold_count}")
        
        if bold_count > 0:
            print("✅ SUCCESS: Preroll formatting is working!")
            return True
        else:
            print("❌ FAILURE: No bold formatting was applied")
            return False
            
    except ImportError as e:
        print(f"❌ Import error: {e}")
        print("Make sure you're running this from the project root directory")
        return False
    except Exception as e:
        print(f"❌ Error testing preroll formatting: {e}")
        return False

def main():
    """Run the preroll formatting test"""
    print("=" * 50)
    print("PREROLL BOLD FORMATTING TEST")
    print("=" * 50)
    
    success = test_preroll_formatting()
    
    print("\n" + "=" * 50)
    if success:
        print("🎉 PREROLL FORMATTING TEST PASSED!")
        print("Preroll labels should now appear in bold formatting.")
    else:
        print("⚠️  PREROLL FORMATTING TEST FAILED!")
        print("Check the error messages above for details.")
    print("=" * 50)

if __name__ == "__main__":
    main()
