from copy import deepcopy
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from io import BytesIO
import logging
import os
from pathlib import Path
import re
from typing import Dict, Any, List, Optional
import traceback
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
# from docx.oxml.shared import OxmlElement, qn  # Duplicate import removed
import time
import pandas as pd
import qrcode
from PIL import Image

# Local imports
from src.core.utils.common import safe_get
from src.core.generation.docx_formatting import (
    apply_lineage_colors,
    enforce_fixed_cell_dimensions,
    clear_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
)
from src.core.generation.unified_font_sizing import (
    get_font_size,
    get_font_size_by_marker,
    set_run_font_size,
    is_classic_type,
    get_line_spacing_by_marker
)
from src.core.generation.text_processing import (
    process_doh_image,
    format_ratio_multiline
)
from src.core.formatting.markers import wrap_with_marker, unwrap_marker, is_already_wrapped

# Performance settings - check if running on PythonAnywhere
import os
IS_PYTHONANYWHERE = 'pythonanywhere.com' in os.environ.get('HTTP_HOST', '')

# Use same settings for both local and PythonAnywhere to ensure consistent generation
MAX_PROCESSING_TIME_PER_CHUNK = 30  # 30 seconds max per chunk
MAX_TOTAL_PROCESSING_TIME = 300     # 5 minutes max total
CHUNK_SIZE_LIMIT = 50               # Limit chunk size for performance

def get_font_scheme(template_type, base_size=12):
    schemes = {
        'default': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'vertical': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'mini': {"base_size": base_size - 2, "min_size": 6, "max_length": 15},
        'horizontal': {"base_size": base_size + 1, "min_size": 7, "max_length": 20},
        'double': {"base_size": base_size - 1, "min_size": 8, "max_length": 30},
        'inventory': {"base_size": base_size, "min_size": 8, "max_length": 40}  # Inventory slips can handle longer text
    }
    return {
        field: {**schemes.get(template_type, schemes['default'])}
        for field in ["Description", "ProductBrand", "Price", "Lineage", "DOH", "Ratio_or_THC_CBD", "Ratio"]
    }

class TemplateProcessor:
    def __init__(self, template_type, font_scheme, scale_factor=1.0):
        self.template_type = template_type
        self.font_scheme = font_scheme
        self.scale_factor = scale_factor
        self.logger = logging.getLogger(__name__)
        self._template_path = self._get_template_path()
        self._expanded_template_buffer = self._expand_template_if_needed()
        
        # Set chunk size based on template type with performance limits
        if not IS_PYTHONANYWHERE:
            self.logger.info(f"DEBUG: Setting chunk size for template_type='{self.template_type}' (type: {type(self.template_type)})")
        
        if self.template_type == 'mini':
            self.chunk_size = min(20, CHUNK_SIZE_LIMIT)  # Fixed: 4x5 grid = 20 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for mini template")
        elif self.template_type == 'double':
            self.chunk_size = min(12, CHUNK_SIZE_LIMIT)  # Fixed: 4x3 grid = 12 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for double template")
        elif self.template_type == 'inventory':
            self.chunk_size = min(4, CHUNK_SIZE_LIMIT)   # Fixed: 2x2 grid = 4 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for inventory template")
        else:
            # For standard templates (horizontal, vertical), use 3x3 grid = 9 labels per page
            self.chunk_size = min(9, CHUNK_SIZE_LIMIT)  # Fixed: 3x3 grid = 9 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for template type '{self.template_type}' (fallback to 3x3)")
        
        self.logger.info(f"Template type: {self.template_type}, Chunk size: {self.chunk_size}")
        
        # Performance tracking
        self.start_time = time.time()
        self.chunk_count = 0

    def _get_template_path(self):
        """Get the template path based on template type."""
        try:
            base_path = Path(__file__).resolve().parent / "templates"
            # Map template types to filenames
            template_files = {
                'horizontal': 'horizontal.docx',
                'vertical': 'vertical.docx',
                'mini': 'mini.docx',
                'double': 'double.docx',
                'inventory': 'inventory.docx'
            }
            template_name = template_files.get(self.template_type, f"{self.template_type}.docx")
            template_path = base_path / template_name
            
            if not template_path.exists():
                # Fallback: case-insensitive match ONLY for non-hidden files (ignore . and ~$ temp files)
                expected_lower = template_name.lower()
                fallback = None
                for p in base_path.iterdir():
                    if not p.is_file():
                        continue
                    name = p.name
                    if name.startswith('.') or name.startswith('~$'):
                        continue
                    if name.lower() == expected_lower:
                        fallback = p
                        break
                if fallback and fallback.exists():
                    self.logger.warning(f"Using fallback template due to case-only mismatch: {fallback}")
                    return fallback
                self.logger.error(f"Template not found: {template_path}")
                raise FileNotFoundError(f"Template not found: {template_path}")
            
            return template_path
        except Exception as e:
            self.logger.error(f"Error getting template path: {e}")
            raise

    def _expand_template_if_needed(self, force_expand=False):
        """Expand template if needed and return buffer."""
        try:
            with open(self._template_path, 'rb') as f:
                buffer = BytesIO(f.read())
            
            # Check if template needs expansion
            doc = Document(buffer)
            text = doc.element.body.xml
            matches = re.findall(r'Label(\d+)\.', text)
            
            # Check if we have all required labels (9 for 3x3, 20 for 4x5, 12 for 4x3, 4 for 2x2)
            if self.template_type == 'mini':
                required_labels = 20  # 4x5 grid
            elif self.template_type == 'double':
                required_labels = 12  # 4x3 grid
            elif self.template_type == 'inventory':
                required_labels = 4   # 2x2 grid
            else:
                required_labels = 9   # 3x3 grid
            
            unique_labels = set(matches)
            
            # For double template, always expand to 4x3 grid
            if self.template_type == 'double':
                self.logger.info("Double template expanding to 4x3 grid")
                return self._expand_template_to_4x3_fixed_double()
            elif len(unique_labels) < required_labels or force_expand:
                # Templates need expansion to create proper grid layouts
                if self.template_type == 'mini':
                    self.logger.info("Calling 4x5 expansion method")
                    return self._expand_template_to_4x5_fixed_scaled()
                elif self.template_type == 'inventory':
                    self.logger.info("Calling 2x2 inventory expansion method")
                    return self._expand_template_to_2x2_inventory()
                elif self.template_type == 'double':
                    self.logger.info("Calling 4x3 expansion method")
                    return self._expand_template_to_4x3_fixed_double()
                else:
                    # horizontal and vertical templates expand to 3x3 grid
                    self.logger.info(f"Calling 3x3 expansion method for template type: '{self.template_type}'")
                    return self._expand_template_to_3x3_fixed()
            
            return buffer
        except Exception as e:
            self.logger.error(f"Error expanding template: {e}")
            raise

    def force_re_expand_template(self):
        """Force re-expansion of template."""
        self._expanded_template_buffer = self._expand_template_if_needed(force_expand=True)

    def _fix_double_template_structure(self):
        """Fix double template structure by adding missing table grid elements without full expansion."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        
        try:
            # Load the original template
            template_path = self._get_template_path()
            self.logger.info(f"Loading double template from: {template_path}")
            doc = Document(template_path)
            
            # Check if template has tables
            if not doc.tables:
                self.logger.warning("Double template has no tables, returning original")
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer
            
            self.logger.info(f"Double template has {len(doc.tables)} tables")
            
            # Fix each table by ensuring it has proper structure
            for table in doc.tables:
                # Check if table has tblGrid element
                tbl_grid = table._element.find(qn('w:tblGrid'))
                if tbl_grid is None:
                    self.logger.info("Adding missing tblGrid element to double template table")
                    
                    # Get the number of columns from the table XML structure instead of table.columns
                    # This avoids the error when table.columns is accessed without proper structure
                    table_element = table._element
                    rows = table_element.findall(qn('w:tr'))
                    if rows:
                        # Count columns from the first row
                        first_row = rows[0]
                        cells = first_row.findall(qn('w:tc'))
                        num_cols = len(cells)
                    else:
                        # Fallback: assume 4 columns for double template
                        num_cols = 4
                    
                    self.logger.info(f"Detected {num_cols} columns in double template table")
                    
                    # Create tblGrid element
                    tbl_grid = OxmlElement('w:tblGrid')
                    for _ in range(num_cols):
                        gc = OxmlElement('w:gridCol')
                        gc.set(qn('w:w'), str(int(1.75 * 1440)))  # 1.75 inches per column
                        tbl_grid.append(gc)
                    
                    # Insert tblGrid at the beginning of the table element
                    table_element.insert(0, tbl_grid)
                    
                    # Also ensure table has proper table properties
                    tbl_pr = table_element.find(qn('w:tblPr'))
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        table_element.insert(0, tbl_pr)
                    
                    # Add table layout
                    layout = tbl_pr.find(qn('w:tblLayout'))
                    if layout is None:
                        layout = OxmlElement('w:tblLayout')
                        layout.set(qn('w:type'), 'fixed')
                        tbl_pr.append(layout)
                    
                    # Ensure the table has at least one row and cell for basic structure
                    if not table_element.findall(qn('w:tr')):
                        self.logger.warning("Double template table has no rows, adding minimal structure")
                        # Add a minimal row with cells
                        row = OxmlElement('w:tr')
                        for _ in range(num_cols):
                            cell = OxmlElement('w:tc')
                            # Add cell properties
                            tc_pr = OxmlElement('w:tcPr')
                            cell.append(tc_pr)
                            # Add a paragraph
                            para = OxmlElement('w:p')
                            cell.append(para)
                            row.append(cell)
                        table_element.append(row)
            
            # Save the fixed template to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            self.logger.error(f"Error fixing double template structure: {e}")
            # Fallback to original template if fixing fails
            try:
                self.logger.info("Attempting fallback to original template")
                with open(self._get_template_path(), 'rb') as f:
                    buffer = BytesIO(f.read())
                self.logger.info("Fallback successful, returning original template")
                return buffer
            except Exception as fallback_error:
                self.logger.error(f"Fallback to original template also failed: {fallback_error}")
                # Last resort: create a minimal working template
                self.logger.warning("Creating minimal working template as last resort")
                try:
                    from docx import Document
                    doc = Document()
                    table = doc.add_table(rows=1, cols=4)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer
                except Exception as create_error:
                    self.logger.error(f"Failed to create minimal template: {create_error}")
                    raise

    def _expand_template_to_2x2_inventory(self):
        """Expand template to 2x2 grid for inventory slips."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 2, 2  # 2x2 grid for inventory
        col_width_inches = 3.75  # Appropriate width for inventory slips
        row_height_inches = 3.5   # Appropriate height for inventory slips
        
        col_width_twips = str(int(col_width_inches * 1440))
        row_height_pts = Pt(row_height_inches * 72)
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Create new table with 2x2 grid
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table properties
        tblPr = tbl._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)

        # Set column widths
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)

        # Set row heights and copy template cell content with proper label numbering
        label_num = 1
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                new_tc = deepcopy(src_tc)
                
                # Update label numbering for 2x2 grid (Label1, Label2, Label3, Label4)
                # Convert the cell XML to string, replace Label1 with current label number
                tc_xml_str = new_tc.xml.decode('utf-8') if isinstance(new_tc.xml, bytes) else str(new_tc.xml)
                tc_xml_str = tc_xml_str.replace('Label1', f'Label{label_num}')
                
                # Parse the updated XML and replace the cell
                from lxml import etree
                new_tc_element = etree.fromstring(tc_xml_str.encode('utf-8'))
                cell._tc.getparent().replace(cell._tc, new_tc_element)
                label_num += 1

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _expand_template_to_4x5_fixed_scaled(self):
        """Expand template to 4x5 grid for mini templates while preserving original design."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 5
        col_width_twips = str(int(1.5 * 1440))  # 1.5 inches per column for equal width
        row_height_pts = Pt(1.5 * 72)  # 1.5 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        
        # Get the original table and its properties
        original_table = doc.tables[0]
        original_table_xml = original_table._element
        
        # Extract original table properties (colors, borders, styling)
        original_tblPr = original_table_xml.find(qn('w:tblPr'))
        original_shd = original_tblPr.find(qn('w:shd')) if original_tblPr is not None else None
        original_borders = original_tblPr.find(qn('w:tblBorders')) if original_tblPr is not None else None
        
        # Get the original cell structure and content
        original_cell = original_table.cell(0, 0)
        src_tc = deepcopy(original_cell._tc)
        
        # Remove the original table
        original_table._element.getparent().remove(original_table._element)

        # Remove empty paragraphs
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Create new 4x5 table
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set up table properties
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        
        # Preserve original shading if it exists
        if original_shd is not None:
            shd = deepcopy(original_shd)
            tblPr.insert(0, shd)
        else:
            # Default shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D3D3D3')
            tblPr.insert(0, shd)
        
        # Set table layout
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        
        # Set up grid columns
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        
        # Set row heights and individual cell widths
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            # Set individual cell widths to ensure exact 1.5" dimensions
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), col_width_twips)
                tcW.set(qn('w:type'), 'dxa')
        
        # Preserve original borders if they exist
        if original_borders is not None:
            borders = deepcopy(original_borders)
            tblPr.append(borders)
        else:
            # Default borders
            borders = OxmlElement('w:tblBorders')
            for side in ('insideH','insideV'):
                b = OxmlElement(f"w:{side}")
                b.set(qn('w:val'), "single")
                b.set(qn('w:sz'), "4")
                b.set(qn('w:color'), "D3D3D3")
                b.set(qn('w:space'), "0")
                borders.append(b)
            tblPr.append(borders)
        
        # Populate cells with original content, updating labels
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r, c)
                cell._tc.clear_content()
                
                # Copy the original cell structure
                tc = deepcopy(src_tc)
                
                # Update all Label1 references to the current label number
                for t in tc.iter(qn('w:t')):
                    if t.text and 'Label1' in t.text:
                        t.text = t.text.replace('Label1', f'Label{cnt}')
                
                # Copy all elements from the original cell
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                
                cnt += 1
        
        # Add cell spacing
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        
        # Save and return
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_4x3_fixed_double(self):
        """Expand template to 4x3 grid for double templates (4 columns, 3 rows)."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 3  # 4 columns, 3 rows for 12 labels total
        
        # Equal width columns: 1.125 inches each for a total of 4.5 inches
        col_width_twips = str(int(1.125 * 1440))  # 1.125 inches per column
        row_height_pts = Pt(2.5 * 72)  # 2.5 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        # Only remove empty paragraphs, preserve content paragraphs
        # This is important for templates that have content outside of tables
        doc_paragraphs = list(doc.paragraphs)  # Create a copy to avoid modification during iteration
        for paragraph in doc_paragraphs:
            if not paragraph.text.strip():
                paragraph._element.getparent().remove(paragraph._element)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Copy the original table properties and styling from the source template
        # instead of creating hardcoded styling
        if hasattr(old, '_element') and old._element is not None:
            old_tblPr = old._element.find(qn('w:tblPr'))
            if old_tblPr is not None:
                # Copy the original table properties
                tbl._element.insert(0, deepcopy(old_tblPr))
            else:
                # Fallback to minimal table properties if none exist
                tblPr = OxmlElement('w:tblPr')
                layout = OxmlElement('w:tblLayout')
                layout.set(qn('w:type'), 'fixed')
                tblPr.append(layout)
                tbl._element.insert(0, tblPr)
        else:
            # Fallback to minimal table properties
            tblPr = OxmlElement('w:tblPr')
            layout = OxmlElement('w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            tblPr.append(layout)
            tbl._element.insert(0, tblPr)
        
        # Set up the grid with proper column widths
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        
        # Set row heights
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Process all cells normally (no gutters)
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                
                # Define all required placeholders for double template
                required_placeholders = [
                    'ProductBrand',
                    'DescAndWeight', 
                    'Price',
                    'Ratio_or_THC_CBD',
                    'DOH',
                    'QR'
                ]
                
                # Keep adding missing placeholders until all are present
                max_iterations = 10  # Prevent infinite loops
                iteration = 0
                
                while iteration < max_iterations:
                    # Get current cell text
                    cell_text = ''
                    for t in tc.iter(qn('w:t')):
                        if t.text:
                            cell_text += t.text
                    
                    self.logger.debug(f"Cell {cnt} iteration {iteration} current cell text: {repr(cell_text)}")
                    
                    # Check which placeholders are missing
                    missing_placeholders = []
                    for placeholder in required_placeholders:
                        placeholder_pattern = f'{{{{Label{cnt}.{placeholder}}}}}'
                        if placeholder_pattern not in cell_text and placeholder not in cell_text:
                            missing_placeholders.append(placeholder)
                    
                    self.logger.debug(f"Cell {cnt} iteration {iteration} checking placeholders:")
                    for placeholder in required_placeholders:
                        placeholder_pattern = f'{{{{Label{cnt}.{placeholder}}}}}'
                        in_cell = placeholder_pattern in cell_text
                        in_cell_simple = placeholder in cell_text
                        self.logger.debug(f"  {placeholder}: pattern={placeholder_pattern}, in_cell={in_cell}, in_cell_simple={in_cell_simple}")
                    
                    if not missing_placeholders:
                        # All placeholders are present, we're done
                        self.logger.debug(f"Cell {cnt} iteration {iteration} - all placeholders present, breaking")
                        break
                    
                    self.logger.debug(f"Cell {cnt} iteration {iteration} missing placeholders: {missing_placeholders}")
                    
                    # Add missing placeholders to the existing cell text with markers
                    # Special handling: ProductBrand should go in a separate paragraph to avoid interfering with vendor/strain
                    brand_placeholder = None
                    other_placeholders = []
                    
                    for placeholder in missing_placeholders:
                        if placeholder == 'ProductBrand':
                            brand_placeholder = placeholder
                        else:
                            other_placeholders.append(placeholder)
                    
                    # Add non-brand placeholders to the current paragraph
                    if other_placeholders:
                        placeholder_text = ''
                        for placeholder in other_placeholders:
                            # Map placeholders to their marker types
                            marker_mapping = {
                                'DescAndWeight': 'DESC',
                                'Price': 'PRICE',
                                'Ratio_or_THC_CBD': 'THC_CBD',
                                'DOH': 'DOH'
                            }
                            
                            marker_type = marker_mapping.get(placeholder, 'DEFAULT')
                            placeholder_text += f'\n{marker_type}_START{{{{Label{cnt}.{placeholder}}}}}{marker_type}_END'
                        
                        # Find the last text element and append the non-brand placeholders
                        text_elements = list(tc.iter(qn('w:t')))
                        if text_elements:
                            last_text_element = text_elements[-1]
                            if last_text_element.text:
                                last_text_element.text += placeholder_text
                            else:
                                last_text_element.text = placeholder_text
                    
                    # Add ProductBrand in a separate paragraph to avoid interference
                    if brand_placeholder:
                        # Create a new paragraph element for ProductBrand
                        new_para = OxmlElement('w:p')
                        new_run = OxmlElement('w:r')
                        new_text = OxmlElement('w:t')
                        new_text.text = f'PRODUCTBRAND_START{{{{Label{cnt}.ProductBrand}}}}PRODUCTBRAND_END'
                        new_run.append(new_text)
                        new_para.append(new_run)
                        tc.append(new_para)
                        
                        self.logger.debug(f"Added ProductBrand to separate paragraph in cell {cnt}")
                    
                    if other_placeholders:
                        self.logger.debug(f"Added other placeholders to existing paragraph in cell {cnt}: {other_placeholders}")
                        
                        self.logger.debug(f"Added missing placeholders to cell {cnt}: {missing_placeholders}")
                    
                    iteration += 1
                
                # CRITICAL FIX: Replace Label1 with Label{cnt} in all placeholders
                # Only replace exact matches to prevent creating extra labels
                for t in tc.iter(qn('w:t')):
                    if t.text and 'Label1' in t.text:
                        # Use regex to ensure we only replace Label1, not Label10, Label11, etc.
                        import re
                        t.text = re.sub(r'\bLabel1\b', f'Label{cnt}', t.text)
                
                # Copy the original cell content and styling exactly as it is
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                cnt += 1
                
                # CRITICAL FIX: Ensure cnt never exceeds 12 for double template
                if cnt > 12:
                    self.logger.error(f"CRITICAL: Counter exceeded 12! Current value: {cnt}")
                    break
                
        # Add minimal spacing between cells
        from docx.oxml.shared import OxmlElement as OE
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_3x3_fixed(self):
        """Expand template to 3x3 grid for standard templates."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 3, 3
        
        # Set dimensions based on template type - use constants for consistency
        from src.core.constants import CELL_DIMENSIONS
        
        cell_dims = CELL_DIMENSIONS.get(self.template_type, {'width': 2.4, 'height': 2.4})
        col_width_twips = str(int(cell_dims['width'] * 1440))  # Use width from constants
        row_height_pts = Pt(cell_dims['height'] * 72)  # Use height from constants
        # Use minimal spacing for vertical template to ensure all 9 labels fit
        if self.template_type == 'vertical':
            cut_line_twips = int(0.0001 * 1440)  # Minimal spacing for vertical
        else:
            cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        # Only remove empty paragraphs, preserve content paragraphs
        # This is important for templates that have content outside of tables
        doc_paragraphs = list(doc.paragraphs)  # Create a copy to avoid modification during iteration
        for paragraph in doc_paragraphs:
            if not paragraph.text.strip():
                paragraph._element.getparent().remove(paragraph._element)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        tblPr.insert(0, shd)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        borders = OxmlElement('w:tblBorders')
        for side in ('insideH','insideV'):
            b = OxmlElement(f"w:{side}")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:color'), "D3D3D3")
            b.set(qn('w:space'), "0")
            borders.append(b)
        tblPr.append(borders)
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                
                # Update Label1 references to Label{cnt} for proper grid expansion
                cell_text = ''
                for t in tc.iter(qn('w:t')):
                    if t.text:
                        cell_text += t.text
                        if 'Label1' in t.text:
                            t.text = t.text.replace('Label1', f'Label{cnt}')
                
                # CRITICAL FIX: The template uses Lineage and ProductVendor, not ProductBrand
                # No need to add missing placeholders - just copy the existing template structure
                
                # Add DOH placeholder if it's missing
                self.logger.debug(f"Cell {cnt} - cell_text: '{cell_text}'")
                self.logger.debug(f"Cell {cnt} - checking for DOH: '{{Label1.DOH}}' not in '{cell_text}' and 'DOH' not in '{cell_text}'")
                if '{{Label1.DOH}}' not in cell_text and 'DOH' not in cell_text:
                    self.logger.debug(f"Adding DOH placeholder to cell {cnt}")
                    # Find the position after the ProductStrain placeholder
                    text_elements = list(tc.iter(qn('w:t')))
                    strain_end_index = -1
                    
                    # Find where the ProductStrain placeholder ends
                    for i, t in enumerate(text_elements):
                        if t.text and 'ProductStrain' in t.text:
                            # Found the ProductStrain text element, look for the closing }}
                            for j in range(i, len(text_elements)):
                                if text_elements[j].text and '}}' in text_elements[j].text:
                                    strain_end_index = j
                                    break
                            break
                    
                    if strain_end_index >= 0:
                        self.logger.debug(f"Found ProductStrain end at index {strain_end_index}")
                        # Insert DOH placeholder after the ProductStrain placeholder
                        new_text = OxmlElement('w:t')
                        new_text.text = f'\n{{{{Label{cnt}.DOH}}}}'
                        
                        # Insert after the strain end element
                        strain_end_element = text_elements[strain_end_index]
                        strain_end_element.getparent().insert(
                            strain_end_element.getparent().index(strain_end_element) + 1, 
                            new_text
                        )
                        self.logger.debug(f"Inserted DOH placeholder: {new_text.text}")
                    else:
                        self.logger.warning(f"Could not find ProductStrain end position for cell {cnt}")
                else:
                    self.logger.debug(f"DOH placeholder already exists in cell {cnt}")
                
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                cnt += 1
        from docx.oxml.shared import OxmlElement as OE
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        
        # CRITICAL: Remove ALL headers and footers from the expanded template
        from src.core.generation.docx_formatting import remove_all_headers_and_footers
        doc = remove_all_headers_and_footers(doc)
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def process_records(self, records):
        """Process records with performance monitoring and timeout protection."""
        try:
            self.start_time = time.time()
            self.chunk_count = 0
            
            # Debug: Log the overall order of records
            overall_order = [record.get('ProductName', 'Unknown') for record in records]
            self.logger.info(f"Processing {len(records)} records in overall order: {overall_order}")
            
            # Deduplicate records by ProductName to prevent multiple outputs
            seen_products = set()
            unique_records = []
            for record in records:
                product_name = record.get('ProductName', 'Unknown')
                if product_name not in seen_products:
                    seen_products.add(product_name)
                    unique_records.append(record)
                else:
                    self.logger.warning(f"Skipping duplicate product: {product_name}")
            
            if len(unique_records) != len(records):
                self.logger.info(f"Deduplicated records: {len(records)} -> {len(unique_records)}")
                records = unique_records
            
            # Performance optimization: Log record count but don't limit
            if len(records) > 200:
                self.logger.info(f"Processing {len(records)} records (performance monitoring enabled)")
            else:
                self.logger.info(f"Processing {len(records)} records")
            
            documents = []
            for i in range(0, len(records), self.chunk_size):
                # Check total processing time
                if time.time() - self.start_time > MAX_TOTAL_PROCESSING_TIME:
                    self.logger.warning(f"Total processing time limit reached ({MAX_TOTAL_PROCESSING_TIME}s), stopping")
                    break
                
                chunk = records[i:i + self.chunk_size]
                self.chunk_count += 1
                
                self.logger.info(f"Processing chunk {self.chunk_count} ({len(chunk)} records)")
                result = self._process_chunk(chunk)
                if result: 
                    documents.append(result)
            
            if not documents: 
                return None
            if len(documents) == 1: 
                return documents[0]
            
            # Combine documents
            self.logger.info(f"Combining {len(documents)} documents")
            composer = Composer(documents[0])
            for doc in documents[1:]:
                composer.append(doc)
            
            final_doc_buffer = BytesIO()
            composer.save(final_doc_buffer)
            final_doc_buffer.seek(0)
            
            # CRITICAL: Remove ALL headers and footers from the final combined document
            final_doc = Document(final_doc_buffer)
            from src.core.generation.docx_formatting import remove_all_headers_and_footers
            final_doc = remove_all_headers_and_footers(final_doc)
            
            total_time = time.time() - self.start_time
            self.logger.info(f"Template processing completed in {total_time:.2f}s for {len(records)} records")
            
            return final_doc
        except Exception as e:
            self.logger.error(f"Error processing records: {e}")
            return None

    def _process_chunk(self, chunk):
        """Process a chunk of records with timeout protection."""
        chunk_start_time = time.time()
        
        try:
            if hasattr(self._expanded_template_buffer, 'seek'):
                self._expanded_template_buffer.seek(0)
            
            doc = DocxTemplate(self._expanded_template_buffer)
            
            # Debug: Log the order of records in this chunk
            chunk_order = [record.get('ProductName', 'Unknown') for record in chunk]
            self.logger.info(f"Processing chunk with {len(chunk)} records in order: {chunk_order}")
            
            # Build context for each record in the chunk
            context = {}
            for i, record in enumerate(chunk):
                # Set current record for brand centering logic
                self.current_record = record
                # Set current product type for brand marker processing
                self.current_product_type = (record.get('ProductType', '').lower() or 
                                          record.get('Product Type*', '').lower())
                if self.template_type == 'inventory':
                    label_context = self._build_inventory_context(record)
                else:
                    label_context = self._build_label_context(record, doc)
                context[f'Label{i+1}'] = label_context
                # Debug logging to check field values and order
                product_name = record.get('ProductName', 'Unknown')
                self.logger.debug(f"Label{i+1} -> {product_name} - ProductBrand: '{label_context.get('ProductBrand', 'NOT_FOUND')}', Price: '{label_context.get('Price', 'NOT_FOUND')}', THC: '{label_context.get('THC', 'NOT_FOUND')}', CBD: '{label_context.get('CBD', 'NOT_FOUND')}'")
            # Leave remaining labels blank instead of duplicating data
            for i in range(len(chunk), self.chunk_size):
                # Create empty context for unfilled labels
                context[f'Label{i+1}'] = {
                    'ProductBrand': '',
                    'DescAndWeight': '',
                    'Price': '',
                    'DOH': '',
                    'Ratio_or_THC_CBD': ''
                }
                self.logger.debug(f"Label{i+1} left blank (no data duplication)")

            # DOH images are already created in _build_label_context, no need for redundant creation here
            
            # Store original QR codes before replacing with placeholders
            original_qr_codes = {}
            if chunk:
                for i in range(1, len(chunk) + 1):
                    label_key = f'Label{i}'
                    if label_key in context and 'QR' in context[label_key]:
                        # Store the original QR code data before replacing with placeholder
                        original_qr_codes[label_key] = context[label_key]['QR']
                        # Replace with placeholder text for template rendering
                        context[label_key]['QR'] = 'QR_PLACEHOLDER'
                        self.logger.debug(f"Added QR placeholder for {label_key}")

            doc.render(context)
            
            # Manual QR code replacement after template rendering for all templates
            # Handle QR code replacement for each label
            if chunk and 'Label1' in context:
                self.logger.info("Attempting QR placeholder replacement for all labels")
                success_count = 0
                total_labels = len([k for k in context.keys() if k.startswith('Label')])
                
                # Replace QR placeholders with actual QR codes using stored original data
                for i in range(1, total_labels + 1):
                    label_key = f'Label{i}'
                    if label_key in original_qr_codes:
                        success = self._replace_qr_placeholder(doc, original_qr_codes[label_key], label_key)
                        if success:
                            success_count += 1
                            self.logger.debug(f"✓ QR code replaced for {label_key}")
                        else:
                            self.logger.warning(f"⚠ QR placeholder replacement failed for {label_key}")
                
                if success_count > 0:
                    self.logger.info(f"✓ QR placeholder replacement completed: {success_count}/{total_labels} labels")
                else:
                    self.logger.warning("⚠ No QR codes were replaced")
            else:
                self.logger.debug("QR replacement skipped - no labels in context")
            
            # For double templates, use manual placeholder replacement as fallback
            # since DocxTemplate was not working reliably
            if self.template_type == 'double':
                self.logger.info("Using manual placeholder replacement for double template")
                self._manual_replace_placeholders(doc, context)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            rendered_doc = Document(buffer)
            
            # Check timeout before post-processing
            if time.time() - chunk_start_time > MAX_PROCESSING_TIME_PER_CHUNK:
                self.logger.warning(f"Chunk processing timeout reached ({MAX_PROCESSING_TIME_PER_CHUNK}s), skipping post-processing")
                return rendered_doc
            
            # PRE-PROCESSING TABLE VALIDATION: Validate and repair all tables before any processing begins
            self.logger.debug("Starting pre-processing table validation")
            for table in rendered_doc.tables:
                try:
                    if not self._safe_table_iteration(table, "pre-processing validation"):
                        self.logger.error(f"Critical: Table validation failed during pre-processing, document may be corrupted")
                        # Continue processing but log the issue
                except Exception as e:
                    self.logger.error(f"Critical: Error during pre-processing table validation: {e}")
            
            # Post-process the document to apply dynamic font sizing first
            self._post_process_and_replace_content(rendered_doc)
            
            # Check timeout before lineage colors
            if time.time() - chunk_start_time > MAX_PROCESSING_TIME_PER_CHUNK:
                self.logger.warning(f"Chunk processing timeout reached ({MAX_PROCESSING_TIME_PER_CHUNK}s), skipping lineage colors")
                return rendered_doc
            
            # Apply lineage colors last to ensure they are not overwritten
            apply_lineage_colors(rendered_doc)
            
            # Final enforcement: prevent any cell/row expansion and force EXACT dimensions
            # Cell widths already standardized
            
            # CRITICAL: Remove ALL headers and footers to prevent unwanted content
            from src.core.generation.docx_formatting import remove_all_headers_and_footers
            rendered_doc = remove_all_headers_and_footers(rendered_doc)
            
            # Ensure proper table centering and document setup
            self._ensure_proper_centering(rendered_doc)

            # All content now uses standard spacing - no special THC_CBD handling
            
            chunk_time = time.time() - chunk_start_time
            self.logger.debug(f"Chunk processed in {chunk_time:.2f}s")
            
            # FINAL MARKER CLEANUP: Remove any lingering *_START and *_END markers AFTER font sizing has been applied
            # This cleanup should only remove markers that weren't processed by the font sizing system
            import re
            marker_pattern = re.compile(r'\b\w+_(START|END)\b')
            prefix_pattern = re.compile(r'^(?:[A-Z0-9_]+_)+')
            
            # Clean in tables
            try:
                for table in rendered_doc.tables:
                    try:
                        # Use safe table iteration to validate and repair if needed
                        if not self._safe_table_iteration(table, "marker cleanup"):
                            self.logger.warning(f"Skipping table with invalid structure during marker cleanup")
                            continue
                        
                        for row in table.rows:
                            try:
                                # Validate row structure before processing
                                if not hasattr(row, 'cells') or not row.cells:
                                    self.logger.warning(f"Skipping row with invalid structure during marker cleanup")
                                    continue
                                    
                                for cell in row.cells:
                                    try:
                                        for para in cell.paragraphs:
                                            # Check if this paragraph was processed by font sizing system
                                            # If it has non-default font sizes, it was processed
                                            was_processed = False
                                            for run in para.runs:
                                                if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                                                    # Check if font size is not the default (12pt)
                                                    if hasattr(run.font.size, 'pt') and run.font.size.pt != 12:
                                                        was_processed = True
                                                        break
                                            
                                            # Only clean markers if the paragraph wasn't processed by font sizing
                                            if not was_processed:
                                                for run in para.runs:
                                                    if marker_pattern.search(run.text):
                                                        run.text = marker_pattern.sub('', run.text)
                                                    if prefix_pattern.search(run.text):
                                                        run.text = prefix_pattern.sub('', run.text)
                                    except Exception as cell_error:
                                        self.logger.warning(f"Skipping cell due to error during marker cleanup: {cell_error}")
                                        continue
                            except Exception as row_error:
                                self.logger.warning(f"Skipping row due to error during marker cleanup: {row_error}")
                                continue
                    except Exception as e:
                        self.logger.warning(f"Skipping table due to error during marker cleanup: {e}")
                        continue
            except Exception as overall_error:
                self.logger.error(f"Critical error during marker cleanup: {overall_error}")
                # Continue processing other parts of the document
            
            # Clean in paragraphs outside tables
            for para in rendered_doc.paragraphs:
                # Check if this paragraph was processed by font sizing system
                was_processed = False
                for run in para.runs:
                    if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                        # Check if font size is not the default (12pt)
                        if hasattr(run.font.size, 'pt') and run.font.size.pt != 12:
                            was_processed = True
                            break
                
                # Only clean markers if the paragraph wasn't processed by font sizing
                if not was_processed:
                    for run in para.runs:
                        if marker_pattern.search(run.text):
                            run.text = marker_pattern.sub('', run.text)
                        if prefix_pattern.search(run.text):
                            run.text = prefix_pattern.sub('', run.text)
            
            return rendered_doc
        except Exception as e:
            self.logger.error(f"Error in _process_chunk: {e}\n{traceback.format_exc()}")
            raise

    def _build_inventory_context(self, record):
        """Build context dictionary for inventory slip template."""
        context = {}
        
        # Map inventory fields with proper formatting
        context['ProductName'] = record.get('Product Name*', '')
        context['Barcode'] = record.get('Barcode*', '')
        context['Quantity'] = record.get('Quantity Received*', '')
        context['AcceptedDate'] = record.get('Accepted Date', '')
        context['Vendor'] = record.get('Vendor', '')
        
        # Add any additional formatting or processing needed for inventory slips
        if context['AcceptedDate']:
            try:
                # Try to parse and reformat the date if needed
                from datetime import datetime
                date_obj = datetime.strptime(context['AcceptedDate'], '%Y-%m-%d')
                context['AcceptedDate'] = date_obj.strftime('%m/%d/%Y')
            except:
                pass  # Keep original format if parsing fails
        
        # Ensure all values are strings
        for key in context:
            if context[key] is None:
                context[key] = ''
            context[key] = str(context[key])
        
        return context

    def _build_label_context(self, record, doc):
        """Ultra-optimized label context building for maximum performance."""
        # Fast dictionary copy
        label_context = dict(record)

        # Fast value cleaning - only process non-empty values
        for key, value in label_context.items():
            if value is not None:
                label_context[key] = str(value).strip()
            else:
                label_context[key] = ""

        # Define product type sets for use throughout the method
        from src.core.constants import CLASSIC_TYPES
        classic_types = CLASSIC_TYPES
        edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}

        # Use DescAndWeight from record if it exists, otherwise construct it
        if 'DescAndWeight' in label_context and label_context['DescAndWeight']:
            # DescAndWeight is already set correctly in the record, use it as-is
            # Just ensure it's wrapped with the DESC marker if not already wrapped
            desc_and_weight = label_context['DescAndWeight']
            if not is_already_wrapped(desc_and_weight, 'DESC'):
                label_context['DescAndWeight'] = wrap_with_marker(desc_and_weight, 'DESC')
        else:
            # Fallback: construct DescAndWeight from Description and WeightUnits
            desc = label_context.get('Description', '') or ''
            weight = (label_context.get('WeightUnits', '') or '').replace('\u202F', '')
            
            # Ultra-fast string operations
            if desc.endswith('- '):
                desc = desc[:-2]
            if weight.startswith('- '):
                weight = weight[2:]
            
            # Template-specific handling
            if self.template_type == 'mini':
                # For mini templates, include both description and weight in DescAndWeight
                # Use regular space instead of non-breaking space to allow proper line breaking
                if desc and weight:
                    # Format for mini template with proper text wrapping
                    combined_text = f"{desc} - {weight}"
                    # Apply mini template text wrapping to prevent line breaks in the middle of words
                    wrapped_text = self._format_mini_template_text(combined_text)
                    label_context['DescAndWeight'] = wrap_with_marker(wrapped_text, 'DESC')
                else:
                    label_context['DescAndWeight'] = wrap_with_marker(desc or weight, 'DESC')
            else:
                if desc and weight:
                    combined_text = f"{desc} -\u00A0{weight}"
                    label_context['DescAndWeight'] = wrap_with_marker(combined_text, 'DESC')
                else:
                    label_context['DescAndWeight'] = wrap_with_marker(desc or weight, 'DESC')

        # Fast DOH image processing - only if needed
        if label_context.get('DOH'):
            doh_value = label_context.get('DOH', '')
            product_type = (label_context.get('ProductType') or 
                          label_context.get('Product Type*') or 
                          record.get('ProductType') or 
                          record.get('Product Type*') or '')

            image_path = process_doh_image(doh_value, product_type)
            if image_path:
                # Fast width selection - reduced by 1mm for all template types
                width_map = {'mini': 8, 'double': 7, 'vertical': 11, 'horizontal': 11}
                image_width = Mm(width_map.get(self.template_type, 11))
                label_context['DOH'] = InlineImage(doc, image_path, width=image_width)
                # Ensure DOH image takes priority - clear any other DOH-related content
                label_context['DOH_TEXT'] = ''  # Clear any text content
            else:
                label_context['DOH'] = ''
                label_context['DOH_TEXT'] = ''
        else:
            label_context['DOH'] = ''
            label_context['DOH_TEXT'] = ''
        
        # CRITICAL: Lineage and ProductVendor logic for classic types
        # This implements the same logic that was in tag_generator
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        product_brand = label_context.get('ProductBrand') or label_context.get('Product Brand', '')
        lineage_text = label_context.get('Lineage', '')
        product_strain = label_context.get('ProductStrain') or label_context.get('Product Strain', '')
        
        # Check if it's a classic type
        is_classic_type = product_type in classic_types
        
        if is_classic_type:
            # For classic types, Lineage should show strain lineage and ProductVendor should show brand
            self.logger.debug(f"Processing classic type '{product_type}' for Lineage and ProductVendor")
            
            # Try to get lineage from database first, then fall back to Excel
            lineage_val = ""
            if product_strain:
                try:
                    from src.core.data.product_database import get_product_database
                    product_db = get_product_database()
                    strain_info = product_db.get_strain_info(product_strain)
                    if strain_info and strain_info.get('canonical_lineage'):
                        lineage_val = strain_info['canonical_lineage'].upper()
                        self.logger.debug(f"Using database lineage: '{lineage_val}'")
                    else:
                        # Fallback to Excel lineage
                        lineage_val = lineage_text.upper() if lineage_text else ""
                        self.logger.debug(f"Using Excel lineage fallback: '{lineage_val}'")
                except Exception as e:
                    # Fallback to Excel lineage if database lookup fails
                    lineage_val = lineage_text.upper() if lineage_text else ""
                    self.logger.debug(f"Using Excel lineage due to error: '{lineage_val}' (error: {e})")
            else:
                # No strain available, try Excel lineage
                lineage_val = lineage_text.upper() if lineage_text else ""
                self.logger.debug(f"No strain available, using Excel lineage: '{lineage_val}'")
            
            # Set Lineage to strain lineage for classic types
            if lineage_val:
                # Debug: Log the lineage value to see if it has leading spaces
                self.logger.debug(f"DEBUG: Original lineage_val: '{repr(lineage_val)}'")
                cleaned_lineage_val = lineage_val.strip()
                self.logger.debug(f"DEBUG: Cleaned lineage_val: '{repr(cleaned_lineage_val)}'")
                
                label_context['Lineage'] = f"LINEAGE_START{cleaned_lineage_val}LINEAGE_END"
                self.logger.debug(f"Set Lineage to strain lineage: '{cleaned_lineage_val}' for classic type '{product_type}'")
            else:
                label_context['Lineage'] = ""
                self.logger.debug(f"No lineage available for classic type '{product_type}', Lineage set to empty")
            
            # Set ProductVendor to actual vendor/supplier for classic types
            # Get vendor from record, not from product_brand
            vendor_val = record.get('Vendor') or record.get('Vendor/Supplier*') or record.get('ProductVendor', '')
            if vendor_val and str(vendor_val).lower() != 'nan':
                label_context['ProductVendor'] = f"PRODUCTVENDOR_START{str(vendor_val)}PRODUCTVENDOR_END"
                self.logger.debug(f"Set ProductVendor to vendor: '{vendor_val}' for classic type '{product_type}'")
            else:
                label_context['ProductVendor'] = ""
                self.logger.debug(f"ProductVendor set to empty for classic type '{product_type}' (no vendor data)")
        else:
            # For non-classic types, Lineage shows brand and ProductVendor is empty
            self.logger.debug(f"Processing non-classic type '{product_type}' for Lineage and ProductVendor")
            if product_brand:
                # For non-classic types, wrap brand with PRODUCTBRAND markers so it gets centered
                label_context['Lineage'] = f"PRODUCTBRAND_START{product_brand}PRODUCTBRAND_END"
                self.logger.debug(f"Set Lineage to brand with PRODUCTBRAND markers: '{product_brand}' for non-classic type '{product_type}'")
            else:
                label_context['Lineage'] = ""
                self.logger.debug(f"Lineage set to empty for non-classic type '{product_type}'")
            label_context['ProductVendor'] = ""
        
        # Fast ratio processing
        ratio_val = label_context.get('Ratio_or_THC_CBD') or label_context.get('Ratio', '')
        
        # If no ratio value but we have individual THC/CBD values, combine them
        # PRIORITY: Excel THC/CBD values are used directly (from Excel sheet processing)
        if not ratio_val:
            thc_val = label_context.get('THC', '')  # Excel THC value
            cbd_val = label_context.get('CBD', '')  # Excel CBD value
            if thc_val or cbd_val:
                # Combine THC and CBD values into ratio format
                ratio_parts = []
                if thc_val:
                    ratio_parts.append(f"THC: {thc_val}%")
                if cbd_val:
                    ratio_parts.append(f"CBD: {cbd_val}%")
                ratio_val = " ".join(ratio_parts)
                
                # Apply the new THC/CBD percentage formatting rules
                from src.core.generation.text_processing import format_thc_cbd_percentages
                ratio_val = format_thc_cbd_percentages(ratio_val)
                
                self.logger.debug(f"Combined Excel THC/CBD values into ratio: '{ratio_val}'")
        
        if ratio_val:
            cleaned_ratio = ratio_val.lstrip('- ')
            product_type = (label_context.get('ProductType', '').lower() or 
                          label_context.get('Product Type*', '').lower())
            
            # Fast product type checking with sets
            is_classic = product_type in classic_types
            is_edible = product_type in edible_types
            
            # Check if this is percentage-based THC/CBD content (not mg-based)
            is_percentage_based = '%' in cleaned_ratio and ('THC:' in cleaned_ratio or 'CBD:' in cleaned_ratio)
            
            if is_classic and 'mg' in cleaned_ratio.lower() and not is_percentage_based:
                cleaned_ratio = format_ratio_multiline(cleaned_ratio)
            elif is_edible and 'mg' in cleaned_ratio.lower() and not is_percentage_based:
                cleaned_ratio = format_ratio_multiline(cleaned_ratio)
            elif is_classic and not is_percentage_based:
                cleaned_ratio = self.format_classic_ratio(cleaned_ratio, record)
            
            # Apply THC/CBD percentage formatting to any existing ratio content
            if cleaned_ratio and ('THC:' in cleaned_ratio or 'CBD:' in cleaned_ratio):
                from src.core.generation.text_processing import format_thc_cbd_percentages
                cleaned_ratio = format_thc_cbd_percentages(cleaned_ratio)
            
            # Fast marker wrapping
            content = cleaned_ratio.replace('|BR|', '\n')
            # Apply new bold label formatting for THC/CBD content
            if content.strip().startswith('THC:') and 'CBD:' in content:
                from src.core.generation.text_processing import format_thc_cbd_bold_labels
                content = format_thc_cbd_bold_labels(content, self.template_type)
            # Force line breaks for vertical, double, and mini templates
            elif self.template_type in ['vertical', 'double', 'mini'] and content.strip().startswith('THC:') and 'CBD:' in content:
                content = content.replace('THC: CBD:', 'THC:\nCBD:').replace('THC:  CBD:', 'THC:\nCBD:')
                # For vertical template, format with right-aligned percentages
                if self.template_type == 'vertical':
                    content = self.format_thc_cbd_vertical_alignment(content)
            
            marker = 'THC_CBD' if is_classic else 'RATIO'
            label_context['Ratio_or_THC_CBD'] = wrap_with_marker(content, marker)
            
            # Also add separate THC_CBD field for template processing
            if is_classic:
                label_context['THC_CBD'] = wrap_with_marker(content, 'THC_CBD')
            else:
                label_context['THC_CBD'] = ''
        else:
            label_context['Ratio_or_THC_CBD'] = ''
            label_context['THC_CBD'] = ''

        # Lineage and ProductVendor logic is now handled earlier in the method for classic types

        # Fast other field processing
        if label_context.get('Price'):
            label_context['Price'] = wrap_with_marker(unwrap_marker(label_context['Price'], 'PRICE'), 'PRICE')
        
        # Always process lineage for classic types, and conditionally for non-classic types
        product_type = (label_context.get('Product Type*', '').lower() or 
                       label_context.get('ProductType', '').lower())
        product_strain = record.get('ProductStrain') or record.get('Product Strain', '')
        
        # For classic types, ALWAYS try to get the strain's canonical lineage from the database
        if is_classic_type and product_strain:
            self.logger.debug(f"DEBUG: Processing classic type '{product_type}' with strain '{product_strain}'")
            try:
                from src.core.data.product_database import get_product_database
                product_db = get_product_database()
                strain_info = product_db.get_strain_info(product_strain)
                self.logger.debug(f"DEBUG: Strain info: {strain_info}")
                if strain_info and strain_info.get('canonical_lineage'):
                    lineage_value = strain_info['canonical_lineage'].upper()
                    self.logger.debug(f"DEBUG: Using database lineage: '{lineage_value}'")
                else:
                    # Fallback to Excel lineage if no database lineage found
                    lineage_value = label_context.get('Lineage', '')
                    self.logger.debug(f"DEBUG: Using Excel lineage fallback: '{lineage_value}'")
            except Exception as e:
                # Fallback to Excel lineage if database lookup fails
                lineage_value = label_context.get('Lineage', '')
                self.logger.debug(f"DEBUG: Using Excel lineage due to error: '{lineage_value}' (error: {e})")
            
            # Lineage logic is now handled earlier in the method for both classic and non-classic types
            
        # Lineage logic is now handled earlier in the method for both classic and non-classic types

        # Fast wrapping for remaining fields
        if label_context.get('DescAndWeight'):
            label_context['DescAndWeight'] = wrap_with_marker(unwrap_marker(label_context['DescAndWeight'], 'DESC'), 'DESC')
        
        if 'ProductType' not in label_context:
            label_context['ProductType'] = record.get('ProductType', '')
        
        # Fast strain handling - always show the actual strain value from Excel
        product_strain = record.get('ProductStrain') or record.get('Product Strain', '')
        
        if product_strain:
            # Always show the actual strain value from the Excel column
            # This ensures the ProductStrain placeholder displays the intended strain information
            label_context['ProductStrain'] = wrap_with_marker(product_strain, 'PRODUCTSTRAIN')
        else:
            label_context['ProductStrain'] = ''

        # Lineage logic is now handled earlier in the method for both classic and non-classic types

        # Add marker strings for template processing
        # These markers will be rendered by DocxTemplate and preserved for font sizing
        label_context['ProductStrain_START'] = 'PRODUCTSTRAIN_START'
        label_context['ProductStrain_END'] = 'PRODUCTSTRAIN_END'
        # Add Lineage markers back for post-processing system to work
        label_context['Lineage_START'] = 'LINEAGE_START'
        label_context['Lineage_END'] = 'LINEAGE_END'
        label_context['ProductBrand_START'] = 'PRODUCTBRAND_START'
        label_context['ProductBrand_END'] = 'PRODUCTBRAND_END'
        label_context['ProductVendor_START'] = 'PRODUCTVENDOR_START'
        label_context['ProductVendor_END'] = 'PRODUCTVENDOR_END'
        label_context['DescAndWeight_START'] = 'DESC_START'
        label_context['DescAndWeight_END'] = 'DESC_END'
        label_context['Ratio_or_THC_CBD_START'] = 'THC_CBD_START'
        label_context['Ratio_or_THC_CBD_END'] = 'THC_CBD_END'
        label_context['Price_START'] = 'PRICE_START'
        label_context['Price_END'] = 'PRICE_END'
        label_context['WeightUnits_START'] = 'WEIGHTUNITS_START'
        label_context['WeightUnits_END'] = 'WEIGHTUNITS_END'
        label_context['Ratio_START'] = 'RATIO_START'
        label_context['Ratio_END'] = 'RATIO_END'
        label_context['JointRatio_START'] = 'JOINT_RATIO_START'
        label_context['JointRatio_END'] = 'JOINT_RATIO_END'
        label_context['THC_START'] = 'THC_START'
        label_context['THC_END'] = 'THC_END'
        label_context['CBD_START'] = 'CBD_START'
        label_context['CBD_END'] = 'CBD_END'

        # Fast joint ratio handling
        if label_context.get('JointRatio'):
            val = label_context['JointRatio']
            # Fix: Handle NaN values in JointRatio
            if pd.isna(val) or str(val).lower() == 'nan':
                val = ''
            marker = 'JOINT_RATIO'
            if is_already_wrapped(val, marker):
                val = unwrap_marker(val, marker)
            formatted_val = self.format_joint_ratio_pack(val)
            label_context['JointRatio'] = wrap_with_marker(formatted_val, marker)

        # Fast description processing
        if label_context.get('Description'):
            label_context['Description'] = self.fix_hyphen_spacing(label_context['Description'])

        # Fast line break processing
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        
        if product_type not in classic_types and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            if desc_weight.endswith(' - '):
                desc_weight = desc_weight[:-3] + '\n- '
            elif desc_weight.endswith(' -'):
                desc_weight = desc_weight[:-2] + '\n- '
            desc_weight = desc_weight.replace(' - ', '\n- ')
            label_context['DescAndWeight'] = desc_weight
        
        # Fast pre-roll processing
        if product_type in {"pre-roll", "infused pre-roll"} and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            desc_weight = desc_weight.replace(' - ', '\n- ')
            label_context['DescAndWeight'] = desc_weight

        # Fast weight and ratio formatting
        for key, marker in [('WeightUnits', 'WEIGHTUNITS'), ('Ratio', 'RATIO')]:
            if label_context.get(key):
                val = label_context[key]
                formatted_val = self.format_with_soft_hyphen(val)
                label_context[key] = wrap_with_marker(unwrap_marker(formatted_val, marker), marker)
        
        # Fast vendor handling - only override if ProductVendor wasn't already set by our logic
        # This preserves the ProductVendor logic for classic types
        if 'ProductVendor' not in label_context:
            product_type = (label_context.get('ProductType', '').lower() or 
                           label_context.get('Product Type*', '').lower())
            
            # Only set vendor from record if ProductVendor wasn't already set by our logic
            product_vendor = record.get('Vendor') or record.get('Vendor/Supplier*', '') or record.get('ProductVendor', '')
            # Handle NaN values in vendor data
            if pd.isna(product_vendor) or str(product_vendor).lower() == 'nan':
                product_vendor = ''
            label_context['ProductVendor'] = wrap_with_marker(product_vendor, 'PRODUCTVENDOR')

        # Generate QR code for Product Name
        product_name = label_context.get('Product Name*') or label_context.get('ProductName') or label_context.get('Product Name', '')
        if product_name and str(product_name).strip():
            qr_code = self._generate_qr_code(product_name, doc)
            if qr_code:
                label_context['QR'] = qr_code
                self.logger.debug(f"Generated QR code for product: '{product_name}'")
            else:
                label_context['QR'] = ''
                self.logger.warning(f"Failed to generate QR code for product: '{product_name}'")
        else:
            label_context['QR'] = ''
            self.logger.debug("No product name available for QR code generation")

        return label_context

    def _generate_qr_code(self, product_name, doc):
        """Generate QR code for the given product name and return as InlineImage."""
        try:
            if not product_name or str(product_name).strip() == '':
                self.logger.warning("Empty product name provided for QR code generation")
                return None
            
            # Clean the product name
            clean_name = str(product_name).strip()
            
            # Create QR code instance
            qr = qrcode.QRCode(
                version=1,  # Auto-determine version based on content
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,  # Size of each box in pixels
                border=4,     # Border size in boxes
            )
            
            # Add data to QR code
            qr.add_data(clean_name)
            qr.make(fit=True)
            
            # Create QR code image
            qr_image = qr.make_image(fill_color="black", back_color="white")
            
            # Convert to BytesIO for InlineImage
            img_buffer = BytesIO()
            qr_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            # Determine QR code size based on template type
            # Smaller QR codes for mini templates, larger for others
            size_map = {
                'mini': 15,      # 15mm for mini templates
                'double': 18,    # 18mm for double templates  
                'vertical': 20,  # 20mm for vertical templates
                'horizontal': 20, # 20mm for horizontal templates
                'inventory': 25  # 25mm for inventory templates
            }
            
            qr_size = Mm(size_map.get(self.template_type, 20))
            
            # Create InlineImage for the document
            qr_inline_image = InlineImage(doc, img_buffer, width=qr_size)
            
            # Store the raw image data for manual replacement
            img_buffer.seek(0)  # Reset buffer position
            qr_inline_image._raw_image_data = img_buffer.read()
            qr_inline_image._raw_image_width = qr_size
            
            self.logger.debug(f"Generated QR code for product: '{clean_name}' with size: {qr_size}")
            return qr_inline_image
            
        except Exception as e:
            self.logger.error(f"Error generating QR code for product '{product_name}': {e}")
            return None

    def _post_process_and_replace_content(self, doc):
        """Post-process the document after template rendering."""
        # Skip unnecessary processing for inventory templates
        if self.template_type == 'inventory':
            self.logger.info("Skipping post-processing for inventory template - just filling placeholders")
            return doc
        """
        Ultra-optimized post-processing for maximum performance.
        """
        # Performance optimization: Skip expensive processing for large documents
        if len(doc.tables) > 10:
            self.logger.warning(f"Skipping expensive post-processing for large document with {len(doc.tables)} tables")
            return doc
        
        # Clean up DOH cells before processing to ensure proper image positioning
        try:
            self._clean_doh_cells_before_processing(doc)
        except Exception as e:
            self.logger.warning(f"DOH cell cleanup failed: {e}")
        
        # Enhanced mini template processing
        if self.template_type == 'mini':
            try:
                self.logger.info("Processing mini template with enhanced design preservation")
                
                # Add markers for proper processing
                self._add_weight_units_markers(doc)
                self._add_brand_markers(doc)
                
                # Ensure proper brand centering for mini templates
                self._ensure_mini_template_brand_centering(doc)
                
                # Clear blank cells that don't have meaningful content
                self._clear_blank_cells_in_mini_template(doc)
                
                # CRITICAL: Enforce fixed cell dimensions to maintain 1.5" x 1.5" cells
                for table in doc.tables:
                    enforce_fixed_cell_dimensions(table, 'mini')
                    self.logger.info("Applied fixed cell dimensions to mini template table")
                
                # Apply mini template specific font sizing
                self._apply_mini_template_font_sizing(doc)
                    
            except Exception as e:
                self.logger.warning(f"Mini template processing failed: {e}")
                # Continue processing even if mini-specific steps fail

        # ProductStrain in Brand cells fix
        try:
            self._fix_productstrain_in_brand_cells(doc)
        except Exception as e:
            self.logger.warning(f"ProductStrain in Brand cells fix failed: {e}")
        
        # Fast double template processing
        if self.template_type == 'double':
            try:
                self._add_brand_markers(doc)
                self._apply_brand_centering_for_double_template(doc)
            except Exception as e:
                self.logger.warning(f"Double template processing failed: {e}")

        # Fast font sizing (with timeout protection)
        try:
            self._post_process_template_specific(doc)
        except Exception as e:
            self.logger.warning(f"Font sizing failed: {e}")

        # Fast BR marker conversion - only process if needed
        try:
            br_found = False
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during BR marker conversion")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '|BR|' in paragraph.text:
                                self._convert_br_markers_to_line_breaks(paragraph)
                                br_found = True
            
            # Only process paragraphs outside tables if BR markers were found
            if br_found:
                for paragraph in doc.paragraphs:
                    if '|BR|' in paragraph.text:
                        self._convert_br_markers_to_line_breaks(paragraph)
        except Exception as e:
            self.logger.warning(f"BR marker conversion failed: {e}")
        
        # Fast ratio spacing fix
        try:
            self._fix_ratio_paragraph_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Ratio spacing failed: {e}")

        # Ensure consistent spacing above lineage/brand section for equal margins
        try:
            self._ensure_consistent_lineage_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Lineage spacing consistency failed: {e}")

        # Add consistent spacing above main content sections for better visual balance
        try:
            self._add_consistent_content_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Content spacing consistency failed: {e}")

        # Fast Arial Bold enforcement
        try:
            from src.core.generation.docx_formatting import enforce_arial_bold_all_text, enforce_ratio_formatting, enforce_thc_cbd_bold_formatting
            enforce_arial_bold_all_text(doc)
            enforce_ratio_formatting(doc)
            enforce_thc_cbd_bold_formatting(doc)
        except Exception as e:
            self.logger.warning(f"Arial bold failed: {e}")

        # Fast DOH image centering
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during DOH centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Fast check for image-only cells
                        if len(cell.paragraphs) > 0 and all(len(paragraph.runs) == 1 and not paragraph.text.strip() for paragraph in cell.paragraphs):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Fast inner table centering
                        for inner_table in cell.tables:
                            inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        # Explicit DOH image centering - check for InlineImage objects
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Check if this run contains an InlineImage (DOH image)
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    # Also center the cell content
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    
            # Additional comprehensive DOH centering pass
            self._ensure_doh_image_centering(doc)
            
            # CRITICAL FIX: Ensure DOH images are properly centered in top-section cells
            self._fix_doh_image_positioning_in_top_section(doc)
            
            # CRITICAL FIX: Ensure DOH images have proper vertical margins to prevent cutoff
            self._ensure_doh_logo_vertical_margins(doc)
            
            # CRITICAL FIX: Final marker cleanup to ensure ALL markers are stripped
            self._final_marker_cleanup(doc)
            
            # FINAL ENFORCEMENT: Absolutely ensure DOH images are centered - this overrides all other positioning
            self._final_doh_positioning_enforcement(doc)
        except Exception as e:
            self.logger.warning(f"DOH centering failed: {e}")
        
        # CRITICAL: Final cell dimension enforcement to prevent any expansion
        try:
            for table in doc.tables:
                enforce_fixed_cell_dimensions(table, self.template_type)
                self.logger.info(f"Applied final fixed cell dimensions to {self.template_type} template table")
        except Exception as e:
            self.logger.warning(f"Final cell dimension enforcement failed: {e}")
            
        return doc

    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method provides improved centering for InlineImage objects.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during DOH image centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_paragraph = None
                        
                        # Improved image detection
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element'):
                                    # Check for drawing elements (InlineImage)
                                    if run._element.find(qn('w:drawing')) is not None:
                                        has_doh_image = True
                                        image_paragraph = paragraph
                                        break
                                    # Check for picture elements
                                    elif run._element.find(qn('w:pict')) is not None:
                                        has_doh_image = True
                                        image_paragraph = paragraph
                                        break
                            if has_doh_image:
                                break
                        
                        if has_doh_image and image_paragraph:
                            self.logger.debug("Found DOH image, applying improved centering")
                            
                            # Apply centering at paragraph level
                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set proper spacing to prevent DOH logo from being cut off
                            image_paragraph.paragraph_format.space_before = Pt(3)
                            image_paragraph.paragraph_format.space_after = Pt(3)
                            image_paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure proper XML-level centering
                            pPr = image_paragraph._element.get_or_add_pPr()
                            
                            # Set paragraph justification to center
                            jc = pPr.find(qn('w:jc'))
                            if jc is None:
                                jc = OxmlElement('w:jc')
                                pPr.append(jc)
                            jc.set(qn('w:val'), 'center')
                            
                            # Remove any existing spacing
                            existing_spacing = pPr.find(qn('w:spacing'))
                            if existing_spacing is not None:
                                pPr.remove(existing_spacing)
                            
                            # Add proper spacing to prevent DOH logo from being cut off
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:before'), '60')  # 3pt = 60 twips
                            spacing.set(qn('w:after'), '60')   # 3pt = 60 twips
                            spacing.set(qn('w:line'), '240')
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            
                            # Ensure proper indentation
                            ind = pPr.find(qn('w:ind'))
                            if ind is None:
                                ind = OxmlElement('w:ind')
                                pPr.append(ind)
                            ind.set(qn('w:left'), '0')
                            ind.set(qn('w:right'), '0')
                            ind.set(qn('w:firstLine'), '0')
                            ind.set(qn('w:hanging'), '0')
                            
                            # CRITICAL FIX: Ensure the image itself is centered within the cell
                            # Check if this is a top-section DOH image (with other content like "100mg THC")
                            cell_text = cell.text.strip()
                            if '100mg THC' in cell_text or '$' in cell_text:
                                # This is a top-section cell with multiple elements
                                # Ensure the DOH image is centered in its own paragraph
                                self.logger.debug("Found top-section DOH image, ensuring proper centering")
                                
                                # Create a dedicated centered paragraph for the DOH image
                                if len(cell.paragraphs) > 1:
                                    # Find the paragraph with the DOH image
                                    for para in cell.paragraphs:
                                        if any(run._element.find(qn('w:drawing')) is not None or 
                                               run._element.find(qn('w:pict')) is not None 
                                               for run in para.runs):
                                            # This is the DOH image paragraph
                                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            # Set XML-level centering
                                            pPr = para._element.get_or_add_pPr()
                                            jc = pPr.find(qn('w:jc'))
                                            if jc is None:
                                                jc = OxmlElement('w:jc')
                                                pPr.append(jc)
                                            jc.set(qn('w:val'), 'center')
                                            
                                            # Remove any indentation that might affect centering
                                            ind = pPr.find(qn('w:ind'))
                                            if ind is not None:
                                                pPr.remove(ind)
                                            
                                            self.logger.debug("Applied top-section DOH image centering")
                                            break
                            
                            self.logger.debug("Applied improved DOH image centering")
                                
        except Exception as e:
            self.logger.warning(f"Error in improved DOH image centering: {e}")

    def _fix_doh_image_positioning_in_top_section(self, doc):
        """
        Fix DOH image positioning in top-section cells that contain multiple elements.
        This ensures DOH images are properly centered even when in cells with other content.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this is a top-section cell with multiple elements
                        cell_text = cell.text.strip()
                        if ('100mg THC' in cell_text or '$' in cell_text) and len(cell.paragraphs) > 1:
                            self.logger.debug("Found top-section cell with multiple elements, fixing DOH positioning")
                            
                            # Find the paragraph with the DOH image
                            for paragraph in cell.paragraphs:
                                has_doh_image = False
                                for run in paragraph.runs:
                                    if hasattr(run, '_element'):
                                        if (run._element.find(qn('w:drawing')) is not None or 
                                            run._element.find(qn('w:pict')) is not None):
                                            has_doh_image = True
                                            break
                                
                                if has_doh_image:
                                    # This is the DOH image paragraph - ensure it's centered
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Set XML-level centering
                                    pPr = paragraph._element.get_or_add_pPr()
                                    jc = pPr.find(qn('w:jc'))
                                    if jc is None:
                                        jc = OxmlElement('w:jc')
                                        pPr.append(jc)
                                    jc.set(qn('w:val'), 'center')
                                    
                                    # Remove any indentation that might affect centering
                                    ind = pPr.find(qn('w:ind'))
                                    if ind is not None:
                                        pPr.remove(ind)
                                    
                                    # Ensure proper spacing to prevent DOH logo from being cut off
                                    spacing = pPr.find(qn('w:spacing'))
                                    if spacing is None:
                                        spacing = OxmlElement('w:spacing')
                                        pPr.append(spacing)
                                    spacing.set(qn('w:before'), '60')  # 3pt = 60 twips
                                    spacing.set(qn('w:after'), '60')   # 3pt = 60 twips
                                    spacing.set(qn('w:line'), '240')
                                    spacing.set(qn('w:lineRule'), 'auto')
                                    
                                    self.logger.debug("Fixed DOH image positioning in top-section cell")
                                    break
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
        except Exception as e:
            self.logger.warning(f"Error fixing DOH image positioning in top-section: {e}")

    def _ensure_doh_logo_vertical_margins(self, doc):
        """
        Ensure DOH logos have proper vertical margins to prevent cutoff at the top.
        This method specifically targets the vertical spacing issue that causes logo clipping.
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element'):
                                    if (run._element.find(qn('w:drawing')) is not None or 
                                        run._element.find(qn('w:pict')) is not None):
                                        has_doh_image = True
                                        break
                            if has_doh_image:
                                break
                        
                        if has_doh_image:
                            # Apply generous vertical margins to prevent DOH logo cutoff
                            for paragraph in cell.paragraphs:
                                # Set generous spacing above and below
                                paragraph.paragraph_format.space_before = Pt(4)
                                paragraph.paragraph_format.space_after = Pt(4)
                                
                                # Set XML-level spacing for maximum compatibility
                                pPr = paragraph._element.get_or_add_pPr()
                                spacing = pPr.find(qn('w:spacing'))
                                if spacing is None:
                                    spacing = OxmlElement('w:spacing')
                                    pPr.append(spacing)
                                
                                # Set generous margins: 4pt = 80 twips
                                spacing.set(qn('w:before'), '80')
                                spacing.set(qn('w:after'), '80')
                                spacing.set(qn('w:line'), '240')
                                spacing.set(qn('w:lineRule'), 'auto')
                                
                                # Ensure no indentation interferes with spacing
                                ind = pPr.find(qn('w:ind'))
                                if ind is not None:
                                    pPr.remove(ind)
                            
                            self.logger.debug("Applied generous vertical margins to prevent DOH logo cutoff")
                            break
                            
        except Exception as e:
            self.logger.warning(f"Error ensuring DOH logo vertical margins: {e}")

    def _final_marker_cleanup(self, doc):
        """
        Final marker cleanup to ensure ALL markers are stripped from the final output.
        This method runs after all other processing to catch any remaining markers.
        """
        try:
            import re
            
            # Enhanced patterns to catch all marker variations
            marker_patterns = [
                r'\b\w+_(START|END)\b',           # Standard markers like PRODUCTBRAND_START
                r'\b\w+_START\b',                 # START markers specifically
                r'\b\w+_END\b',                   # END markers specifically
                r'PRODUCTBRAND_START\s*',         # PRODUCTBRAND_START with optional spaces
                r'\s*PRODUCTBRAND_END\b',         # PRODUCTBRAND_END with optional spaces
                r'PRODUCTSTRAIN_START\s*',        # PRODUCTSTRAIN_START with optional spaces
                r'\s*PRODUCTSTRAIN_END\b',        # PRODUCTSTRAIN_END with optional spaces
                r'LINEAGE_START\s*',              # LINEAGE_START with optional spaces
                r'\s*LINEAGE_END\b',              # LINEAGE_END with optional spaces
                r'PRODUCTVENDOR_START\s*',        # PRODUCTVENDOR_START with optional spaces
                r'\s*PRODUCTVENDOR_END\b',        # PRODUCTVENDOR_END with optional spaces
                r'\bPRODUCTBRAND\b',              # Standalone PRODUCTBRAND
                r'\bPRODUCTSTRAIN\b',             # Standalone PRODUCTSTRAIN
                r'\bLINEAGE\b',                   # Standalone LINEAGE
                r'\bPRODUCTVENDOR\b',             # Standalone PRODUCTVENDOR
            ]
            
            def clean_text(text):
                """Clean text by removing all marker patterns."""
                cleaned = text
                for pattern in marker_patterns:
                    cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
                
                # CRITICAL FIX: Remove partial marker remnants like "bis" from "PRODUCTBRAND_END"
                partial_remnants = [
                    r'\bbis\b',                    # "bis" from PRODUCTBRAND_END
                    r'\bSTART\b',                  # Any remaining START
                    r'\bEND\b',                    # Any remaining END
                    r'\bPRODUCT\b',                # Any remaining PRODUCT
                    r'\bBRAND\b',                  # Any remaining BRAND
                    r'\bSTRAIN\b',                 # Any remaining STRAIN
                    r'\bVENDOR\b',                 # Any remaining VENDOR
                    r'\bLINEAGE\b',                # Any remaining LINEAGE
                ]
                
                for remnant in partial_remnants:
                    cleaned = re.sub(remnant, '', cleaned, flags=re.IGNORECASE)
                
                # Clean up any double spaces, leading/trailing spaces
                cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                return cleaned
            
            # Clean markers in all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                original_text = run.text
                                cleaned_text = clean_text(original_text)
                                if cleaned_text != original_text:
                                    run.text = cleaned_text
                                    self.logger.debug(f"Cleaned markers from run: '{original_text}' -> '{cleaned_text}'")
            
            # Clean markers in paragraphs outside tables
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original_text = run.text
                    cleaned_text = clean_text(original_text)
                    if cleaned_text != original_text:
                        run.text = cleaned_text
                        self.logger.debug(f"Cleaned markers from paragraph run: '{original_text}' -> '{cleaned_text}'")
            
            # FINAL LINEAGE CLEANUP: Remove any leading spaces from lineage content
            self._final_lineage_cleanup(doc)
            
            self.logger.debug("Enhanced final marker cleanup completed - all markers should be stripped")
            
        except Exception as e:
            self.logger.warning(f"Error in enhanced final marker cleanup: {e}")

    def _final_lineage_cleanup(self, doc):
        """
        Final cleanup to remove any leading spaces from lineage content.
        This runs after all other processing to ensure clean lineage display.
        """
        try:
            # Define lineage values that should be cleaned
            lineage_values = [
                "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                "CBD", "MIXED", "PARAPHERNALIA", "PARA"
            ]
            
            # Clean lineage content in all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                original_text = run.text
                                
                                # Check if this run contains lineage content
                                for lineage in lineage_values:
                                    if lineage in original_text.upper():
                                        # Aggressively clean leading spaces
                                        cleaned_text = original_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                                        
                                        if cleaned_text != original_text:
                                            run.text = cleaned_text
                                            self.logger.debug(f"Final lineage cleanup: '{original_text}' -> '{cleaned_text}'")
                                        break
            
            # Clean lineage content in paragraphs outside tables
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original_text = run.text
                    
                    # Check if this run contains lineage content
                    for lineage in lineage_values:
                        if lineage in original_text.upper():
                            # Aggressively clean leading spaces
                            cleaned_text = original_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                            
                            if cleaned_text != original_text:
                                run.text = cleaned_text
                                self.logger.debug(f"Final lineage cleanup: '{original_text}' -> '{cleaned_text}'")
                            break
            
            self.logger.debug("Final lineage cleanup completed - all leading spaces should be removed")
            
        except Exception as e:
            self.logger.warning(f"Error in final lineage cleanup: {e}")

    def _clear_blank_cells_in_mini_template(self, doc):
        """
        Clear blank cells in mini templates when they run out of values.
        This removes empty cells that don't have any meaningful content.
        """
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during blank cell clearing")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Check if cell is essentially empty
                        cell_text = cell.text.strip()
                        
                        # Consider a cell blank if it has no text or only contains template placeholders
                        is_blank = (
                            not cell_text or 
                            cell_text == '' or
                            # Check for empty template placeholders like {{LabelX.Description}}
                            (cell_text.startswith('{{Label') and cell_text.endswith('}}') and 
                             any(field in cell_text for field in ['.Description}}', '.Price}}', '.Lineage}}', '.ProductBrand}}', '.Ratio_or_THC_CBD}}', '.DOH}}', '.ProductStrain}}']))
                        )
                        
                        if is_blank:
                            # Clear the cell content
                            cell._tc.clear_content()
                            
                            # Add a single empty paragraph to maintain cell structure
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set cell background to white/transparent to ensure it's visually clean
                            from src.core.generation.docx_formatting import clear_cell_background
                            clear_cell_background(cell)
                            
                            self.logger.debug(f"Cleared blank cell in mini template")
                            
        except Exception as e:
            self.logger.error(f"Error clearing blank cells in mini template: {e}")
            # Don't raise the exception - this is a cleanup operation that shouldn't break the main process

    def _post_process_template_specific(self, doc):
        """
        Apply template-type-specific font sizing to all markers in the document.
        Uses the original font-sizing functions based on template type.
        """
        # Define marker processing for all template types (including double)
        markers = [
            'DESC', 'PRODUCTBRAND', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 
            'THC_CBD', 'THC_CBD_LABEL', 'RATIO', 'WEIGHTUNITS', 'PRODUCTSTRAIN', 'DOH', 'PRODUCTVENDOR'
        ]
        
        # Process all markers in a single pass to avoid conflicts
        self._recursive_autosize_template_specific_multi(doc, markers)
        
        # Apply vertical template specific optimizations for minimal spacing
        if self.template_type in ['vertical', 'double']:
            self._optimize_vertical_template_spacing(doc)



    def _optimize_vertical_template_spacing(self, doc):
        """
        Apply minimal spacing optimizations specifically for vertical and double templates
        to ensure all labels fit on one page.
        """
        try:
            from docx.shared import Pt
            
            def optimize_paragraph_spacing(paragraph):
                """Set minimal spacing for all paragraphs in vertical and double templates."""
                # Set absolute minimum spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
                # All content now uses standard spacing
                
                # Default spacing for non-THC_CBD content
                paragraph.paragraph_format.line_spacing = 1.0
                
                # Set at XML level for maximum compatibility
                pPr = paragraph._element.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            optimize_paragraph_spacing(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                optimize_paragraph_spacing(paragraph)
            
            self.logger.debug("Applied vertical/double template spacing optimizations")
            
        except Exception as e:
            self.logger.error(f"Error optimizing vertical/double template spacing: {e}")
            # Don't raise the exception - this is an optimization that shouldn't break the main process

    def _recursive_autosize_template_specific(self, element, marker_name):
        """
        Recursively find and replace markers in paragraphs and tables using template-specific font sizing.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_marker_template_specific(p, marker_name)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific(cell, marker_name)

    def _recursive_autosize_template_specific_multi(self, element, markers):
        """
        Recursively find and replace all markers in paragraphs and tables using template-specific font sizing.
        Processes all markers in a single pass to avoid conflicts.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_markers_template_specific(p, markers)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific_multi(cell, markers)

    def _process_paragraph_for_markers_template_specific(self, paragraph, markers):
        """
        Process a single paragraph for multiple markers using template-specific font sizing.
        Handles all markers in a single pass to avoid conflicts.
        """
        full_text = "".join(run.text for run in paragraph.runs)
        
        # DEBUG: Log the input parameters
        self.logger.debug(f"Processing paragraph for markers: {markers}")
        self.logger.debug(f"Full text: '{full_text}'")
        
        # First, check if this is a combined lineage/vendor paragraph
        if self._detect_and_process_combined_lineage_vendor(paragraph):
            return
        
        # Check if any markers are present
        found_markers = []
        for marker_name in markers:
            start_marker = f'{marker_name}_START'
            end_marker = f'{marker_name}_END'
            if start_marker in full_text and end_marker in full_text:
                found_markers.append(marker_name)
                self.logger.debug(f"Found marker: {marker_name} (start: '{start_marker}', end: '{end_marker}')")
            else:
                self.logger.debug(f"Marker not found: {marker_name} (start: '{start_marker}' in text: {start_marker in full_text}, end: '{end_marker}' in text: {end_marker in full_text})")
        
        self.logger.debug(f"Total markers found: {found_markers}")
        
        if found_markers:
            # Process all markers and build the final content
            final_content = full_text
            processed_content = {}
            
            for marker_name in found_markers:
                start_marker = f'{marker_name}_START'
                end_marker = f'{marker_name}_END'
                
                # Extract content for this marker
                start_idx = final_content.find(start_marker)
                end_idx = final_content.find(end_marker) + len(end_marker)
                
                if start_idx != -1 and end_idx != -1:
                    marker_start = final_content.find(start_marker) + len(start_marker)
                    marker_end = final_content.find(end_marker)
                    content = final_content[marker_start:marker_end]
                    
                    # Get font size for this marker
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    self.logger.debug(f"Processing marker '{marker_name}': content='{content}', font_size={font_size.pt}pt")
                    
                    processed_content[marker_name] = {
                        'content': content,
                        'font_size': font_size,
                        'start_pos': start_idx,
                        'end_pos': end_idx
                    }
                    
                    # Remove this marker from final_content so subsequent markers can find their correct positions
                    final_content = final_content[:start_idx] + final_content[end_idx:]
            
            self.logger.debug(f"Processed content: {processed_content}")
            
            # Clear paragraph and rebuild with all processed content
            paragraph.clear()
            
            # Ensure consistent spacing above all marker sections for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # Sort markers by position in text
            sorted_markers = sorted(processed_content.items(), key=lambda x: x[1]['start_pos'])
            
            current_pos = 0
            for marker_name, marker_data in sorted_markers:
                # Add any text before this marker
                if marker_data['start_pos'] > current_pos:
                    text_before = full_text[current_pos:marker_data['start_pos']]
                    # Preserve line breaks and whitespace, but skip if completely empty
                    if text_before or text_before.strip():
                        run = paragraph.add_run(text_before)
                        run.font.name = "Arial"
                        run.font.bold = True
                        run.font.size = Pt(12)  # Default size for non-marker text
                        self.logger.debug(f"Added text before '{marker_name}': '{text_before}' -> 12pt")
                # Add the processed marker content (use the potentially modified content)
                display_content = marker_data.get('display_content', marker_data['content'])
                # --- BULLETPROOF: Only one run for the entire marker content, preserving line breaks ---
                run = paragraph.add_run()
                run.font.name = "Arial"
                
                # Special handling for PRODUCTVENDOR - don't make it bold
                if marker_name == 'PRODUCTVENDOR':
                    run.font.bold = False
                else:
                    run.font.bold = True
                
                run.font.size = marker_data['font_size']
                set_run_font_size(run, marker_data['font_size'])
                self.logger.debug(f"Added marker '{marker_name}': '{display_content}' -> {marker_data['font_size'].pt}pt")
                
                lines = display_content.splitlines()
                for i, line in enumerate(lines):
                    if i > 0:
                        run.add_break()
                    run.add_text(line)
                current_pos = marker_data['end_pos']
            
            # Add any remaining text
            if current_pos < len(full_text):
                text_after = full_text[current_pos:]
                # Preserve line breaks and whitespace, but skip if completely empty
                if text_after or text_after.strip():
                    run = paragraph.add_run(text_after)
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = Pt(12)  # Default size for non-marker text
                    self.logger.debug(f"Added text after: '{text_after}' -> 12pt")
            
            # Convert |BR| markers to actual line breaks after marker processing
            self._convert_br_markers_to_line_breaks(paragraph)
            
            # Apply special formatting for specific markers
            for marker_name, marker_data in processed_content.items():
                # Special handling for ProductBrand markers in Double template
                if ('PRODUCTBRAND' in marker_name) and self.template_type == 'double':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                    continue
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                if marker_name == 'RATIO':
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], 'RATIO', self.template_type, self.scale_factor, product_type))
                        # Ensure ratio values are bold
                        run.font.bold = True
                    continue
                if marker_name == 'LINEAGE':
                    content = marker_data['content']
                    product_type = None
                    if hasattr(self, 'current_product_type'):
                        product_type = self.current_product_type
                    elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                        product_type = self.label_context['ProductType']
                    
                    # Use unified LINEAGE font sizing for all templates including double
                    for run in paragraph.runs:
                        # Use unified font sizing system instead of old get_font_size_by_marker
                        from src.core.generation.unified_font_sizing import get_font_size
                        font_size = get_font_size(content, 'lineage', self.template_type, self.scale_factor)
                        set_run_font_size(run, font_size)
                    
                    # Handle alignment based on PRODUCT TYPE, not just lineage content
                    from src.core.constants import CLASSIC_TYPES
                    is_classic_product = product_type and product_type.lower() in CLASSIC_TYPES
                    
                    # Debug logging for vape cartridge lineage alignment
                    if product_type and 'vape' in product_type.lower():
                        self.logger.debug(f"VAPE CARTRIDGE DEBUG: product_type='{product_type}', is_classic_product={is_classic_product}, CLASSIC_TYPES={CLASSIC_TYPES}")
                    
                    # Classic product types should have LEFT alignment for lineage
                    if is_classic_product:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # NO LEFT INDENT - this was causing lineage indentation
                        paragraph.paragraph_format.left_indent = Inches(0)
                        # Ensure consistent spacing above lineage section for equal margins
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(1)
                    else:
                        # Non-classic product types should have CENTER alignment for lineage
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Ensure consistent spacing above lineage section for equal margins
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(1)
                    
                    # SPECIFIC OVERRIDE: Ensure Vape Cartridge products always have LEFT-aligned lineage
                    if product_type and 'vape' in product_type.lower():
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.left_indent = Inches(0)
                        self.logger.debug(f"VAPE CARTRIDGE OVERRIDE: Forced LEFT alignment for lineage")
                    
                    continue
                # Always center ProductBrand and ProductBrand_Center markers
                if marker_name in ('PRODUCTBRAND', 'PRODUCTBRAND_CENTER') or 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Ensure consistent spacing above product brand section for equal margins
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                    continue
                # Right-align PRODUCTVENDOR markers
                if marker_name == 'PRODUCTVENDOR':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # Ensure consistent spacing above vendor section for equal margins
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    # Use unified font sizing for vendor text
                    for run in paragraph.runs:
                        # Apply unified font sizing using 'vendor' field type
                        from src.core.generation.unified_font_sizing import get_font_size
                        vendor_font_size = get_font_size(marker_data['content'], 'vendor', self.template_type, self.scale_factor)
                        set_run_font_size(run, vendor_font_size)
                        # Set vendor text to italic and light gray color
                        run.font.italic = True
                        from docx.shared import RGBColor
                        run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                        run.font.color.theme_color = None  # Clear any theme color
                    continue
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                else:
                    product_type = None
                
                # Special handling for ProductStrain marker - always use 1pt font
                if marker_name in ('PRODUCTSTRAIN', 'STRAIN'):
                    for run in paragraph.runs:
                        # Only apply 1pt font to runs that contain strain content AND are actually ProductStrain markers
                        # This prevents ProductStrain font sizing from affecting other fields like ProductBrand
                        if (marker_data['content'] in run.text and 
                            any(marker in run.text for marker in ['PRODUCTSTRAIN_START', 'STRAIN_START'])):
                            # Force 1pt font size for ProductStrain markers
                            strain_font_size = Pt(1 * self.scale_factor)
                            set_run_font_size(run, strain_font_size)
                    continue
                
                # Apply normal font sizing for other markers
                set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                # Special handling for ProductVendor marker - now handled above with unified font sizing
                # This section removed to prevent conflicts with the unified font sizing system
            
            self.logger.debug(f"Applied multi-marker processing for: {list(processed_content.keys())}")
        try:
            pass
        except Exception as e:
            self.logger.error(f"Error processing multi-marker template: {e}")
            # Fallback: remove all markers and use default size
            for run in paragraph.runs:
                for marker_name in markers:
                    start_marker = f'{marker_name}_START'
                    end_marker = f'{marker_name}_END'
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                # Use appropriate default size based on template type
                if self.template_type == 'mini':
                    default_size = Pt(8 * self.scale_factor)
                elif self.template_type == 'vertical':
                    default_size = Pt(10 * self.scale_factor)
                else:  # horizontal
                    default_size = Pt(12 * self.scale_factor)
                run.font.size = default_size
        finally:
            # Always check for |BR| markers regardless of success/failure
            self._convert_br_markers_to_line_breaks(paragraph)

    def _process_paragraph_for_marker_template_specific(self, paragraph, marker_name):
        """
        Process a single paragraph for a specific marker using template-type-specific font sizing.
        """
        start_marker = f'{marker_name}_START'
        end_marker = f'{marker_name}_END'
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        if start_marker in full_text and end_marker in full_text:
            try:
                # Extract content
                start_idx = full_text.find(start_marker) + len(start_marker)
                end_idx = full_text.find(end_marker)
                content = full_text[start_idx:end_idx]
                
                # For THC_CBD markers, calculate font size before any splitting to ensure consistency
                if marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL'] and ('\n' in content or '|BR|' in content):
                    # Calculate font size based on the original unsplit content to ensure consistency
                    original_content = content.replace('\n', ' ').replace('|BR|', ' ')
                    font_size = self._get_template_specific_font_size(original_content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with original content '{original_content}' -> font_size: {font_size}")
                    
                    # Clear and recreate with single run approach
                    paragraph.clear()
                    
                    # Create a single run with the entire content
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = font_size
                    set_run_font_size(run, font_size)
                    
                    # Add the content with line breaks as text
                    run.add_text(content)
                    
                    # Convert line breaks to actual line breaks, passing the font size
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                else:
                    # Use template-type-specific font sizing based on original functions
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with content '{content}' -> font_size: {font_size}")
                    
                    # Clear paragraph and re-add content with template-optimized formatting
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    # Special handling for PRODUCTVENDOR - don't make it bold
                    if marker_name == 'PRODUCTVENDOR':
                        run.font.bold = False
                    else:
                        run.font.bold = True
                    run.font.size = font_size
                    
                    # Apply template-specific font size setting
                    set_run_font_size(run, font_size)
                    
                    # Add the content to the run
                    run.add_text(content)
                    
                    # Convert |BR| markers to actual line breaks for other markers
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                
                # Handle special formatting for specific markers
                if marker_name in ['PRODUCTBRAND', 'PRODUCTBRAND_CENTER']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Also ensure all runs in this paragraph are properly sized
                    for run in paragraph.runs:
                        set_run_font_size(run, font_size)
                elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                    # Ensure THC_CBD and RATIO values are bold
                    for run in paragraph.runs:
                        run.font.bold = True
                    
                    # For vertical template, apply line spacing from unified font sizing
                    line_spacing = get_line_spacing_by_marker(marker_name, self.template_type)
                    if line_spacing:
                        paragraph.paragraph_format.line_spacing = line_spacing
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Set at XML level for maximum compatibility
                        pPr = paragraph._element.get_or_add_pPr()
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                        spacing.set(qn('w:lineRule'), 'auto')
                    
                    # For vertical template THC_CBD content, use right alignment for percentage values
                    if self.template_type == 'vertical' and marker_name == 'THC_CBD':
                        # Set paragraph alignment to right for proper percentage alignment
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        self.logger.debug(f"Set right alignment for vertical template THC_CBD content")
                    # All content now uses standard spacing
                    # For all other Ratio content in horizontal template, set vertical alignment to top
                    elif self.template_type == 'horizontal' and marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Set vertical alignment to top for the cell containing this paragraph
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    # For all other THC/CBD content in other templates, set vertical alignment to top
                    elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Set vertical alignment to top for the cell containing this paragraph
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                
                # Center alignment for brand names
                if 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Center alignment for DOH (Date of Harvest)
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Special handling for lineage markers
                if marker_name == 'LINEAGE':
                    self.logger.debug(f"Processing LINEAGE marker with content: '{content}'")
                    # Extract product type information from the content
                    if '_PRODUCT_TYPE_' in content and '_IS_CLASSIC_' in content:
                        parts = content.split('_PRODUCT_TYPE_')
                        if len(parts) == 2:
                            actual_lineage = parts[0]
                            type_info = parts[1]
                            type_parts = type_info.split('_IS_CLASSIC_')
                            if len(type_parts) == 2:
                                product_type = type_parts[0]
                                is_classic_raw = type_parts[1]
                                # Remove LINEAGE_END marker if present
                                if is_classic_raw.endswith('LINEAGE_END'):
                                    is_classic_raw = is_classic_raw[:-len('LINEAGE_END')]
                                is_classic = is_classic_raw.lower() == 'true'
                                
                                # Check if the lineage content itself is a classic lineage value
                                classic_lineages = [
                                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                                ]
                                is_classic_lineage = actual_lineage.upper() in classic_lineages
                                
                                # Don't center if it's either a classic product type OR a classic lineage value
                                if is_classic or is_classic_lineage:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Update the content to only show the actual lineage (remove any markers)
                                if actual_lineage.startswith('LINEAGE_START'):
                                    actual_lineage = actual_lineage[len('LINEAGE_START'):]
                                content = actual_lineage
                    else:
                        # Fallback: check if this is a classic product type by using the context
                        # Import constants to check against CLASSIC_TYPES
                        from src.core.constants import CLASSIC_TYPES
                        
                        # Get product type from context, not from content
                        is_classic_product = False
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        else:
                            product_type = None
                        
                        # Check if the product type is classic
                        if product_type:
                            is_classic_product = product_type.lower() in CLASSIC_TYPES
                            # Debug logging for vape cartridge lineage alignment in fallback
                            if 'vape' in product_type.lower():
                                self.logger.debug(f"VAPE CARTRIDGE FALLBACK DEBUG: product_type='{product_type}', is_classic_product={is_classic_product}, CLASSIC_TYPES={CLASSIC_TYPES}")
                        
                        # Also check if the content is a classic lineage value
                        content_upper = content.upper()
                        classic_lineages = [
                            "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                            "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                        ]
                        is_classic_lineage = content_upper in classic_lineages
                        
                        # If it's either a classic product type OR a classic lineage value, left-justify
                        if is_classic_product or is_classic_lineage:
                            # For Classic Types or Classic Lineages, left-justify the text
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            self.logger.debug(f"Left-justified lineage for classic content: '{content}' (classic_product: {is_classic_product}, classic_lineage: {is_classic_lineage}, product_type: {product_type})")
                        else:
                            # Only center if it's NOT a classic type and NOT a classic lineage
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self.logger.debug(f"Centered lineage for non-classic content: '{content}' (product_type: {product_type})")
                        
                        # SPECIFIC OVERRIDE: Ensure Vape Cartridge products always have LEFT-aligned lineage (fallback)
                        if product_type and 'vape' in product_type.lower():
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.left_indent = Inches(0)
                            self.logger.debug(f"VAPE CARTRIDGE FALLBACK OVERRIDE: Forced LEFT alignment for lineage")
                
                self.logger.debug(f"Applied template-specific font sizing: {font_size.pt}pt for {marker_name} marker")

            except Exception as e:
                self.logger.error(f"Error processing template-specific marker {marker_name}: {e}")
                # Fallback: remove markers and use default size based on template type
                for run in paragraph.runs:
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                    # Use appropriate default size based on template type
                    if self.template_type == 'mini':
                        default_size = Pt(8 * self.scale_factor)
                    elif self.template_type == 'vertical':
                        default_size = Pt(10 * self.scale_factor)
                    else:  # horizontal
                        default_size = Pt(12 * self.scale_factor)
                    run.font.size = default_size
        elif start_marker in full_text or end_marker in full_text:
            # Log partial markers for debugging
            self.logger.debug(f"Found partial {marker_name} marker in text: '{full_text[:100]}...'")

    def _convert_br_markers_to_line_breaks(self, paragraph, font_size=None):
        """
        Convert |BR| markers and \n characters in paragraph text to actual line breaks.
        This splits the text at |BR| markers or \n characters and creates separate runs for each part.
        """
        try:
            # Get all text from the paragraph and store existing font sizes
            full_text = "".join(run.text for run in paragraph.runs)
            
            # Store existing font sizes for each run
            existing_sizes = []
            for run in paragraph.runs:
                if run.text.strip():
                    existing_sizes.append(run.font.size)
            
            # If we have existing sizes, use the first one for all runs to ensure consistency
            # Or use the passed font_size parameter if provided
            consistent_font_size = None
            if font_size is not None:
                consistent_font_size = font_size
            elif existing_sizes:
                consistent_font_size = existing_sizes[0]
            
            # Check if there are any |BR| markers or \n characters
            if '|BR|' not in full_text and '\n' not in full_text:
                return
            
            # First split by |BR| markers, then by \n characters
            if '|BR|' in full_text:
                parts = full_text.split('|BR|')
            else:
                parts = full_text.split('\n')
            
            # Clear the paragraph
            paragraph.clear()
            
            # Set tight paragraph spacing to prevent excessive gaps
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # Only set line spacing if it's not already set (to preserve custom line spacing)
                        # Use standard line spacing for all content
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Add each part as a separate run, with line breaks between them
            size_index = 0
            for i, part in enumerate(parts):
                if part.strip():  # Only add non-empty parts
                    # Strip whitespace for all content to remove extra spaces
                    run = paragraph.add_run(part.strip())
                    run.font.name = "Arial"
                    
                    # Check if this paragraph contains ratio content and should be bold
                    # This ensures multi-line ratio content stays bold
                    if any(pattern in full_text for pattern in [
                        'mg THC', 'mg CBD', 'mg CBC', 'mg CBG', 'mg CBN',
                        'THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:',
                        '1:1', '2:1', '3:1', '1:1:1', '2:1:1',
                        'RATIO_START', 'THC_CBD_START'
                    ]):
                        run.font.bold = True
                    
                    # Use consistent font size for all runs
                    if consistent_font_size:
                        run.font.size = consistent_font_size
                    else:
                        # Use a default size only if no existing size is available
                        run.font.size = Pt(12)
                    
                    # Add a line break after this part only if the next part is not empty
                    if i < len(parts) - 1 and parts[i + 1].strip():
                        # Use add_break() with WD_BREAK.LINE to create proper line breaks within the same paragraph
                        run.add_break(WD_BREAK.LINE)
            
            # All content now uses standard 1.0 line spacing
            
            self.logger.debug(f"Converted {len(parts)-1} |BR| markers to line breaks")
            
        except Exception as e:
            self.logger.error(f"Error converting BR markers to line breaks: {e}")
            # Fallback: just remove the BR markers
            for run in paragraph.runs:
                run.text = run.text.replace('|BR|', ' ')

    def _ensure_consistent_lineage_spacing(self, doc):
        """
        Ensure consistent spacing above lineage/brand sections for equal margins across all labels.
        This creates uniform visual spacing above the colored lineage/brand bars.
        """
        try:
            def process_paragraph(paragraph):
                # Check if this paragraph contains lineage or vendor content
                text = paragraph.text.lower()
                if any(keyword in text for keyword in ['indica', 'sativa', 'hybrid', 'cbd', 'alpha crux', 'constellation']):
                    # Set consistent spacing for lineage/brand sections
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    
                    # Also set at XML level for maximum compatibility
                    pPr = paragraph._element.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:before'), '40')  # 2pt = 40 twips
                    spacing.set(qn('w:after'), '20')   # 1pt = 20 twips
                    spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            self.logger.debug("Ensured consistent spacing above lineage/brand sections")
            
        except Exception as e:
            self.logger.error(f"Error ensuring consistent lineage spacing: {e}")

    def _add_consistent_content_spacing(self, doc):
        """
        Add consistent spacing above main content sections for better visual balance.
        This ensures uniform spacing above product names, prices, and other key content.
        """
        try:
            def process_paragraph(paragraph):
                # Check if this paragraph contains main content
                text = paragraph.text.lower()
                if any(keyword in text for keyword in ['$', 'thc:', 'cbd:', 'mg', 'oz', 'g', 'pack']):
                    # Add consistent spacing above main content sections
                    current_before = paragraph.paragraph_format.space_before
                    current_after = paragraph.paragraph_format.space_after
                    
                    # Only add spacing if it's not already set to our target values
                    if current_before == Pt(0) or current_before is None:
                        paragraph.paragraph_format.space_before = Pt(1)
                    
                    if current_after == Pt(0) or current_after is None:
                        paragraph.paragraph_format.space_after = Pt(0.5)
                    
                    # Also set at XML level for maximum compatibility
                    pPr = paragraph._element.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    
                    if current_before == Pt(0) or current_before is None:
                        spacing.set(qn('w:before'), '20')  # 1pt = 20 twips
                    if current_after == Pt(0) or current_after is None:
                        spacing.set(qn('w:after'), '10')   # 0.5pt = 10 twips
                    
                    spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            self.logger.debug("Added consistent spacing above main content sections")
            
        except Exception as e:
            self.logger.error(f"Error adding consistent content spacing: {e}")

    def _fix_ratio_paragraph_spacing(self, doc):
        """
        Fix paragraph spacing for ratio content to prevent excessive gaps between lines.
        This ensures tight spacing for multi-line ratio content.
        """
        try:
            # Define patterns that indicate ratio content
            ratio_patterns = [
                'mg THC', 'mg CBD', 'mg CBG', 'mg CBN', 'mg CBC',
                'THC:', 'CBD:', 'CBG:', 'CBN:', 'CBC:',
                '1:1', '2:1', '3:1', '1:1:1', '2:1:1'
            ]
            
            def process_paragraph(paragraph):
                # Check if this paragraph contains ratio content
                text = paragraph.text.lower()
                if any(pattern.lower() in text for pattern in ratio_patterns):
                    # Set tight spacing for all ratio content (including THC_CBD)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    # Also set tight spacing for any child paragraphs (in case of nested content)
                    for child_para in paragraph._element.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if hasattr(child_para, 'pPr') and child_para.pPr is not None:
                            # Set spacing properties at XML level for maximum compatibility
                            spacing = child_para.pPr.find('.//w:spacing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                child_para.pPr.append(spacing)
                            
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing (240 twips)
                            spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            self.logger.debug("Fixed paragraph spacing for ratio content")
            
        except Exception as e:
            self.logger.error(f"Error fixing ratio paragraph spacing: {e}")
            # Don't raise the exception - this is a formatting enhancement that shouldn't break the main process

    def _validate_and_repair_table_structure(self, table):
        """
        Validate and repair table structure to ensure it has required elements.
        Returns True if table is valid, False if it cannot be repaired.
        """
        try:
            # First, try to access table properties to see if there's an actual error
            try:
                _ = table.rows
                _ = table.columns
                # If we can access these without error, the table is fine
                return True
            except Exception as e:
                # Table is corrupted, try to repair
                self.logger.debug(f"Table access failed, attempting repair: {e}")
            
            # Check if table has the required tblGrid element
            tblGrid = table._element.find(qn('w:tblGrid'))
            if tblGrid is None:
                # Create tblGrid element
                tblGrid = OxmlElement('w:tblGrid')
                
                # Try to get column count from the XML structure directly
                try:
                    # Look for rows in the XML to determine column count
                    rows = table._element.findall(qn('w:tr'))
                    if rows:
                        # Count cells in the first row
                        first_row_cells = rows[0].findall(qn('w:tc'))
                        col_count = len(first_row_cells)
                        
                        # Create grid columns
                        for _ in range(col_count):
                            gridCol = OxmlElement('w:gridCol')
                            gridCol.set(qn('w:w'), '1440')  # Default width of 1 inch
                            tblGrid.append(gridCol)
                        
                        # Insert tblGrid at the beginning of the table element
                        table._element.insert(0, tblGrid)
                        self.logger.debug(f"Repaired missing tblGrid for table with {col_count} columns")
                        
                        # Now test if the repair worked
                        try:
                            _ = table.rows
                            _ = table.columns
                            return True
                        except Exception as test_error:
                            self.logger.error(f"Table repair failed validation test: {test_error}")
                            return False
                    else:
                        self.logger.warning("Cannot repair table: no rows found in XML")
                        return False
                except Exception as repair_error:
                    self.logger.error(f"Failed to repair table structure: {repair_error}")
                    return False
            else:
                # Table already has tblGrid, but let's verify it's working
                try:
                    _ = table.rows
                    _ = table.columns
                    return True
                except Exception as verify_error:
                    self.logger.error(f"Table has tblGrid but still corrupted: {verify_error}")
                    # Try to repair the existing tblGrid
                    try:
                        # Remove and recreate tblGrid
                        old_tblGrid = table._element.find(qn('w:tblGrid'))
                        if old_tblGrid is not None:
                            old_tblGrid.getparent().remove(old_tblGrid)
                        
                        # Create new tblGrid
                        new_tblGrid = OxmlElement('w:tblGrid')
                        rows = table._element.findall(qn('w:tr'))
                        if rows:
                            first_row_cells = rows[0].findall(qn('w:tc'))
                            col_count = len(first_row_cells)
                            
                            for _ in range(col_count):
                                gridCol = OxmlElement('w:gridCol')
                                gridCol.set(qn('w:w'), '1440')
                                new_tblGrid.append(gridCol)
                            
                            table._element.insert(0, new_tblGrid)
                            
                            # Test the repair
                            try:
                                _ = table.rows
                                _ = table.columns
                                self.logger.debug(f"Successfully repaired corrupted tblGrid for table with {col_count} columns")
                                return True
                            except Exception as final_test_error:
                                self.logger.error(f"Final repair attempt failed: {final_test_error}")
                                return False
                        else:
                            return False
                    except Exception as final_repair_error:
                        self.logger.error(f"Final repair attempt failed: {final_repair_error}")
                        # Last resort: try to rebuild the entire table
                        try:
                            if self._rebuild_corrupted_table(table, self.template_type):
                                self.logger.info("Table successfully rebuilt after all repair attempts failed")
                                return True
                            else:
                                self.logger.error("Table rebuild failed, table is beyond repair")
                                return False
                        except Exception as rebuild_error:
                            self.logger.error(f"Table rebuild attempt failed: {rebuild_error}")
                            return False
                
        except Exception as e:
            self.logger.error(f"Error validating/repairing table structure: {e}")
            return False

    def _safe_table_iteration(self, table, operation_name="table operation"):
        """
        Safely iterate through table rows and cells with comprehensive error handling.
        Returns True if successful, False if table is corrupted beyond repair.
        """
        try:
            # First validate the table structure
            if not self._validate_and_repair_table_structure(table):
                self.logger.warning(f"Table validation failed for {operation_name}")
                return False
            
            # Test basic table access
            try:
                rows = table.rows
                if not rows:
                    self.logger.warning(f"Table has no rows for {operation_name}")
                    return False
            except Exception as e:
                self.logger.error(f"Table rows access failed for {operation_name}: {e}")
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error in safe table iteration for {operation_name}: {e}")
            return False

    def _safe_table_processing(self, table, operation_name="table operation", processor_func=None):
        """
        Safely process a table with comprehensive error handling and repair attempts.
        
        Args:
            table: The table to process
            operation_name: Name of the operation for logging
            processor_func: Function to apply to each row/cell (optional)
            
        Returns:
            bool: True if processing was successful, False otherwise
        """
        try:
            # First validate the table structure
            if not self._safe_table_iteration(table, operation_name):
                return False
            
            # If no processor function provided, just return success
            if processor_func is None:
                return True
            
            # Process the table safely
            try:
                for row in table.rows:
                    try:
                        # Validate row structure
                        if not hasattr(row, 'cells') or not row.cells:
                            self.logger.warning(f"Skipping row with invalid structure during {operation_name}")
                            continue
                        
                        for cell in row.cells:
                            try:
                                # Apply the processor function to the cell
                                processor_func(cell)
                            except Exception as cell_error:
                                self.logger.warning(f"Skipping cell due to error during {operation_name}: {cell_error}")
                                continue
                    except Exception as row_error:
                        self.logger.warning(f"Skipping row due to error during {operation_name}: {row_error}")
                        continue
                
                return True
                
            except Exception as processing_error:
                self.logger.error(f"Error during table processing for {operation_name}: {processing_error}")
                return False
                
        except Exception as e:
            self.logger.error(f"Error in safe table processing for {operation_name}: {e}")
            return False

    def _rebuild_corrupted_table(self, table, template_type):
        """
        Attempt to completely rebuild a corrupted table structure.
        This is a last resort when other repair methods fail.
        
        Args:
            table: The corrupted table
            template_type: Type of template to determine structure
            
        Returns:
            bool: True if rebuild was successful, False otherwise
        """
        try:
            self.logger.warning(f"Attempting to rebuild corrupted table for template type: {template_type}")
            
            # Get the table element
            table_element = table._element
            
            # Determine expected structure based on template type
            if template_type == 'vertical':
                expected_rows, expected_cols = 3, 3
            elif template_type == 'double':
                expected_rows, expected_cols = 4, 3
            elif template_type == 'mini':
                expected_rows, expected_cols = 4, 5
            elif template_type == 'inventory':
                expected_rows, expected_cols = 2, 2
            else:
                expected_rows, expected_cols = 3, 3  # Default
            
            # Create new table structure
            new_table_element = OxmlElement('w:tbl')
            
            # Add table properties
            tblPr = OxmlElement('w:tblPr')
            new_table_element.append(tblPr)
            
            # Add table layout
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Add table width
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(expected_cols * 1440))  # 1 inch per column
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Add table grid
            tblGrid = OxmlElement('w:tblGrid')
            for _ in range(expected_cols):
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), '1440')
                tblGrid.append(gridCol)
            new_table_element.append(tblGrid)
            
            # Add table rows
            for _ in range(expected_rows):
                tr = OxmlElement('w:tr')
                for _ in range(expected_cols):
                    tc = OxmlElement('w:tc')
                    # Add cell properties
                    tcPr = OxmlElement('w:tcPr')
                    tc.append(tcPr)
                    # Add cell width
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), '1440')
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
                    # Add empty paragraph
                    p = OxmlElement('w:p')
                    tc.append(p)
                    tr.append(tc)
                new_table_element.append(tr)
            
            # Replace the old table element
            parent = table_element.getparent()
            if parent is not None:
                parent.replace(table_element, new_table_element)
                self.logger.info(f"Successfully rebuilt table structure for {template_type} template")
                return True
            else:
                self.logger.error("Cannot rebuild table: no parent element found")
                return False
                
        except Exception as e:
            self.logger.error(f"Failed to rebuild corrupted table: {e}")
            return False

    def _ensure_proper_centering(self, doc):
        """
        Ensure tables are properly centered in the document with correct margins and spacing.
        """
        try:
            # Set document margins to ensure proper centering
            for section in doc.sections:
                # Use smaller margins for vertical template to fit all 9 labels
                if self.template_type == 'vertical':
                    section.left_margin = Inches(0.25)
                    section.right_margin = Inches(0.25)
                    section.top_margin = Inches(0.25)
                    section.bottom_margin = Inches(0.25)
                else:
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
            
            # Remove any extra paragraphs that might affect centering
            paragraphs_to_remove = []
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip() and not paragraph.runs:
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                paragraph._element.getparent().remove(paragraph._element)
            
            # Ensure all tables are properly centered and have valid structure
            for table in doc.tables:
                # Use safe table iteration to validate and repair if needed
                if not self._safe_table_iteration(table, "centering setup"):
                    self.logger.warning(f"Skipping table that cannot be repaired during centering setup")
                    continue
                
                # Set table alignment to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Ensure table properties are set correctly
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                
                # Set table to fixed layout
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                
                # Ensure table is not auto-fit
                table.autofit = False
                if hasattr(table, 'allow_autofit'):
                    table.allow_autofit = False
                

                
                # Calculate and set proper table width for perfect centering
                from src.core.constants import CELL_DIMENSIONS, GRID_LAYOUTS
                
                # Get individual cell dimensions and grid layout
                cell_dims = CELL_DIMENSIONS.get(self.template_type, {'width': 2.4, 'height': 2.4})
                grid_layout = GRID_LAYOUTS.get(self.template_type, {'rows': 3, 'cols': 3})
                
                # Calculate total table width: individual cell width * number of columns
                individual_cell_width = cell_dims['width']
                num_columns = grid_layout['cols']
                total_table_width = individual_cell_width * num_columns
                
                # Set table width to ensure proper centering
                table.width = Inches(total_table_width)
                
                # Also set the table width property in XML to ensure it's properly applied
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)
                
                # Set table width property
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), str(int(total_table_width * 1440)))  # Convert to twips
                tblW.set(qn('w:type'), 'dxa')
                
                # For double templates, ensure proper table grid structure without accessing table.columns
                if self.template_type == 'double':
                    # Use safe table iteration to get column count
                    if self._safe_table_iteration(table, "double template grid setup"):
                        # Get column count from XML structure instead of table.columns
                        table_element = table._element
                        rows = table_element.findall(qn('w:tr'))
                        if rows:
                            first_row = rows[0]
                            cells = first_row.findall(qn('w:tc'))
                            col_count = len(cells)
                            
                            # Ensure tblGrid exists and has correct structure
                            tblGrid = table_element.find(qn('w:tblGrid'))
                            if tblGrid is None:
                                # Create new tblGrid
                                tblGrid = OxmlElement('w:tblGrid')
                                table_element.insert(0, tblGrid)
                            
                            # Clear existing grid columns and recreate
                            for existing_gc in tblGrid.findall(qn('w:gridCol')):
                                tblGrid.remove(existing_gc)
                            
                            # Add grid columns with proper widths for double template
                            col_width = 1.75  # 1.75 inches per column for double template
                            for _ in range(col_count):
                                gc = OxmlElement('w:gridCol')
                                gc.set(qn('w:w'), str(int(col_width * 1440)))  # Convert to twips
                                tblGrid.append(gc)
                            
                            # Also ensure each cell has the correct width property
                            for row in table.rows:
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcW = tcPr.find(qn('w:tcW'))
                                    if tcW is None:
                                        tcW = OxmlElement('w:tcW')
                                        tcPr.append(tcW)
                                    tcW.set(qn('w:w'), str(int(col_width * 1440)))
                                    tcW.set(qn('w:type'), 'dxa')
                
                # For other template types, use the existing logic but with safety checks
                elif self.template_type not in ['horizontal', 'mini', 'vertical']:
                    # Use safe table iteration to get column count
                    if self._safe_table_iteration(table, "grid setup"):
                        # Get column count from XML structure
                        table_element = table._element
                        rows = table_element.findall(qn('w:tr'))
                        if rows:
                            first_row = rows[0]
                            cells = first_row.findall(qn('w:tc'))
                            col_count = len(cells)
                            
                            # Create new grid with proper column widths
                            tblGrid = OxmlElement('w:tblGrid')
                            col_width = cell_dims['width']
                            
                            for _ in range(col_count):
                                gridCol = OxmlElement('w:gridCol')
                                gridCol.set(qn('w:w'), str(int(col_width * 1440)))  # Convert to twips
                                tblGrid.append(gridCol)
                            
                            # Remove existing grid if present
                            existing_grid = table_element.find(qn('w:tblGrid'))
                            if existing_grid is not None:
                                existing_grid.getparent().remove(existing_grid)
                            
                            # Insert the grid at the beginning of the table element
                            table_element.insert(0, tblGrid)
                            
                            # Also ensure each cell has the correct width property
                            for row in table.rows:
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcW = tcPr.find(qn('w:tcW'))
                                    if tcW is None:
                                        tcW = OxmlElement('w:tcW')
                                        tcPr.append(tcW)
                                    tcW.set(qn('w:w'), str(int(col_width * 1440)))
                                    tcW.set(qn('w:type'), 'dxa')
            
            self.logger.debug("Ensured proper table centering and document setup")
            
        except Exception as e:
            self.logger.error(f"Error ensuring proper centering: {e}")

    def _add_weight_units_markers(self, doc):
        """
        Add RATIO markers around weight units content for mini templates with classic types.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during weight units marker addition")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for weight units content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                # Check if this run contains weight units content (ends with 'g' or 'mg', or contains specific patterns)
                                # More specific check to avoid marking brand names that contain 'g'
                                is_weight_unit = (
                                    run_text.strip().endswith('g') or 
                                    run_text.strip().endswith('mg') or
                                    re.match(r'^\d+\.?\d*\s*g$', run_text.strip()) or  # "1g", "1.5g"
                                    re.match(r'^\d+\.?\d*\s*mg$', run_text.strip()) or  # "100mg", "50.5mg"
                                    re.match(r'^\d+\.?\d*\s*g\s*x\s*\d+', run_text.strip()) or  # "1g x 2"
                                    re.match(r'^\d+\.?\d*\s*mg\s*x\s*\d+', run_text.strip())  # "100mg x 2"
                                )
                                
                                if is_weight_unit and 'RATIO_START' not in run_text:
                                    # This is likely weight units content that needs markers
                                    # Replace the run text with marked content
                                    run.text = f"RATIO_START{run_text}RATIO_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Default size, will be adjusted by post-processing
                                    
                                    self.logger.debug(f"Added RATIO markers around weight units: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding weight units markers: {e}")

    def _add_brand_markers(self, doc):
        """
        Add PRODUCTBRAND_CENTER markers around brand content for mini templates.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            # Import CLASSIC_TYPES to check if current product type is classic
            from src.core.constants import CLASSIC_TYPES
            
            # Get current product type if available
            current_product_type = None
            if hasattr(self, 'current_product_type'):
                current_product_type = self.current_product_type
            elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                current_product_type = self.label_context['ProductType']
            
            # Check if current product type is a classic type
            is_classic_type = False
            if current_product_type:
                is_classic_type = current_product_type.lower() in [ct.lower() for ct in CLASSIC_TYPES]
                self.logger.debug(f"Product type: {current_product_type}, Is classic: {is_classic_type}")
            else:
                self.logger.debug(f"No current_product_type available")
            
            # For mini templates, always add brand markers regardless of product type
            # For other templates, skip brand marker addition for classic types (they should show lineage instead of brand)
            if self.template_type != 'mini' and is_classic_type:
                self.logger.debug(f"Skipping brand marker addition for classic type: {current_product_type}")
                return
            
            self.logger.debug(f"Processing brand markers for {'mini template' if self.template_type == 'mini' else f'non-classic type: {current_product_type}'}")
            
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during brand marker addition")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for brand content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                self.logger.debug(f"Processing run text: '{run_text}'")
                                # Check if this run contains brand content (not empty and not already marked)
                                # Only add markers to text that looks like brand names (not empty, not marked, not placeholders)
                                # IMPORTANT: Don't add brand markers if content is already wrapped in RATIO markers
                                # IMPORTANT: Don't add brand markers if content is already wrapped in PRODUCTSTRAIN markers
                                if (run_text.strip() and 
                                    'PRODUCTBRAND_CENTER_START' not in run_text and 
                                    'PRODUCTSTRAIN_START' not in run_text and  # Don't mark content already in PRODUCTSTRAIN markers
                                    'PRODUCTSTRAIN_END' not in run_text and    # Don't mark content already in PRODUCTSTRAIN markers
                                    'RATIO_START' not in run_text and  # Don't mark content already in RATIO markers
                                    'RATIO_END' not in run_text and    # Don't mark content already in RATIO markers
                                    '{{' not in run_text and 
                                    '}}' not in run_text and
                                    'QR_PLACEHOLDER' not in run_text and  # Don't mark QR placeholders
                                    len(run_text.strip()) > 0 and
                                    # Only mark content that looks like brand names (not numbers, not empty)
                                    not run_text.strip().isdigit() and
                                    not run_text.strip().startswith('$') and
                                    not run_text.strip().endswith('g') and
                                    not run_text.strip().endswith('mg')):
                                    # This is likely brand content that needs markers
                                    # Replace the run text with marked content
                                    run.text = f"PRODUCTBRAND_CENTER_START{run_text}PRODUCTBRAND_CENTER_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Default size, will be adjusted by post-processing
                                    
                                    # Ensure brand content is centered for mini templates
                                    if self.template_type == 'mini':
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        self.logger.debug(f"Set center alignment for mini template brand: {run_text}")
                                    
                                    self.logger.debug(f"Added PRODUCTBRAND_CENTER markers around brand: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding brand markers: {e}")

    def _ensure_mini_template_brand_centering(self, doc):
        """
        Ensure all brand content in mini templates is properly centered.
        This method specifically handles mini template brand alignment.
        """
        try:
            if self.template_type != 'mini':
                return
                
            self.logger.debug("Ensuring brand content centering for mini template")
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Check if this paragraph contains brand content
                            paragraph_text = paragraph.text
                            
                            # For mini templates, be more aggressive about finding brand content
                            is_brand_content = False
                            
                            # Check for explicit brand markers
                            if ('PRODUCTBRAND_CENTER_START' in paragraph_text or 
                                'PRODUCTBRAND_CENTER_END' in paragraph_text or
                                'PRODUCTBRAND_START' in paragraph_text or
                                'PRODUCTBRAND_END' in paragraph_text):
                                is_brand_content = True
                            
                            # For mini templates, also check for content that looks like brand names
                            elif self.template_type == 'mini':
                                # Look for content that appears to be brand names (not empty, not numbers, not prices, not weights)
                                clean_text = paragraph_text.strip()
                                if (clean_text and 
                                    not clean_text.startswith('$') and
                                    not clean_text.endswith('g') and
                                    not clean_text.endswith('mg') and
                                    not clean_text.isdigit() and
                                    not ('THC:' in clean_text and 'CBD:' in clean_text) and
                                    len(clean_text) < 50 and
                                    # Not lineage values
                                    clean_text.upper() not in ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED"]):
                                    is_brand_content = True
                                    self.logger.debug(f"Identified potential brand content in mini template: {clean_text}")
                            
                            if is_brand_content:
                                # Set center alignment for brand paragraphs
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                self.logger.debug(f"Centered brand paragraph: {paragraph_text[:50]}...")
                                
                                # Also ensure all runs in brand paragraphs are properly formatted
                                for run in paragraph.runs:
                                    if run.text.strip():
                                        run.font.name = "Arial"
                                        run.font.bold = True
                                        
            self.logger.debug("Completed mini template brand centering")
            
        except Exception as e:
            self.logger.error(f"Error ensuring mini template brand centering: {e}")

    def _apply_brand_centering_for_double_template(self, doc):
        """
        Apply brand centering logic specifically for the double template.
        This method ensures that non-classic types get centered brand names,
        while classic types maintain their default alignment (left-aligned).
        """
        try:
            # Import CLASSIC_TYPES to check if current product type is classic
            from src.core.constants import CLASSIC_TYPES
            
            # Get current product type if available
            current_product_type = None
            if hasattr(self, 'current_record') and self.current_record:
                current_product_type = (self.current_record.get('ProductType', '').lower() or 
                                      self.current_record.get('Product Type*', '').lower())
            elif hasattr(self, 'current_product_type'):
                current_product_type = self.current_product_type
            
            # Check if current product type is a classic type
            is_classic_type = False
            if current_product_type:
                is_classic_type = current_product_type.lower() in [ct.lower() for ct in CLASSIC_TYPES]
                self.logger.debug(f"Double template brand centering - Product type: {current_product_type}, Is classic: {is_classic_type}")
            else:
                self.logger.debug(f"Double template brand centering - No current_product_type available")
            
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during brand centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.strip()
                            
                            # Skip empty paragraphs
                            if not paragraph_text:
                                continue
                            
                            # Check if this paragraph contains brand content
                            # Look for common brand indicators (not lineage, not price, not weight)
                            is_brand_content = (
                                paragraph_text and
                                not paragraph_text.startswith('$') and
                                not paragraph_text.endswith('g') and
                                not paragraph_text.endswith('mg') and
                                not paragraph_text.isdigit() and
                                # Not classic lineage values
                                paragraph_text.upper() not in ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED"] and
                                # Not THC/CBD content
                                not ('THC:' in paragraph_text and 'CBD:' in paragraph_text) and
                                # Not description content (usually longer)
                                len(paragraph_text) < 50
                            )
                            
                            if is_brand_content:
                                if is_classic_type:
                                    # For classic types, ensure they are NOT centered (left-aligned)
                                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        self.logger.debug(f"Fixed classic type alignment: set to LEFT for '{paragraph_text[:30]}...'")
                                else:
                                    # For non-classic types, center the brand content
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    self.logger.debug(f"Applied brand centering for non-classic type: '{paragraph_text[:30]}...'")
            
        except Exception as e:
            self.logger.error(f"Error applying brand centering for double template: {e}")

    def _get_template_specific_font_size(self, content, marker_name):
        """
        Get font size using the unified font sizing system.
        """
        # Import get_font_size locally to avoid scoping issues
        from src.core.generation.unified_font_sizing import get_font_size
        
        # Special handling for RATIO marker: if content contains THC/CBD data, use THC_CBD field type
        if marker_name == 'RATIO' and ('THC:' in content or 'CBD:' in content):
            # Use THC_CBD field type for THC/CBD content
            return get_font_size(content, 'thc_cbd', self.template_type, self.scale_factor)
        
        # Use unified font sizing system for all templates
        return get_font_size_by_marker(content, marker_name, self.template_type, self.scale_factor)

    def fix_hyphen_spacing(self, text):
        """Replace regular hyphens with non-breaking hyphens to prevent line breaks, 
        but add line breaks before hanging hyphens.
        Used for general text formatting to prevent unwanted line breaks at hyphens."""
        if not text:
            return text
        
        # First, normalize hyphen spacing to ensure consistent format
        text = re.sub(r'\s*-\s*', ' - ', text)
        
        # Check for hanging hyphens (hyphen at the end of a line or followed by a space and then end)
        # Pattern: space + hyphen + space + end of string, or space + hyphen + end of string
        if re.search(r' - $', text) or re.search(r' - \s*$', text):
            # Add line break before the hanging hyphen
            text = re.sub(r' - (\s*)$', r'\n- \1', text)
        
        return text

    def format_with_soft_hyphen(self, text):
        """Format text with soft hyphen + nonbreaking space + value pattern.
        Used for specific formatting where you want a soft hyphen followed by nonbreaking space."""
        if not text:
            return text
        # Replace any leading hyphens/spaces with a single soft hyphen + nonbreaking space
        text = re.sub(r'^[\s\-]+', '\u00AD\u00A0', text)
        # If it didn't start with hyphen/space, prepend
        if not text.startswith('\u00AD\u00A0'):
            text = f'\u00AD\u00A0{text}'
        return text

    def format_classic_ratio(self, text, record=None):
        """
        Format ratio for classic types. Handles various input formats and converts them to the standard display format.
        """
        if not text:
            return text
        
        # Clean the text and normalize
        text = text.strip()
        
        # Handle the default "THC:|BR|CBD:" format from excel processor
        if text == "THC:|BR|CBD:" or text == "THC: | BR | CBD:":
            self.logger.debug(f"Processing THC:|BR|CBD: placeholder for record: {record.get('Product Name*', 'Unknown')}")
            if record:
                # Show all relevant THC fields for debugging
                self.logger.debug(f"RECORD THC FIELDS: THC: '{record.get('THC', '')}', THC test result: '{record.get('THC test result', '')}', Total THC: '{record.get('Total THC', '')}', THCA: '{record.get('THCA', '')}'")
                # Always use Excel 'THC test result' if present and valid
                excel_thc = str(record.get('THC test result', '')).strip()
                try:
                    excel_thc_float = float(excel_thc)
                    if excel_thc not in ['0', '0.0', '', 'nan', 'NaN']:
                        thc_value = excel_thc
                        self.logger.debug(f"USING EXCEL THC test result: '{thc_value}'")
                    else:
                        raise ValueError
                except Exception:
                    # Fallback to highest of all other fields if Excel value is not valid
                    thc_candidates = []
                    thc_debug_vals = {}
                    for key in ['Total THC', 'THCA', 'THC']:
                        val = str(record.get(key, '')).strip()
                        thc_debug_vals[key] = val
                        try:
                            val_float = float(val)
                            if val not in ['0', '0.0', '', 'nan', 'NaN']:
                                thc_candidates.append(val_float)
                        except Exception:
                            continue
                    self.logger.debug(f"THC candidate values (fallback): {thc_debug_vals}, numeric candidates: {thc_candidates}")
                    if thc_candidates:
                        max_thc = max(thc_candidates)
                        thc_value = str(max_thc)
                        self.logger.debug(f"USING HIGHEST THC VALUE (fallback): '{thc_value}' from candidates: {thc_candidates}")
                    else:
                        thc_value = '0'
                        self.logger.debug("No valid THC value found, defaulting to 0")

                # Always use the highest value for both THC and CBD from all relevant fields
                # Always use Excel 'CBD test result' if present and valid
                excel_cbd = str(record.get('CBD test result', '')).strip()
                try:
                    excel_cbd_float = float(excel_cbd)
                    if excel_cbd not in ['0', '0.0', '', 'nan', 'NaN']:
                        cbd_value = excel_cbd
                        self.logger.debug(f"USING EXCEL CBD test result: '{cbd_value}'")
                    else:
                        raise ValueError
                except Exception:
                    # Fallback to highest of all other fields if Excel value is not valid
                    cbd_candidates = []
                    cbd_debug_vals = {}
                    for key in ['Total CBD', 'CBDA', 'CBD']:
                        val = str(record.get(key, '')).strip()
                        cbd_debug_vals[key] = val
                        try:
                            val_float = float(val)
                            if val not in ['0', '0.0', '', 'nan', 'NaN']:
                                cbd_candidates.append(val_float)
                        except Exception:
                            continue
                    self.logger.debug(f"CBD candidate values (fallback): {cbd_debug_vals}, numeric candidates: {cbd_candidates}")
                    if cbd_candidates:
                        max_cbd = max(cbd_candidates)
                        cbd_value = str(max_cbd)
                        self.logger.debug(f"USING HIGHEST CBD VALUE (fallback): '{cbd_value}' from candidates: {cbd_candidates}")
                    else:
                        cbd_value = '0'
                        self.logger.debug("No valid CBD value found, defaulting to 0")
                
                if not cbd_value or cbd_value == '0' or cbd_value == '0.0':
                    total_cbd_value = str(record.get('Total CBD', '')).strip()
                    cbd_test_result = str(record.get('CBD test result', '')).strip()
                    
                    if total_cbd_value and total_cbd_value != '0' and total_cbd_value != '0.0':
                        cbd_value = total_cbd_value
                        self.logger.debug(f"Using database Total CBD: '{cbd_value}'")
                    elif cbd_test_result and cbd_test_result != '0' and cbd_test_result != '0.0':
                        cbd_value = cbd_test_result
                        self.logger.debug(f"Using database CBD test result: '{cbd_value}'")
                
                # Clean up values (remove 'nan', empty strings, etc.)
                if thc_value in ['nan', 'NaN', '']:
                    thc_value = '0'
                if cbd_value in ['nan', 'NaN', '']:
                    cbd_value = '0'
                
                self.logger.debug(f"THC/CBD values found - THC: '{thc_value}', CBD: '{cbd_value}'")
                
                # Format with actual values - always show both THC and CBD
                if thc_value and cbd_value:
                    result = f"THC: {thc_value}% CBD: {cbd_value}%"
                elif thc_value:
                    result = f"THC: {thc_value}% CBD: 0%"
                elif cbd_value:
                    result = f"THC: 0% CBD: {cbd_value}%"
                else:
                    result = f"THC: 0% CBD: 0%"
                
                self.logger.debug(f"Formatted THC/CBD result: '{result}'")
                return result
            
            # Fallback to default format if no record data or no values
            return "THC: 0% CBD: 0%"
        
        # If the text already contains THC/CBD format, return as-is
        if 'THC:' in text and 'CBD:' in text:
            return text
        
        # If the text contains mg values, return as-is (let text_processing handle it)
        if 'mg' in text.lower():
            return text
        
        # If the text contains simple ratios (like 1:1:1), format with spaces
        if ':' in text and any(c.isdigit() for c in text):
            # Add spaces around colons for better readability
            # Handle 3-part ratios first to avoid conflicts
            text = re.sub(r'(\d+):(\d+):(\d+)', r'\1: \2: \3', text)
            # Then handle 2-part ratios
            text = re.sub(r'(\d+):(\d+)', r'\1: \2', text)
            return text
        
        # Common patterns for THC/CBD ratios
        thc_patterns = [
            r'THC[:\s]*([0-9.]+)%?',
            r'([0-9.]+)%?\s*THC',
            r'([0-9.]+)\s*THC'
        ]
        
        cbd_patterns = [
            r'CBD[:\s]*([0-9.]+)%?',
            r'([0-9.]+)%?\s*CBD',
            r'([0-9.]+)\s*CBD'
        ]
        
        thc_value = None
        cbd_value = None
        
        # Extract THC value
        for pattern in thc_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                thc_value = match.group(1)
                break
        
        # Extract CBD value
        for pattern in cbd_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                cbd_value = match.group(1)
                break
        
        # If we found both values, format them
        if thc_value and cbd_value:
            # Keep on same line without line breaks
            formatted_text = f"THC: {thc_value}% CBD: {cbd_value}%"
            # Apply the new THC/CBD percentage formatting rules
            from src.core.generation.text_processing import format_thc_cbd_percentages
            return format_thc_cbd_percentages(formatted_text)
        elif thc_value:
            formatted_text = f"THC: {thc_value}%"
            from src.core.generation.text_processing import format_thc_cbd_percentages
            return format_thc_cbd_percentages(formatted_text)
        elif cbd_value:
            formatted_text = f"CBD: {cbd_value}%"
            from src.core.generation.text_processing import format_thc_cbd_percentages
            return format_thc_cbd_percentages(formatted_text)
        else:
            # If no clear THC/CBD pattern found, return the original text
            return text

    def format_joint_ratio_pack(self, text):
        """
        Format JointRatio as: [amount]g x [count] Pack
        Handles various input formats and normalizes them to standard format.
        For single units, shows just the weight (e.g., "1g" instead of "1g x 1 Pack").
        """
        if not text:
            return text
            
        # Convert to string and clean up
        text = str(text).strip()
        
        # Remove any leading/trailing spaces and hyphens
        text = re.sub(r'^[\s\-]+', '', text)
        text = re.sub(r'[\s\-]+$', '', text)
        
        # Handle various input patterns
        patterns = [
            # Standard format: "1g x 2 Pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Compact format: "1gx2Pack"
            r"([0-9.]+)g\s*x?\s*([0-9]+)pack",
            # With spaces: "1g x 2 pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Just weight: "1g"
            r"([0-9.]+)g",
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                amount = match.group(1).strip()
                # Try to get count, default to 1 if not found
                try:
                    count = match.group(2).strip()
                    if count and count.isdigit():
                        count_int = int(count)
                        if count_int == 1:
                            # For single units, just show the weight
                            formatted = f"{amount}g"
                        else:
                            # For multiple units, show the full pack format
                            formatted = f"{amount}g x {count} Pack"
                    else:
                        # Only amount found (like "1g") - show just the weight
                        formatted = f"{amount}g"
                except IndexError:
                    # Only amount found (like "1g") - show just the weight
                    formatted = f"{amount}g"
                return formatted
        
        # If no pattern matches, return the original text
        return text

    def format_thc_cbd_vertical_alignment(self, text):
        """
        Format THC_CBD content for vertical templates with right-aligned percentages.
        Splits THC and CBD into separate lines and right-aligns the percentage values.
        Adds extra line spacing between THC percentage and CBD line.
        """
        if not text:
            return text
        
        # Split into lines
        lines = text.split('\n')
        formatted_lines = []
        
        # Check if this contains THC/CBD content with percentages
        if not any('%' in line for line in lines):
            return text
        
        # First pass: collect all percentage values to determine maximum width
        all_percentages = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Extract percentage values using regex
            import re
            percentages = re.findall(r'([0-9.]+)%', line)
            all_percentages.extend(percentages)
        
        # Calculate maximum percentage width for right-alignment
        max_percentage_width = 0
        if all_percentages:
            max_percentage_width = max(len(percentage) for percentage in all_percentages)
        
        # Second pass: format each line with proper right-alignment
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line contains both THC and CBD (same line)
            if 'THC:' in line and 'CBD:' in line and '%' in line:
                # Split the line into THC and CBD parts
                cbd_start = line.find('CBD:')
                thc_part = line[:cbd_start].strip()
                cbd_part = line[cbd_start:].strip()
                
                # Check if there are other cannabinoids after CBD
                remaining_content = ''
                if 'CBC:' in cbd_part:
                    cbc_start = cbd_part.find('CBC:')
                    cbd_part_only = cbd_part[:cbc_start].strip()
                    remaining_content = cbd_part[cbc_start:].strip()
                    cbd_part = cbd_part_only
                elif 'CBG:' in cbd_part:
                    cbg_start = cbd_part.find('CBG:')
                    cbd_part_only = cbd_part[:cbg_start].strip()
                    remaining_content = cbd_part[cbg_start:].strip()
                    cbd_part = cbd_part_only
                
                # Format THC part with simple formatting
                formatted_thc = self._format_thc_cbd_simple(thc_part, max_percentage_width)
                
                # Format CBD part with simple formatting
                formatted_cbd = self._format_thc_cbd_simple(cbd_part, max_percentage_width)
                
                # Combine with line break between THC and CBD
                if remaining_content:
                    # Format remaining content (like CBC) with right-alignment
                    formatted_remaining = self._format_percentage_right_alignment(remaining_content, max_percentage_width)
                    # Add line break between THC and CBD
                    formatted_line = f"{formatted_thc}\n{formatted_cbd}\n{formatted_remaining}"
                else:
                    # Add line break between THC and CBD
                    formatted_line = f"{formatted_thc}\n{formatted_cbd}"
                formatted_lines.append(formatted_line)
            else:
                # For single cannabinoid lines, apply right-alignment
                formatted_line = self._format_percentage_right_alignment(line, max_percentage_width)
                formatted_lines.append(formatted_line)
        
        return '\n'.join(formatted_lines)
    
    def _format_thc_cbd_simple(self, text, max_percentage_width):
        """
        Helper function to format THC/CBD with simple line break between values.
        Returns format: "THC: x%\nCBD: x%"
        """
        if not text or '%' not in text:
            return text
        
        import re
        
        # Split the text into parts: label, percentage, and any remaining text
        # Pattern to match: "THC: " + percentage + "%" + remaining
        match = re.match(r'^([^0-9]*?)([0-9.]+)%(.*)$', text)
        if not match:
            return text
        
        label = match.group(1).strip()  # e.g., "THC:"
        percentage = match.group(2)  # e.g., "21.0"
        remaining = match.group(3)  # e.g., " CBD: 0.25%"
        
        # Simple format: label and percentage on same line
        formatted_group = f"{label} {percentage}%"
        
        # Add remaining content if any
        if remaining.strip():
            formatted_group += f"\n{remaining.strip()}"
        
        return formatted_group

    def _identify_marker_type(self, text):
        """Identify the marker type from text content for proper font sizing."""
        if not text:
            return 'default'
            
        text_upper = text.upper()
        
        # Check for specific content patterns in order of specificity
        if any(word in text_upper for word in ['THC', 'CBD', 'RATIO', ':', '%']):
            return 'THC_CBD'
        elif any(word in text_upper for word in ['SATIVA', 'INDICA', 'HYBRID', 'MIXED', 'PARA']):
            return 'LINEAGE'
        elif any(word in text_upper for word in ['$', 'PRICE', 'COST']):
            return 'PRICE'
        elif any(word in text_upper for word in ['WEIGHT', 'G', 'OZ', 'LB', 'KG']) and not any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY']):
            return 'WEIGHT'
        elif any(word in text_upper for word in ['DOH', 'DATE', 'EXP', '/']) and len(text) <= 12:
            return 'DOH'
        elif any(word in text_upper for word in ['VENDOR', 'SUPPLIER']):
            return 'VENDOR'
        elif any(word in text_upper for word in ['STRAIN', 'VARIETY']) or (len(text.split()) <= 2 and len(text) <= 15 and not any(char in text for char in ['$', '%', ':', '/']) and not any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY'])):
            return 'STRAIN'
        elif any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY']) or (len(text.split()) <= 3 and len(text) <= 25 and not any(word in text_upper for word in ['DESCRIPTION', 'LONG', 'DETAILED'])):
            return 'BRAND'
        else:
            # Check if it looks like a description (longer text, multiple words)
            if len(text.split()) > 3 or len(text) > 25:
                return 'DESCRIPTION'
            else:
                return 'default'

    def _apply_mini_template_font_sizing(self, doc):
        """Apply mini template specific font sizing to all content."""
        try:
            from src.core.generation.unified_font_sizing import get_font_size_by_marker, set_run_font_size
            
            # Process all tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    # Identify the marker type from the text content for proper font sizing
                                    marker_type = self._identify_marker_type(run.text)
                                    
                                    # Get appropriate font size for mini template using proper field type
                                    font_size = get_font_size_by_marker(
                                        run.text, 
                                        marker_type, 
                                        template_type='mini', 
                                        scale_factor=self.scale_factor
                                    )
                                    set_run_font_size(run, font_size)
                                    
            self.logger.info("Applied mini template specific font sizing with proper field type identification")
            
        except Exception as e:
            self.logger.warning(f"Error applying mini template font sizing: {e}")

    def _format_mini_template_text(self, text):
        """Format text for mini template to prevent improper line breaks."""
        if not text:
            return text
        
        # For mini templates, we want to prevent line breaks in the middle of weight units
        # Replace spaces around hyphens with non-breaking spaces to keep " - 1g" together
        # But allow breaks between the main description and the weight unit
        text = text.replace(' - ', ' -\u00A0')  # Non-breaking space after hyphen
        
        return text

    def _format_percentage_right_alignment(self, text, max_percentage_width):
        """
        Helper function to right-align percentage values in a single line.
        """
        if not text or '%' not in text:
            return text
        
        import re
        
        # Split the text into parts: label, percentage, and any remaining text
        # Pattern to match: "THC: " + percentage + "%" + remaining
        match = re.match(r'^([^0-9]*?)([0-9.]+)%(.*)$', text)
        if not match:
            return text
        
        label = match.group(1)  # e.g., "THC: "
        percentage = match.group(2)  # e.g., "21.0"
        remaining = match.group(3)  # e.g., " CBD: 0.25%"
        
        # Calculate spacing needed for right-alignment
        spacing_needed = max_percentage_width - len(percentage)
        spaces = ' ' * max(0, spacing_needed)
        
        # Return the formatted string with proper spacing
        return f"{label}{spaces}{percentage}%{remaining}"

    def _process_combined_lineage_vendor(self, paragraph, lineage_content, vendor_content):
        """
        Process combined lineage and vendor text with different font sizes.
        This handles the case where lineage and product vendor are on the same line.
        Lineage is left-aligned, vendor is right-aligned.
        IMPORTANT: Product Vendor should never be split up - if Lineage is too long, it should break to new line.
        SPECIAL RULE: For Vertical template, if Lineage is "Hybrid/Indica" or "Hybrid/Sativa", automatically put ProductVendor on next line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # Ensure consistent spacing above lineage/vendor section for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # SPECIAL RULE: For Vertical template, automatically force vendor to next line for specific lineages
            if (self.template_type == 'vertical' and 
                lineage_content and 
                lineage_content.strip().upper() in ['HYBRID/INDICA', 'HYBRID/SATIVA'] and
                vendor_content and vendor_content.strip()):
                
                self.logger.debug(f"Vertical template: Forcing vendor to next line for lineage '{lineage_content}'")
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Check if we need to split to multiple lines due to content length
            # Calculate approximate character limits based on template type
            if self.template_type == 'mini':
                max_chars_per_line = 25
            elif self.template_type == 'vertical':
                max_chars_per_line = 35
            else:  # horizontal, double
                max_chars_per_line = 45
            
            # Check if combined content would be too long for one line
            combined_length = len(lineage_content or '') + len(vendor_content or '')
            
            if combined_length > max_chars_per_line and vendor_content and vendor_content.strip():
                # Split to two lines: lineage on first line, vendor on second line
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Original single-line processing
            # Set paragraph to justified alignment to allow for right-aligned vendor
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add lineage with larger font size (left-aligned)
            if lineage_content and lineage_content.strip():
                # Debug: Log the lineage content to see what we're working with
                self.logger.debug(f"DEBUG: Original lineage_content: '{repr(lineage_content)}'")
                
                # Ensure no leading spaces in lineage content - aggressive cleaning
                clean_lineage = lineage_content.strip().lstrip().lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                self.logger.debug(f"DEBUG: Cleaned lineage_content: '{repr(clean_lineage)}'")
                
                lineage_run = paragraph.add_run(clean_lineage)
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Get lineage font size
                product_type = None
                if hasattr(self, 'current_product_type'):
                    product_type = self.current_product_type
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                
                lineage_font_size = get_font_size_by_marker(lineage_content, 'LINEAGE', self.template_type, self.scale_factor, product_type)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add tab character to push vendor to the right (only if vendor content exists)
            if lineage_content and vendor_content:
                tab_run = paragraph.add_run("\t")
                tab_run.font.name = "Arial"
                tab_run.font.bold = True
                # Use lineage font size for tab to maintain alignment
                set_run_font_size(tab_run, lineage_font_size)
            
            # Add vendor with smaller font size (right-aligned)
            if vendor_content and vendor_content.strip():
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                vendor_run.font.bold = False
                vendor_run.font.italic = True  # Make vendor text italic
                
                # Set vendor color to light gray (#CCCCCC)
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Ensure the color is applied by setting it explicitly
                vendor_run.font.color.theme_color = None  # Clear any theme color
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Get vendor font size using unified font sizing system
                from src.core.generation.unified_font_sizing import get_font_size
                vendor_font_size = get_font_size(vendor_content, 'vendor', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)
            
            # Set tab stops to position vendor on the right (only if vendor content exists)
            if vendor_content:
                # Clear existing tab stops
                paragraph.paragraph_format.tab_stops.clear_all()
                # Add right-aligned tab stop at the right margin - positioned further right for full justification
                if self.template_type == 'mini':
                    tab_position = Inches(1.2)  # Adjusted for 1.5x1.5 mini template
                elif self.template_type == 'vertical':
                    tab_position = Inches(2.3)  # Increased for vertical template
                else:  # horizontal, double
                    tab_position = Inches(3.2)  # Further increased for horizontal/double templates
                
                paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                
                # Alternative: Use multiple tab stops for more aggressive right positioning
                # This creates additional tab stops to ensure the vendor text reaches the right edge
                if self.template_type in ['horizontal', 'double']:
                    # Add an additional tab stop even further right as backup
                    backup_tab_position = Inches(3.5)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
            else:
                # For non-classic products without vendor, use left alignment
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0)
            
            self.logger.debug(f"Processed combined lineage/vendor with right-aligned vendor: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing combined lineage/vendor: {e}")
            # Fallback: use default processing
            paragraph.clear()
            # Don't strip leading spaces for LINEAGE to preserve our spacing fix
            combined_text = f"{lineage_content or ''}  {vendor_content or ''}".rstrip()
            if combined_text:
                run = paragraph.add_run(combined_text)

    def _process_lineage_vendor_two_lines(self, paragraph, lineage_content, vendor_content):
        """
        Process lineage and vendor on two separate lines to prevent vendor splitting.
        Lineage goes on the first line, vendor goes on the second line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # Ensure consistent spacing above lineage/vendor section for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # Set paragraph to left alignment for two-line layout
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add lineage on first line with larger font size
            if lineage_content and lineage_content.strip():
                # Ensure no leading spaces in lineage content - aggressive cleaning
                clean_lineage = lineage_content.strip().lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                lineage_run = paragraph.add_run(clean_lineage)
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Get lineage font size
                product_type = None
                if hasattr(self, 'current_product_type'):
                    product_type = self.current_product_type
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                
                lineage_font_size = get_font_size_by_marker(lineage_content, 'LINEAGE', self.template_type, self.scale_factor, product_type)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add line break
            if lineage_content and vendor_content:
                paragraph.add_run("\n")
            
            # Add vendor on second line with smaller font size
            if vendor_content and vendor_content.strip():
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                vendor_run.font.bold = False
                vendor_run.font.italic = True  # Make vendor text italic
                
                # Set vendor color to light gray (#CCCCCC)
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Ensure the color is applied by setting it explicitly
                vendor_run.font.color.theme_color = None  # Clear any theme color
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Get vendor font size using unified font sizing system
                from src.core.generation.unified_font_sizing import get_font_size
                vendor_font_size = get_font_size(vendor_content, 'vendor', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0)
            
            self.logger.debug(f"Processed lineage/vendor on two lines: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing lineage/vendor on two lines: {e}")
            # Fallback: use single line processing
            self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)

    def _detect_and_process_combined_lineage_vendor(self, paragraph):
        """
        Detect if paragraph contains combined lineage and vendor markers and process them separately.
        Remove vendor for non-classic product types.
        """
        # Check if this paragraph has already been processed for combined lineage/vendor
        if hasattr(paragraph, '_combined_lineage_vendor_processed'):
            return True
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Check if both lineage and vendor markers are present
        lineage_start = "LINEAGE_START"
        lineage_end = "LINEAGE_END"
        vendor_start = "PRODUCTVENDOR_START"
        vendor_end = "PRODUCTVENDOR_END"
        
        if (lineage_start in full_text and lineage_end in full_text and 
            vendor_start in full_text and vendor_end in full_text):
            
            try:
                # Extract lineage content
                lineage_start_idx = full_text.find(lineage_start) + len(lineage_start)
                lineage_end_idx = full_text.find(lineage_end)
                lineage_content = full_text[lineage_start_idx:lineage_end_idx].strip()
                
                # Extract vendor content
                vendor_start_idx = full_text.find(vendor_start) + len(vendor_start)
                vendor_end_idx = full_text.find(vendor_end)
                vendor_content = full_text[vendor_start_idx:vendor_end_idx]
                
                # Note: Product type filtering is now handled in _build_label_context
                # This method only processes the content that's already been filtered
                
                # Process with different font sizes
                self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)
                
                # Mark this paragraph as processed to prevent re-processing
                paragraph._combined_lineage_vendor_processed = True
                
                return True
                
            except Exception as e:
                self.logger.error(f"Error detecting combined lineage/vendor: {e}")
                return False
        
        return False

    def _fix_productstrain_in_brand_cells(self, doc):
        """Fix ProductStrain appearing in ProductBrand cells for non-classic types."""
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Check if this paragraph contains both ProductBrand and ProductStrain content
                            full_text = "".join(run.text for run in paragraph.runs)
                            
                            # Look for ProductStrain markers in ProductBrand cells (legacy approach)
                            if ('PRODUCTSTRAIN_START' in full_text and 
                                ('PRODUCTBRAND_CENTER_START' in full_text or 'PRODUCTBRAND_CENTER_END' in full_text)):
                                self.logger.debug(f"Found ProductStrain markers in ProductBrand cell, fixing...")
                                self._fix_productstrain_markers_in_brand_cells(paragraph, full_text)
                            
                            # NEW: Look for actual ProductStrain content appearing in cells with ProductBrand
                            # This handles the real-world scenario where markers are removed during rendering
                            elif self._detect_productstrain_in_brand_content(full_text):
                                self.logger.debug(f"Found ProductStrain content in ProductBrand cell, fixing...")
                                self._fix_productstrain_content_in_brand_cells(paragraph, full_text)
                                    
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain in Brand cells: {e}")
    
    def _detect_productstrain_in_brand_content(self, text):
        """Detect if ProductStrain content appears in the same cell as ProductBrand content."""
        # CRITICAL FIX: Don't interfere with Lineage field content
        # The Lineage field is SUPPOSED to contain strain content (INDICA, SATIVA, HYBRID) 
        # alongside vendor names for classic types - this is correct behavior
        
        # Only apply this fix if we're in a ProductBrand-specific context
        # Check if this is actually a ProductBrand cell, not a Lineage cell
        if 'PRODUCTBRAND' in text.upper() or 'BRAND' in text.upper():
            # Common strain values that shouldn't appear in brand cells
            common_strains = [
                'HYBRID', 'INDICA', 'SATIVA', 'MIXED', 'CBD', 'CBD BLEND', 'PARAPHERNALIA', 'PARA'
            ]
            
            # Check if any strain content appears in the text
            for strain in common_strains:
                if strain in text.upper():
                    # Also check if there's likely brand content (words that look like brand names)
                    # Brand names typically have mixed case, multiple words, and may contain LLC, Inc, etc.
                    brand_indicators = ['LLC', 'INC', 'CORP', 'CO', 'COMPANY', 'BRANDS', 'BRAND']
                    has_brand_indicators = any(indicator in text.upper() for indicator in brand_indicators)
                    
                    # If we have both strain content and brand indicators, this needs fixing
                    if has_brand_indicators:
                        return True
        
        return False
    
    def _fix_productstrain_content_in_brand_cells(self, paragraph, full_text):
        """Fix ProductStrain content appearing in ProductBrand cells by removing strain content."""
        try:
            # Common strain values to remove
            common_strains = [
                'HYBRID', 'INDICA', 'SATIVA', 'MIXED', 'CBD', 'CBD BLEND', 'PARAPHERNALIA', 'PARA'
            ]
            
            # Find and remove strain content
            new_text = full_text
            removed_strains = []
            
            for strain in common_strains:
                if strain.upper() in new_text.upper():
                    # Remove the strain content (case-insensitive)
                    strain_pattern = re.compile(re.escape(strain), re.IGNORECASE)
                    new_text = strain_pattern.sub('', new_text)
                    removed_strains.append(strain)
            
            # Clean up extra whitespace and line breaks
            new_text = re.sub(r'\n\s*\n', '\n', new_text)  # Remove double line breaks
            new_text = new_text.strip()  # Remove leading/trailing whitespace
            
            if removed_strains:
                self.logger.debug(f"Removed ProductStrain content: {removed_strains}")
                self.logger.debug(f"Original text: '{full_text}'")
                self.logger.debug(f"Cleaned text: '{new_text}'")
                
                # Update the paragraph with cleaned text
                paragraph.clear()
                run = paragraph.add_run()
                run.text = new_text
                run.font.name = "Arial"
                run.font.bold = True
                
                # Apply appropriate font sizing for ProductBrand using the unified font sizing
                from src.core.generation.unified_font_sizing import get_font_size
                font_size = get_font_size(new_text, 'brand', self.template_type, self.scale_factor)
                run.font.size = font_size
                
                return True
            
            return False
            
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain content in brand cells: {e}")
            return False
    
    def _fix_productstrain_markers_in_brand_cells(self, paragraph, full_text):
        """Legacy method to fix ProductStrain markers in ProductBrand cells."""
        try:
            # Extract the ProductStrain content
            strain_start = full_text.find('PRODUCTSTRAIN_START')
            strain_end = full_text.find('PRODUCTSTRAIN_END')
            
            if strain_start >= 0 and strain_end >= 0:
                # Extract the strain content (without markers)
                strain_content_start = strain_start + len('PRODUCTSTRAIN_START')
                strain_content = full_text[strain_content_start:strain_end]
                
                # Extract the ProductBrand content (without ProductStrain)
                brand_start = full_text.find('PRODUCTBRAND_CENTER_START')
                brand_end = full_text.find('PRODUCTBRAND_CENTER_END')
                
                if brand_start >= 0 and brand_end >= 0:
                    # Get the brand content between the markers
                    brand_content_start = brand_start + len('PRODUCTBRAND_CENTER_START')
                    brand_content = full_text[brand_content_start:brand_end]
                    
                    # Remove the ProductStrain content from this paragraph
                    # Keep only the ProductBrand content
                    new_text = full_text[:strain_start] + full_text[strain_end + len('PRODUCTSTRAIN_END'):]
                    
                    # Clear the paragraph and recreate with just the ProductBrand content
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.text = new_text
                    run.font.name = "Arial"
                    run.font.bold = True
                    
                    # Apply appropriate font sizing for ProductBrand
                    from src.core.generation.unified_font_sizing import get_font_size_by_marker
                    font_size = get_font_size_by_marker(new_text, 'PRODUCTBRAND_CENTER', self.template_type, self.scale_factor)
                    if font_size:
                        run.font.size = font_size
                    
                    self.logger.debug(f"Separated ProductStrain '{strain_content}' from ProductBrand cell")
                    
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain markers in brand cells: {e}")

    def _clean_doh_cells_before_processing(self, doc):
        """
        Clean up DOH cells before processing to ensure no content interferes with image positioning.
        This should be called before DOH images are inserted.
        """
        try:
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains DOH placeholder
                        cell_text = cell.text.strip()
                        if '{{Label' in cell_text and '.DOH}}' in cell_text:
                            # Clear the cell content to prepare for image insertion
                            cell._tc.clear_content()
                            
                            # Add a single empty paragraph to maintain cell structure
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set minimal spacing
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            self.logger.debug("Cleaned DOH cell for image insertion")
                            
        except Exception as e:
            self.logger.warning(f"Error cleaning DOH cells: {e}")

    def _replace_qr_placeholder(self, doc, qr_inline_image, label_key="Label1"):
        """Replace QR placeholder with actual QR code image for specific label."""
        try:
            self.logger.debug(f"Starting QR placeholder replacement for {label_key}")
            
            # Find and replace QR_PLACEHOLDER text with actual QR code
            # We'll replace the first QR_PLACEHOLDER we find for this label
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if 'QR_PLACEHOLDER' in para.text:
                                self.logger.debug(f"Found QR placeholder in paragraph for {label_key}")
                                
                                # Clear the paragraph and add the QR code
                                para.clear()
                                para.alignment = 1  # Center alignment
                                
                                # Add the QR code image using raw image data
                                run = para.add_run()
                                from io import BytesIO
                                img_buffer = BytesIO(qr_inline_image._raw_image_data)
                                run.add_picture(img_buffer, width=qr_inline_image._raw_image_width)
                                
                                self.logger.debug(f"✓ QR code replaced successfully for {label_key}")
                                return True
            
            self.logger.warning(f"QR placeholder not found in document for {label_key}")
            return False
            
        except Exception as e:
            self.logger.error(f"Error replacing QR placeholder: {e}")
            return False

    def _manual_replace_placeholders(self, doc, context):
        """Manually replace placeholders in the document when DocxTemplate fails."""
        try:
            self.logger.info("Starting manual placeholder replacement")
            
            # Debug: Log the context structure
            self.logger.info(f"Context keys: {list(context.keys())}")
            for label_key, label_context in context.items():
                if isinstance(label_context, dict):
                    self.logger.info(f"{label_key} fields: {list(label_context.keys())}")
                    # Log a few sample values
                    for field_key, field_value in list(label_context.items())[:5]:
                        self.logger.info(f"  {field_key}: {repr(str(field_value)[:100])}")
            
            # Process each table in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Process runs for placeholder replacement
                            for run in paragraph.runs:
                                text = run.text
                                for label_key, label_context in context.items():
                                    for field_key, field_value in label_context.items():
                                        placeholder = f"{{{{{label_key}.{field_key}}}}}"
                                        if placeholder in text:
                                            # Special handling for QR codes (InlineImage objects)
                                            if field_key == 'QR' and hasattr(field_value, 'image'):
                                                # Clear the run and add the QR code image
                                                run.clear()
                                                run.add_picture(field_value.image, width=field_value.width)
                                                self.logger.debug(f"Replaced {placeholder} with QR code image")
                                                continue
                                            
                                            # Handle simple QR placeholder {{QR}}
                                            if placeholder == '{{QR}}' and field_key == 'QR' and hasattr(field_value, 'image'):
                                                # Clear the run and add the QR code image
                                                run.clear()
                                                run.add_picture(field_value.image, width=field_value.width)
                                                self.logger.debug(f"Replaced simple {{QR}} placeholder with QR code image")
                                                continue
                                            
                                            # For all templates, unwrap markers to get clean content
                                            if True:  # Apply to all templates
                                                # Unwrap common markers to get clean content
                                                clean_value = str(field_value)
                                                if 'DESC_START' in clean_value and 'DESC_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'DESC')
                                                elif 'PRICE_START' in clean_value and 'PRICE_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'PRICE')
                                                elif 'THC_CBD_START' in clean_value and 'THC_CBD_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'THC_CBD')
                                                elif 'RATIO_START' in clean_value and 'RATIO_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'RATIO')
                                                elif 'PRODUCTNAME_START' in clean_value and 'PRODUCTNAME_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'PRODUCTNAME')
                                                elif 'PRODUCTBRAND_START' in clean_value and 'PRODUCTBRAND_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'PRODUCTBRAND')
                                                elif 'WEIGHTUNITS_START' in clean_value and 'WEIGHTUNITS_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'WEIGHTUNITS')
                                                elif 'PRODUCTVENDOR_START' in clean_value and 'PRODUCTVENDOR_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'PRODUCTVENDOR')
                                                elif 'PRODUCTSTRAIN_START' in clean_value and 'PRODUCTSTRAIN_END' in clean_value:
                                                    clean_value = unwrap_marker(clean_value, 'PRODUCTSTRAIN')
                                                
                                                # Replace placeholder with clean value
                                                text = text.replace(placeholder, clean_value)
                                                self.logger.debug(f"Replaced {placeholder} with unwrapped value: {clean_value}")
                                                
                                                # Apply unified font sizing for double templates
                                                field_type_mapping = {
                                                    'DescAndWeight': 'description',
                                                    'Price': 'price',
                                                    'Ratio_or_THC_CBD': 'thc_cbd',
                                                    'ProductBrand': 'brand',
                                                    'Lineage': 'lineage',
                                                    'ProductStrain': 'strain',
                                                    'ProductVendor': 'vendor',
                                                    'DOH': 'doh',
                                                    'QR': 'qr'
                                                }
                                                
                                                field_type = field_type_mapping.get(field_key, 'default')
                                                from src.core.generation.unified_font_sizing import get_font_size
                                                font_size = get_font_size(clean_value, field_type, self.template_type, self.scale_factor)
                                                set_run_font_size(run, font_size)
                                                self.logger.debug(f"Applied font sizing: {field_key} -> {field_type} -> {font_size.pt}pt")
                                            else:
                                                # For other templates, use the original value
                                                text = text.replace(placeholder, str(field_value))
                                                self.logger.debug(f"Replaced {placeholder} with {field_value}")
                                
                                run.text = text
                            
                            # Also check paragraph text for remaining placeholders
                            paragraph_text = paragraph.text
                            for label_key, label_context in context.items():
                                for field_key, field_value in label_context.items():
                                    placeholder = f"{{{{{label_key}.{field_key}}}}}"
                                    if placeholder in paragraph_text:
                                        # Special handling for QR codes (InlineImage objects)
                                        if field_key == 'QR' and hasattr(field_value, 'image'):
                                            # Clear the paragraph and add the QR code image
                                            paragraph.clear()
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            paragraph.add_run().add_picture(field_value.image, width=field_value.width)
                                            self.logger.debug(f"Replaced {placeholder} with QR code image in paragraph")
                                            continue
                                        
                                        # Handle simple QR placeholder {{QR}}
                                        if placeholder == '{{QR}}' and field_key == 'QR' and hasattr(field_value, 'image'):
                                            # Clear the paragraph and add the QR code image
                                            paragraph.clear()
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            paragraph.add_run().add_picture(field_value.image, width=field_value.width)
                                            self.logger.debug(f"Replaced simple {{QR}} placeholder with QR code image in paragraph")
                                            continue
                                        
                                        # For all templates, unwrap markers to get clean content
                                        if True:  # Apply to all templates
                                            # Unwrap common markers to get clean content
                                            clean_value = str(field_value)
                                            if 'DESC_START' in clean_value and 'DESC_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'DESC')
                                            elif 'PRICE_START' in clean_value and 'PRICE_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'PRICE')
                                            elif 'THC_CBD_START' in clean_value and 'THC_CBD_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'THC_CBD')
                                            elif 'RATIO_START' in clean_value and 'RATIO_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'RATIO')
                                            elif 'PRODUCTNAME_START' in clean_value and 'PRODUCTNAME_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'PRODUCTNAME')
                                            elif 'PRODUCTBRAND_START' in clean_value and 'PRODUCTBRAND_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'PRODUCTBRAND')
                                            elif 'WEIGHTUNITS_START' in clean_value and 'WEIGHTUNITS_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'WEIGHTUNITS')
                                            elif 'PRODUCTVENDOR_START' in clean_value and 'PRODUCTVENDOR_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'PRODUCTVENDOR')
                                            elif 'PRODUCTSTRAIN_START' in clean_value and 'PRODUCTSTRAIN_END' in clean_value:
                                                clean_value = unwrap_marker(clean_value, 'PRODUCTSTRAIN')
                                            
                                            # Replace placeholder with clean value
                                            paragraph_text = paragraph_text.replace(placeholder, clean_value)
                                            self.logger.debug(f"Replaced {placeholder} with unwrapped value in paragraph: {clean_value}")
                                            
                                            # Apply unified font sizing for double templates
                                            field_type_mapping = {
                                                'DescAndWeight': 'description',
                                                'Price': 'price',
                                                'Ratio_or_THC_CBD': 'thc_cbd',
                                                'ProductBrand': 'brand',
                                                'Lineage': 'lineage',
                                                'ProductStrain': 'strain',
                                                'ProductVendor': 'vendor',
                                                'DOH': 'doh',
                                                'QR': 'qr'
                                            }
                                            
                                            field_type = field_type_mapping.get(field_key, 'default')
                                            from src.core.generation.unified_font_sizing import get_font_size
                                            font_size = get_font_size(clean_value, field_type, self.template_type, self.scale_factor)
                                            
                                            # Apply font sizing to all runs in the paragraph
                                            for run in paragraph.runs:
                                                set_run_font_size(run, font_size)
                                            self.logger.debug(f"Applied font sizing to paragraph: {field_key} -> {field_type} -> {font_size.pt}pt")
                                        else:
                                            # For other templates, use the original value
                                            paragraph_text = paragraph_text.replace(placeholder, str(field_value))
                                            self.logger.debug(f"Replaced {placeholder} with {field_value} in paragraph")
                            
                            paragraph.text = paragraph_text
            
            self.logger.info("Manual placeholder replacement completed")
            
        except Exception as e:
            self.logger.error(f"Error during manual placeholder replacement: {e}")
            raise
    
    def _cell_contains_doh_image(self, cell):
        """
        Check if a cell contains a DOH image.
        This is used to preserve center alignment for DOH images when setting other content to TOP alignment.
        
        Args:
            cell: The table cell to check
            
        Returns:
            bool: True if the cell contains a DOH image, False otherwise
        """
        # Use the common utility function
        from src.core.utils.common import cell_contains_doh_image
        return cell_contains_doh_image(cell)

    def _final_doh_positioning_enforcement(self, doc):
        """
        Final enforcement pass to ensure DOH images are ALWAYS centered.
        This method runs at the very end and overrides any other positioning logic.
        It addresses the issue where Advanced Layout gets set to "top" instead of "center".
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            doh_cells_fixed = 0
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        doh_paragraphs = []
                        
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element'):
                                    # Check for drawing elements (InlineImage) or picture elements
                                    if (run._element.find(qn('w:drawing')) is not None or 
                                        run._element.find(qn('w:pict')) is not None):
                                        has_doh_image = True
                                        doh_paragraphs.append(paragraph)
                                        break
                            if has_doh_image:
                                break
                        
                        if has_doh_image:
                            doh_cells_fixed += 1
                            
                            # FORCE cell vertical alignment to CENTER
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # FORCE paragraph alignment to CENTER for all DOH image paragraphs
                            for paragraph in doh_paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # FORCE XML-level centering to override any Word template defaults
                                try:
                                    pPr = paragraph._element.get_or_add_pPr()
                                    
                                    # Remove any existing justification and set to center
                                    existing_jc = pPr.find(qn('w:jc'))
                                    if existing_jc is not None:
                                        pPr.remove(existing_jc)
                                    
                                    jc = OxmlElement('w:jc')
                                    jc.set(qn('w:val'), 'center')
                                    pPr.append(jc)
                                    
                                    # Ensure proper spacing
                                    existing_spacing = pPr.find(qn('w:spacing'))
                                    if existing_spacing is not None:
                                        pPr.remove(existing_spacing)
                                    
                                    spacing = OxmlElement('w:spacing')
                                    spacing.set(qn('w:before'), '60')  # 3pt before
                                    spacing.set(qn('w:after'), '60')   # 3pt after
                                    spacing.set(qn('w:line'), '240')   # Single line spacing
                                    spacing.set(qn('w:lineRule'), 'auto')
                                    pPr.append(spacing)
                                    
                                    # Remove any indentation that might affect centering
                                    existing_ind = pPr.find(qn('w:ind'))
                                    if existing_ind is not None:
                                        pPr.remove(existing_ind)
                                    
                                except Exception as xml_error:
                                    self.logger.warning(f"Error applying XML-level DOH centering: {xml_error}")
            
            if doh_cells_fixed > 0:
                self.logger.info(f"Final DOH positioning enforcement: Fixed {doh_cells_fixed} DOH image cells to ensure CENTER alignment")
                
        except Exception as e:
            self.logger.warning(f"Error in final DOH positioning enforcement: {e}")

__all__ = ['get_font_scheme', 'TemplateProcessor']