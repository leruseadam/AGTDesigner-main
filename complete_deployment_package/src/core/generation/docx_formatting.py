from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

# Define colors for lineage
COLORS = {
    'SATIVA': 'ED4123',
    'INDICA': '9900FF',
    'HYBRID': '009900',
    'HYBRID_INDICA': '9900FF',
    'HYBRID_SATIVA': 'ED4123',
    'CBD': 'F1C232',
    'CBD_BLEND': 'F1C232',  # Same color as CBD
    'MIXED': '0021F5',
    'PARA': 'FFC0CB'
}

def apply_lineage_colors(doc):
    """Apply lineage colors to all cells based on keywords in cell text."""
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.upper()
                    color_hex = None
                    
                    # Remove marker wrappers for robust matching
                    for marker in ["LINEAGE_START", "LINEAGE_END", "PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END", "PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END"]:
                        text = text.replace(marker, "")
                    text = text.strip()
                    
                    # Apply lineage coloring logic based on clean text
                    if "PARAPHERNALIA" in text:
                        color_hex = COLORS['PARA']
                    elif "HYBRID/INDICA" in text or "HYBRID INDICA" in text:
                        color_hex = COLORS['HYBRID_INDICA']
                    elif "HYBRID/SATIVA" in text or "HYBRID SATIVA" in text:
                        color_hex = COLORS['HYBRID_SATIVA']
                    elif "SATIVA" in text:
                        color_hex = COLORS['SATIVA']
                    elif "INDICA" in text:
                        color_hex = COLORS['INDICA']
                    elif "HYBRID" in text:
                        color_hex = COLORS['HYBRID']
                    elif "CBD" in text or "CBD_BLEND" in text:
                        color_hex = COLORS['CBD']
                    elif "CBD BLEND" in text:
                        color_hex = COLORS['CBD_BLEND']
                    elif "MIXED" in text:
                        # MIXED lineage always gets blue bars (this covers non-classic types like edibles)
                        color_hex = COLORS['MIXED']  # Blue for Mixed
                    elif any(brand in text.upper() for brand in ["MOONSHOT", "PLATINUM", "PREMIUM", "GOLD", "SILVER", "ELITE", "SELECT", "RESERVE", "CRAFT", "ARTISAN", "BOUTIQUE", "SIGNATURE", "LIMITED", "EXCLUSIVE", "PRIVATE", "CUSTOM", "SPECIAL", "DELUXE", "ULTRA", "SUPER", "MEGA", "MAX", "PRO", "PLUS", "X", "CONSTELLATION"]):
                        # Product Brand values (like "MOONSHOT", "PLATINUM", "CONSTELLATION", etc.) get blue bars for non-classic types
                        color_hex = COLORS['MIXED']  # Blue for Product Brand
                    
                    if color_hex:
                        # Set cell background color
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        for old_shd in tcPr.findall(qn('w:shd')):
                            tcPr.remove(old_shd)
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), color_hex)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        tcPr.append(shd)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold = True
                                run.font.name = "Arial"
        # FINAL LINEAGE CLEANUP: Remove any leading spaces from lineage content after coloring
        _final_lineage_cleanup_after_coloring(doc)
        
        logger.debug("Applied lineage colors to document")
        return doc
    except Exception as e:
        logger.error(f"Error applying lineage colors: {str(e)}")
        raise

def _final_lineage_cleanup_after_coloring(doc):
    """
    Final cleanup to remove any leading spaces from lineage content after coloring is applied.
    This runs after all other processing to ensure clean lineage display.
    """
    try:
        # Define lineage values that should be cleaned
        lineage_values = [
            "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
            "CBD", "CBD BLEND", "MIXED", "PARAPHERNALIA", "PARA"
        ]
        
        # Clean lineage content in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            original_text = run.text
                            
                            # Check if this run contains lineage content
                            for lineage in lineage_values:
                                if lineage in original_text.upper():
                                    # Aggressively clean leading spaces
                                    cleaned_text = original_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                                    
                                    if cleaned_text != original_text:
                                        run.text = cleaned_text
                                        logger.debug(f"Final lineage cleanup after coloring: '{original_text}' -> '{cleaned_text}'")
                                    break
        
        # Clean lineage content in paragraphs outside tables
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                original_text = run.text
                
                # Check if this run contains lineage content
                for lineage in lineage_values:
                    if lineage in original_text.upper():
                        # Aggressively clean leading spaces
                        cleaned_text = original_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                        
                        if cleaned_text != original_text:
                            run.text = cleaned_text
                            logger.debug(f"Final lineage cleanup after coloring: '{original_text}' -> '{cleaned_text}'")
                        break
        
        logger.debug("Final lineage cleanup after coloring completed - all leading spaces should be removed")
        
    except Exception as e:
        logger.warning(f"Error in final lineage cleanup after coloring: {e}")

def fix_table_row_heights(doc, template_type):
    """Fix table row heights based on template type."""
    try:
        row_height = {
            'horizontal': 2.4,
            'vertical': 3.4,
            'mini': 1.5,
            'inventory': 2.0
        }.get(template_type, 2.4)
        for table in doc.tables:
            for row in table.rows:
                row.height = Inches(row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        logger.debug(f"Fixed table row heights for template type: {template_type}")
        return doc
    except Exception as e:
        logger.error(f"Error fixing table row heights: {str(e)}")
        raise

def safe_fix_paragraph_spacing(doc):
    """
    Safely adjust paragraph spacing in document without affecting cell colors.
    """
    try:
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
        logger.debug("Successfully fixed paragraph spacing")
    except Exception as e:
        logger.error(f"Error in safe_fix_paragraph_spacing: {str(e)}")
        raise

def apply_conditional_formatting(doc, conditions=None):
    """
    Apply conditional formatting to document elements based on specified conditions.
    """
    try:
        if not conditions:
            conditions = {
                'PRICE_START': {
                    'bold': True,
                    'color': RGBColor(0, 0, 0),
                    'size': Pt(12)
                },
                'LINEAGE_START': {
                    'bold': True,
                    'color': RGBColor(255, 255, 255),
                    'size': Pt(11)
                }
            }
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for marker, formatting in conditions.items():
                                if marker in run.text:
                                    if formatting.get('bold') is not None:
                                        run.font.bold = formatting['bold']
                                    if formatting.get('color') is not None:
                                        run.font.color.rgb = formatting['color']
                                    if formatting.get('size') is not None:
                                        run.font.size = formatting['size']
        logger.debug("Applied conditional formatting to document")
        return doc
    except Exception as e:
        logger.error(f"Error applying conditional formatting: {str(e)}")
        raise

def set_cell_background(cell, color_hex):
    """Set cell background color with white text."""
    try:
        color_hex = color_hex.upper().strip('#')
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for element in tcPr.findall(qn('w:shd')):
            tcPr.remove(element)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:themeFill'), '0')
        tcPr.append(shd)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = "Arial"
        return
    except Exception as e:
        logger.error(f"Error in set_cell_background: {str(e)}")
        raise

def clear_cell_background(cell):
    """Clear cell background color and reset to default."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove any existing shading
        for element in tcPr.findall(qn('w:shd')):
            tcPr.remove(element)
        # Set to clear/transparent background
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'auto')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        tcPr.append(shd)
        # Reset text color to black
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        return
    except Exception as e:
        logger.error(f"Error in clear_cell_background: {str(e)}")
        raise

def clear_cell_margins(cell):
    """Remove cell margins."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'right', 'bottom', 'left']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        for element in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(element)
        tcPr.append(tcMar)
    except Exception as e:
        logger.error(f"Error in clear_cell_margins: {str(e)}")
        raise

def clear_table_cell_padding(cell):
    """Clear padding from a table cell."""
    try:
        if not cell or not hasattr(cell, '_tc'):
            return
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        for old_mar in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(old_mar)
        tcPr.append(tcMar)
    except Exception as e:
        logger.error(f"Error in clear_table_cell_padding: {str(e)}")
        raise

def enforce_ratio_formatting(doc):
    """Enforce Arial Bold and consistent font size for ratio-related content."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    ratio_patterns = [
        'THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:',
        '100mg', '500mg', '50mg', '25mg', '10mg', '5mg',
        '1:1:1', '1:1', '2:1', '3:1', '4:1', '5:1',
        'CBC/CBG', 'THC/CBD', 'CBD/THC',
        'RATIO_START', 'THC_CBD_START',
        'mg THC', 'mg CBD', 'mg CBC', 'mg CBG', 'mg CBN',
        # Add patterns for longer ratio values
        'RATIO_END', 'THC_CBD_END',
        # Add patterns for any content containing mg values
        'mg', 'MG',
        # Add patterns for ratio values with spaces
        'THC mg', 'CBD mg', 'CBC mg', 'CBG mg', 'CBN mg',
        # Add patterns for any content that looks like cannabinoid ratios
        'THC/CBD', 'CBD/THC', 'THC/CBC', 'CBC/THC', 'THC/CBG', 'CBG/THC'
    ]
    
    def is_ratio_content(text):
        """Check if text contains ratio-like content."""
        # Check for specific patterns
        if any(pattern in text for pattern in ratio_patterns):
            return True
        
        # Check for RATIO markers
        if 'RATIO_START' in text or 'RATIO_END' in text:
            return True
            
        # Check for THC_CBD markers
        if 'THC_CBD_START' in text or 'THC_CBD_END' in text:
            return True
            
        # Check for cannabinoid patterns with numbers and mg
        cannabinoid_pattern = r'\b(THC|CBD|CBC|CBG|CBN)\s*\d+mg\b'
        if re.search(cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for ratio patterns like "X:Y"
        ratio_pattern = r'\b\d+:\d+\b'
        if re.search(ratio_pattern, text):
            return True
            
        # Check for partial ratio content (individual cannabinoid values)
        # This catches cases where "50mg CBC" might be in a separate run
        partial_cannabinoid_pattern = r'\b\d+mg\s+(THC|CBD|CBC|CBG|CBN)\b'
        if re.search(partial_cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for any text containing "mg" followed by a cannabinoid
        mg_cannabinoid_pattern = r'mg\s+(THC|CBD|CBC|CBG|CBN)'
        if re.search(mg_cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for any text containing a cannabinoid followed by "mg"
        cannabinoid_mg_pattern = r'(THC|CBD|CBC|CBG|CBN)\s+mg'
        if re.search(cannabinoid_mg_pattern, text, re.IGNORECASE):
            return True
            
        return False

    def process_paragraph(paragraph):
        # Use comprehensive ratio content detection
        if is_ratio_content(paragraph.text):
            # Store text content and existing font sizes
            text = paragraph.text
            existing_runs = []
            for run in paragraph.runs:
                existing_runs.append({
                    'text': run.text,
                    'size': run.font.size,
                    'bold': run.font.bold
                })
            
            # Store existing paragraph alignment
            existing_alignment = paragraph.alignment
            
            # Clear paragraph
            paragraph.clear()
            
            # Add new run with proper formatting
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.bold = True
            
            # Restore paragraph alignment if it was set to right (for THC/CBD content)
            if existing_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Set font properties at XML level
            rPr = run._element.get_or_add_rPr()
            
            # Force Arial font
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:eastAsia'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.append(rFonts)
            
            # Force bold
            b = OxmlElement('w:b')
            b.set(qn('w:val'), '1')
            rPr.append(b)
            
            # Preserve existing font size if available, otherwise use default
            if existing_runs and existing_runs[0]['size']:
                font_size = existing_runs[0]['size']
                run.font.size = font_size
                # Set at XML level too
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # Word uses half-points
                rPr.append(sz)
            else:
                # Use default size (12pt)
                run.font.size = Pt(12)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(12 * 2)))  # Word uses half-points
                rPr.append(sz)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                    # Also process individual runs for longer ratio values
                    # First check if the paragraph contains any ratio content
                    paragraph_has_ratio = is_ratio_content(paragraph.text)
                    
                    for run in paragraph.runs:
                        # If paragraph has ratio content, make all runs bold
                        # OR if individual run has ratio content, make it bold
                        if paragraph_has_ratio or is_ratio_content(run.text):
                            run.font.name = "Arial"
                            run.font.bold = True
                            # Set at XML level for maximum compatibility
                            rPr = run._element.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Arial')
                            rFonts.set(qn('w:hAnsi'), 'Arial')
                            rFonts.set(qn('w:eastAsia'), 'Arial')
                            rFonts.set(qn('w:cs'), 'Arial')
                            rPr.append(rFonts)
                            b = OxmlElement('w:b')
                            b.set(qn('w:val'), '1')
                            rPr.append(b)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
        # Also process individual runs for longer ratio values
        # First check if the paragraph contains any ratio content
        paragraph_has_ratio = is_ratio_content(paragraph.text)
        
        for run in paragraph.runs:
            # If paragraph has ratio content, make all runs bold
            # OR if individual run has ratio content, make it bold
            if paragraph_has_ratio or is_ratio_content(run.text):
                run.font.name = "Arial"
                run.font.bold = True
                # Set at XML level for maximum compatibility
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                rFonts.set(qn('w:cs'), 'Arial')
                rPr.append(rFonts)
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                rPr.append(b)

    return doc

def enforce_arial_bold_all_text(doc):
    """Enforce Arial Bold font for ALL text in the document - NO EXCEPTIONS."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def process_run(run):
        """Apply Arial Bold formatting to a single run - NO EXCEPTIONS."""
        # Store existing font size only (we don't care about existing bold state)
        existing_size = run.font.size
        
        # FORCE Arial Bold for EVERYTHING - NO EXCEPTIONS
        run.font.name = "Arial"
        run.font.bold = True
        
        # Remove any italic formatting
        run.font.italic = False
        
        # Remove any other font properties that might interfere
        if hasattr(run.font, 'underline'):
            run.font.underline = None
        
        # Restore font size if it existed
        if existing_size:
            run.font.size = existing_size

        # Force Arial at XML level for maximum compatibility - NO EXCEPTIONS
        rPr = run._element.get_or_add_rPr()
        
        # Remove any existing font properties
        for element in list(rPr):
            if element.tag.endswith('}rFonts') or element.tag.endswith('}b') or element.tag.endswith('}i'):
                rPr.remove(element)
        
        # Set font family - FORCE Arial
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr.append(rFonts)
        
        # Force bold - NO EXCEPTIONS
        b = OxmlElement('w:b')
        b.set(qn('w:val'), '1')
        rPr.append(b)
        
        # Remove italic at XML level
        i = OxmlElement('w:i')
        i.set(qn('w:val'), '0')
        rPr.append(i)
        
        # Set font size at XML level if it exists
        if existing_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:w'), str(int(existing_size.pt * 2)))  # Word uses half-points
            rPr.append(sz)
            
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:w'), str(int(existing_size.pt * 2)))
            rPr.append(szCs)

    # Process ALL runs in ALL paragraphs in ALL tables - NO EXCEPTIONS
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Process ALL runs regardless of text content - NO EXCEPTIONS
                        process_run(run)
    
    # Process ALL runs in ALL paragraphs outside tables - NO EXCEPTIONS
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Process ALL runs regardless of text content - NO EXCEPTIONS
            process_run(run)
    
    # Process ALL runs in ALL headers and footers - NO EXCEPTIONS
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                # Process ALL runs regardless of text content - NO EXCEPTIONS
                process_run(run)
        for footer in section.footer.paragraphs:
            for run in footer.runs:
                # Process ALL runs regardless of text content - NO EXCEPTIONS
                process_run(run)
    
    return doc

def enforce_thc_cbd_bold_formatting(doc):
    """Enforce bold formatting for THC/CBD labels and values in the new format."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    def process_paragraph(paragraph):
        text = paragraph.text.strip()
        
        # Check if this paragraph contains the new THC/CBD format
        if text.startswith('THC:') or text.startswith('CBD:'):
            # Clear the paragraph
            paragraph.clear()
            
            # Split the text into lines
            lines = text.split('\n')
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Add a new run for each line
                run = paragraph.add_run(line)
                run.font.name = "Arial"
                run.font.bold = True
                
                # Set font size
                run.font.size = Pt(10)
                
                # Add line break if not the last line
                if i < len(lines) - 1:
                    paragraph.add_run('\n')
    
    # Process all paragraphs in the document
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    # Process all paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

def cleanup_all_price_markers(doc):
    """Remove all price markers while preserving formatting."""
    price_patterns = [
        'PRICE_START',
        'PRICE_END',
        '{{PRICE}}',
        '{{/PRICE}}'
    ]
    
    def process_paragraph(paragraph):
        # Store original text and check if it contains price markers
        text = paragraph.text
        if not any(pattern in text for pattern in price_patterns):
            return
        
        # Remove all price markers
        for pattern in price_patterns:
            text = text.replace(pattern, '')
        
        # Clear paragraph and add cleaned text
        paragraph.clear()
        if text.strip():
            run = paragraph.add_run(text.strip())
            # Apply price formatting
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(14)  # Standard price font size

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def remove_extra_spacing(doc):
    """Remove extra spacing between paragraphs and set consistent line spacing while preserving font sizes."""
    from docx.shared import Pt
    
    def process_paragraph(paragraph):
        # Store text content and existing font sizes
        if not paragraph.text.strip():
            return
            
        text = paragraph.text
        
        # Store existing font sizes for each run
        existing_sizes = []
        for run in paragraph.runs:
            # Process ALL runs regardless of text content - NO EXCEPTIONS
            existing_sizes.append(run.font.size)
        
        # Clear and reset paragraph
        paragraph.clear()
        
        # Split text into runs and preserve font sizes
        words = text.split()
        size_index = 0
        
        for i, word in enumerate(words):
            run = paragraph.add_run(word)
            
            # Set consistent font
            run.font.name = "Arial"
            run.font.bold = True
            
            # Restore font size if available
            if size_index < len(existing_sizes) and existing_sizes[size_index]:
                run.font.size = existing_sizes[size_index]
                size_index += 1
            
            # Add space between words (except for last word)
            if i < len(words) - 1:
                run.add_text(' ')
        
        # Set paragraph spacing to minimum to prevent cell expansion
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def apply_type_formatting(doc, product_type, template_type='vertical'):
    """Apply specific formatting based on product type."""
    def process_paragraph(paragraph):
        # Skip empty paragraphs
        if not paragraph.text.strip():
            return
            
        # Apply type-specific formatting
        if product_type.lower() == 'paraphernalia':
            # Make paraphernalia text smaller and less prominent
            for run in paragraph.runs:
                if run.font.size:
                    current_size = run.font.size.pt
                    run.font.size = Pt(max(6, current_size * 0.8))  # Reduce size by 20%
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        elif product_type.lower() in ['concentrate', 'wax', 'shatter', 'rosin']:
            # Make concentrate text bold and prominent
            for run in paragraph.runs:
                run.font.bold = True
                if run.font.size:
                    current_size = run.font.size.pt
                    run.font.size = Pt(min(24, current_size * 1.1))  # Increase size by 10%

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def create_3x3_grid(doc, template_type='vertical'):
    """Create a 3x3 grid table for label layout."""
    try:
        # Fix page margins first to ensure the grid fits
        doc = fix_page_margins_for_3x3_grid(doc)
        
        # Remove existing tables
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        
        # Remove default paragraph if it exists
        if doc.paragraphs:
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        
        # Create new table
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table properties
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        table._element.insert(0, tblPr)
        
        # Calculate optimal column width based on available space
        # Standard letter paper is 8.5" wide, with 0.25" margins = 8.0" available
        # We need 3 columns, so each column should be about 2.6" wide
        # Leave some space for borders and spacing
        available_width = 8.0  # 8.5" - 0.5" total margins
        col_width_inches = min(2.6, available_width / 3)  # Don't exceed available space
        
        # Set column widths
        col_width = Inches(col_width_inches)
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(3):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(col_width.inches * 1440)))  # Convert to twips
            tblGrid.append(gridCol)
        table._element.insert(0, tblGrid)
        
        # Calculate optimal row height based on available space
        # Standard letter paper is 11" tall, with 0.25" margins = 10.5" available
        # We need 3 rows, so each row should fit within available space
        available_height = 10.5  # 11" - 0.5" total margins
        # Leave minimal buffer space for borders and spacing
        usable_height = available_height - 0.1  # 0.1" buffer for borders/spacing
        row_height_inches = min(3.47, usable_height / 3)  # Ensure total height fits with minimal buffer
        
        # Set row heights
        row_height = Inches(row_height_inches)
        for row in table.rows:
            row.height = row_height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Enforce fixed cell dimensions to prevent any growth
        enforce_fixed_cell_dimensions(table, template_type)
        
        logger.debug(f"Created 3x3 grid table with {col_width_inches:.2f}\" columns and {row_height_inches:.2f}\" rows")
        return table
    except Exception as e:
        logger.error(f"Error creating 3x3 grid: {str(e)}")
        raise

def disable_autofit(table):
    """Disable autofit for a table."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        table._element.insert(0, tblPr)
        logger.debug("Disabled autofit for table")
    except Exception as e:
        logger.error(f"Error disabling autofit: {str(e)}")
        raise

def truncate_text_for_cell(text, max_length=50):
    """Truncate text to fit within cell boundaries."""
    if not text or len(text) <= max_length:
        return text
    
    # Truncate and add ellipsis
    return text[:max_length-3] + "..."


def enforce_fixed_cell_dimensions(table, template_type=None):
    """Enforce fixed cell dimensions to prevent any cell growth with text."""
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Safety check: ensure table has valid structure
        if not table or not table.rows or len(table.rows) == 0:
            logger.warning("Cannot enforce dimensions on empty or invalid table")
            return table
        
        # Additional safety check: ensure table has valid XML structure
        try:
            # Test if we can access the first row's cells safely
            if table.rows and len(table.rows) > 0:
                test_row = table.rows[0]
                if not hasattr(test_row, '_element') or not hasattr(test_row._element, 'tc_lst'):
                    logger.warning("Table row missing required XML structure, attempting to repair")
                    return table  # Return table as-is if we can't repair it
        except Exception as e:
            logger.warning(f"Table structure validation failed: {e}, returning table as-is")
            return table
        
        # Disable autofit
        disable_autofit(table)
        
        # Set table properties to prevent any auto-sizing
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        
        # Ensure fixed layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Set table to not auto-fit
        table.autofit = False
        if hasattr(table, 'allow_autofit'):
            table.allow_autofit = False
        
        # If template_type is provided, enforce the correct dimensions
        if template_type:
            # For mini templates, enforce the exact 1.5" x 1.5" dimensions
            if template_type == 'mini':
                logger.info("Enforcing exact 1.5\" x 1.5\" dimensions for mini template")
                try:
                    from src.core.constants import CELL_DIMENSIONS
                    cell_dims = CELL_DIMENSIONS.get(template_type)
                    if cell_dims and table.rows:
                        # Set column widths based on template type
                        tblGrid = table._element.find(qn('w:tblGrid'))
                        if tblGrid is not None:
                            # Remove existing grid
                            tblGrid.getparent().remove(tblGrid)
                        
                        # Create new grid with correct dimensions
                        tblGrid = OxmlElement('w:tblGrid')
                        # Use the actual number of columns in the table, not table.columns
                        num_cols = len(table.rows[0].cells) if table.rows else 0
                        if num_cols > 0:
                            for _ in range(num_cols):
                                gridCol = OxmlElement('w:gridCol')
                                gridCol.set(qn('w:w'), str(int(cell_dims['width'] * 1440)))
                                tblGrid.append(gridCol)
                            table._element.insert(0, tblGrid)
                            
                            # Set row heights and cell widths based on template type
                            for row in table.rows:
                                row.height = Inches(cell_dims['height'])
                                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                                
                                # Set individual cell widths
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcW = tcPr.find(qn('w:tcW'))
                                    if tcW is None:
                                        tcW = OxmlElement('w:tcW')
                                        tcPr.append(tcW)
                                    tcW.set(qn('w:w'), str(int(cell_dims['width'] * 1440)))
                                    tcW.set(qn('w:type'), 'dxa')
                except Exception as e:
                    logger.warning(f"Could not enforce mini template dimensions: {e}")
                    # Continue with general dimension enforcement
            
            try:
                from src.core.constants import CELL_DIMENSIONS
                cell_dims = CELL_DIMENSIONS.get(template_type)
                if cell_dims and table.rows:
                    # Set column widths based on template type
                    tblGrid = table._element.find(qn('w:tblGrid'))
                    if tblGrid is not None:
                        # Remove existing grid
                        tblGrid.getparent().remove(tblGrid)
                    
                    # Create new grid with correct dimensions
                    tblGrid = OxmlElement('w:tblGrid')
                    # Use the actual number of columns in the table, not table.columns
                    num_cols = len(table.rows[0].cells) if table.rows else 0
                    if num_cols > 0:
                        for _ in range(num_cols):
                            gridCol = OxmlElement('w:gridCol')
                            gridCol.set(qn('w:w'), str(int(cell_dims['width'] * 1440)))
                            tblGrid.append(gridCol)
                        table._element.insert(0, tblGrid)
                        
                        # Set row heights and cell widths based on template type
                        for row in table.rows:
                            row.height = Inches(cell_dims['height'])
                            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                            
                            # Set individual cell widths
                            for cell in row.cells:
                                tcPr = cell._tc.get_or_add_tcPr()
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is None:
                                    tcW = OxmlElement('w:tcW')
                                    tcPr.append(tcW)
                                tcW.set(qn('w:w'), str(int(cell_dims['width'] * 1440)))
                                tcW.set(qn('w:type'), 'dxa')
            except Exception as e:
                logger.warning(f"Could not enforce template-specific dimensions for {template_type}: {e}")
                # Continue with general dimension enforcement
        
        # Ensure table has a valid tblGrid before processing cells
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is None:
            # Create tblGrid if it doesn't exist
            tblGrid = OxmlElement('w:tblGrid')
            if table.rows and len(table.rows) > 0:
                num_cols = len(table.rows[0]._element.tc_lst)
                for _ in range(num_cols):
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), '1440')  # Default width
                    tblGrid.append(gridCol)
                table._element.insert(0, tblGrid)
        
        # Process each cell to enforce fixed dimensions
        for row in table.rows:
            try:
                # Set exact row height rule if not already set
                if not template_type:
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                # Use _element.tc_lst to safely access cells
                if hasattr(row, '_element') and hasattr(row._element, 'tc_lst'):
                    for cell_element in row._element.tc_lst:
                        try:
                            cell = row.cells[cell_element._index] if hasattr(cell_element, '_index') else None
                            if cell is None:
                                continue
                            
                            # Check if this cell contains a DOH image before setting TOP alignment
                            from src.core.utils.common import cell_contains_doh_image
                            has_doh_image = cell_contains_doh_image(cell)
                            
                            # Set cell vertical alignment to top to prevent content from expanding cell
                            # BUT preserve center alignment for cells with DOH images
                            if not has_doh_image:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                            else:
                                # AGGRESSIVELY preserve center alignment for DOH images
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                
                                # Apply XML-level vertical alignment to be absolutely sure
                                try:
                                    from docx.oxml import OxmlElement
                                    from docx.oxml.ns import qn
                                    tc_element = cell._tc
                                    tcPr = tc_element.get_or_add_tcPr()
                                    
                                    # Force vertical alignment at cell level
                                    vAlign = tcPr.find(qn('w:vAlign'))
                                    if vAlign is None:
                                        vAlign = OxmlElement('w:vAlign')
                                        tcPr.append(vAlign)
                                    vAlign.set(qn('w:val'), 'center')
                                    
                                    # Also ensure paragraph centering
                                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        
                                except Exception as e:
                                    logger.warning(f"Error enforcing DOH center alignment in docx_formatting: {e}")
                                
                                logger.debug("Preserved center alignment for cell with DOH image in enforce_fixed_cell_dimensions")
                            
                            # Clear any cell margins that might allow expansion
                            clear_cell_margins(cell)
                            
                            # CRITICAL: Force cell to maintain exact dimensions
                            tcPr = cell._tc.get_or_add_tcPr()
                            
                            # Set cell width to exact value to prevent expansion
                            tcW = tcPr.find(qn('w:tcW'))
                            if tcW is None:
                                tcW = OxmlElement('w:tcW')
                                tcPr.append(tcW)
                            tcW.set(qn('w:w'), '1440')  # Fixed width in twips
                            tcW.set(qn('w:type'), 'dxa')  # Fixed width type
                            
                            # Disable cell auto-sizing
                            tcFitText = tcPr.find(qn('w:tcFitText'))
                            if tcFitText is None:
                                tcFitText = OxmlElement('w:tcFitText')
                                tcPr.append(tcFitText)
                            tcFitText.set(qn('w:val'), '0')  # Disable fit text
                            
                            # Set cell height to exact value
                            tcH = tcPr.find(qn('w:tcH'))
                            if tcH is None:
                                tcH = OxmlElement('w:tcH')
                                tcPr.append(tcH)
                            tcH.set(qn('w:w'), '1440')  # Fixed height in twips
                            tcH.set(qn('w:hRule'), 'exact')  # Exact height rule
                            
                            # Process paragraphs in the cell to prevent text overflow
                            for paragraph in cell.paragraphs:
                                # Set paragraph spacing to minimum
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.paragraph_format.line_spacing = 1.0
                                
                                # CRITICAL: Set paragraph alignment to prevent expansion
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # CRITICAL: Enable text wrapping and prevent overflow
                                pPr = paragraph._element.get_or_add_pPr()
                                
                                # Add text wrapping control
                                wrap = pPr.find(qn('w:wordWrap'))
                                if wrap is None:
                                    wrap = OxmlElement('w:wordWrap')
                                    pPr.append(wrap)
                                wrap.set(qn('w:val'), '1')  # Enable word wrapping
                                
                                # Add overflow control
                                overflow = pPr.find(qn('w:overflowPunct'))
                                if overflow is None:
                                    overflow = OxmlElement('w:overflowPunct')
                                    pPr.append(overflow)
                                overflow.set(qn('w:val'), '0')  # Disable overflow punctuation
                                
                                # CRITICAL: Truncate text if it's too long for the cell
                                full_text = paragraph.text
                                if full_text and len(full_text) > 50:  # Adjust max length as needed
                                    truncated_text = truncate_text_for_cell(full_text, 50)
                                    # Clear existing runs and add truncated text
                                    paragraph.clear()
                                    run = paragraph.add_run(truncated_text)
                                    
                                    # Set font properties to prevent text expansion
                                    if not run.font.size:
                                        run.font.size = Pt(12)  # Set default size if none
                                    
                                    # CRITICAL: Force text to wrap within cell boundaries
                                    rPr = run._element.get_or_add_rPr()
                                    
                                    # Add text wrapping control to run
                                    wrap_run = rPr.find(qn('w:wordWrap'))
                                    if wrap_run is None:
                                        wrap_run = OxmlElement('w:wordWrap')
                                        rPr.append(wrap_run)
                                    wrap_run.set(qn('w:val'), '1')  # Enable word wrapping
                                else:
                                    # Ensure text doesn't wrap beyond cell boundaries
                                    for run in paragraph.runs:
                                        # Set font properties to prevent text expansion
                                        if not run.font.size:
                                            run.font.size = Pt(12)  # Set default size if none
                                        
                                        # CRITICAL: Force text to wrap within cell boundaries
                                        rPr = run._element.get_or_add_rPr()
                                        
                                        # Add text wrapping control to run
                                        wrap_run = rPr.find(qn('w:wordWrap'))
                                        if wrap_run is None:
                                            wrap_run = OxmlElement('w:wordWrap')
                                            rPr.append(wrap_run)
                                        wrap_run.set(qn('w:val'), '1')  # Enable word wrapping
                        except Exception as cell_error:
                            logger.warning(f"Error processing cell in row: {cell_error}")
                            continue
                else:
                    logger.warning(f"Row missing required XML structure: {row}")
            except Exception as row_error:
                logger.warning(f"Error processing row: {row_error}")
                continue
        
        # CRITICAL: Final autofit disabling to ensure no expansion
        table.autofit = False
        if hasattr(table, 'allow_autofit'):
            table.allow_autofit = False
        
        # CRITICAL: Verify table layout is fixed
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None or tblLayout.get(qn('w:type')) != 'fixed':
                # Force fixed layout
                if tblLayout is not None:
                    tblLayout.getparent().remove(tblLayout)
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            # CRITICAL: Add table width constraint to prevent expansion
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), '4320')  # Fixed table width in twips (3 inches)
            tblW.set(qn('w:type'), 'dxa')  # Fixed width type
            
            # CRITICAL: Disable table auto-sizing completely
            tblLook = tblPr.find(qn('w:tblLook'))
            if tblLook is None:
                tblLook = OxmlElement('w:tblLook')
                tblPr.append(tblLook)
            tblLook.set(qn('w:val'), '0000')  # Disable all auto-sizing
            tblLook.set(qn('w:firstRow'), '0')
            tblLook.set(qn('w:lastRow'), '0')
            tblLook.set(qn('w:firstColumn'), '0')
            tblLook.set(qn('w:lastColumn'), '0')
            tblLook.set(qn('w:noHBand'), '0')
            tblLook.set(qn('w:noVBand'), '0')
        
        logger.debug(f"Enforced fixed cell dimensions for table (template: {template_type})")
        return table
    except Exception as e:
        logger.error(f"Error enforcing fixed cell dimensions: {str(e)}")
        raise

def fix_table(doc, num_rows=3, num_cols=3, template_type='horizontal'):
    """Fix table with proper cell dimensions based on template type."""
    from src.core.constants import CELL_DIMENSIONS
    
    # Get individual cell dimensions
    cell_dims = CELL_DIMENSIONS.get(template_type, {'width': 2.4, 'height': 2.4})
    cell_width = Inches(cell_dims['width'])
    cell_height = Inches(cell_dims['height'])
    
    # Remove all existing tables
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    
    # Remove default paragraph if it exists
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    # Create new table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    table._element.insert(0, tblPr)
    
    # Set column widths
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(cell_width.inches * 1440)))
        tblGrid.append(gridCol)
    table._element.insert(0, tblGrid)
    
    # Set row heights
    for row in table.rows:
        row.height = cell_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Enforce fixed cell dimensions to prevent any growth
    enforce_fixed_cell_dimensions(table, template_type)
    
    return table

def rebuild_3x3_grid(doc, template_type='horizontal'):
    """Rebuild 3x3 grid with proper cell dimensions based on template type."""
    from src.core.constants import CELL_DIMENSIONS
    
    # Get individual cell dimensions
    cell_dims = CELL_DIMENSIONS.get(template_type, {'width': 2.4, 'height': 2.4})
    cell_width = Inches(cell_dims['width'])
    cell_height = Inches(cell_dims['height'])
    
    # Remove all existing tables
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    
    # Remove default paragraph if it exists
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    # Create new 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    table._element.insert(0, tblPr)
    
    # Set column widths
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(cell_width.inches * 1440)))
        tblGrid.append(gridCol)
    table._element.insert(0, tblGrid)
    
    # Set row heights
    for row in table.rows:
        row.height = cell_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Enforce fixed cell dimensions to prevent any growth
    enforce_fixed_cell_dimensions(table, template_type)
    
    return table 

def apply_custom_formatting(doc, template_settings):
    """Apply custom formatting based on template settings."""
    try:
        font_family = template_settings.get('font', 'Arial')
        bold_headers = template_settings.get('boldHeaders', False)
        italic_descriptions = template_settings.get('italicDescriptions', False)
        line_spacing = float(template_settings.get('lineSpacing', '1.0'))
        paragraph_spacing = int(template_settings.get('paragraphSpacing', '0'))
        text_color = template_settings.get('textColor', '#000000')
        background_color = template_settings.get('backgroundColor', '#ffffff')
        header_color = template_settings.get('headerColor', '#333333')
        accent_color = template_settings.get('accentColor', '#007bff')
        auto_resize = template_settings.get('autoResize', True)
        smart_truncation = template_settings.get('smartTruncation', True)
        
        # Apply formatting to all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Set line spacing
            paragraph.paragraph_format.line_spacing = line_spacing
            
            # Set paragraph spacing
            if paragraph_spacing > 0:
                paragraph.paragraph_format.space_after = Pt(paragraph_spacing)
            
            # Apply formatting to runs
            for run in paragraph.runs:
                # Set font family
                run.font.name = font_family
                
                # Set font color
                if text_color != '#000000':
                    run.font.color.rgb = RGBColor.from_string(text_color[1:])  # Remove # from hex
                
                # Apply bold to headers (if enabled)
                if bold_headers and any(keyword in run.text.lower() for keyword in ['brand', 'price', 'lineage', 'thc', 'cbd']):
                    run.font.bold = True
                    if header_color != '#333333':
                        run.font.color.rgb = RGBColor.from_string(header_color[1:])
                
                # Apply italic to descriptions (if enabled)
                if italic_descriptions and len(run.text) > 20:  # Assume long text is description
                    run.font.italic = True
        
        # Apply background color to tables if specified
        if background_color != '#ffffff':
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_background(cell, background_color)
        
        # Apply accent color to specific elements
        if accent_color != '#007bff':
            # Apply to price elements
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if any(keyword in run.text.lower() for keyword in ['$', 'price', 'cost']):
                        run.font.color.rgb = RGBColor.from_string(accent_color[1:])
        
        # Apply auto-resize if enabled
        if auto_resize:
            # Adjust font sizes to fit content
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.size and run.font.size.pt > 6:
                        # Reduce font size if text is too long
                        if len(run.text) > 50:
                            run.font.size = Pt(max(6, run.font.size.pt - 2))
        
        # Apply smart truncation if enabled
        if smart_truncation:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if len(run.text) > 100:
                        # Truncate very long text
                        run.text = run.text[:97] + "..."
        
        logger.debug(f"Applied custom formatting with font: {font_family}, line spacing: {line_spacing}")
        
    except Exception as e:
        logger.error(f"Error applying custom formatting: {str(e)}")
        # Fall back to default formatting
        enforce_arial_bold_all_text(doc) 

def remove_all_headers_and_footers(doc):
    """CRITICAL: Remove ALL headers and footers from the document to prevent unwanted content."""
    try:
        # Process all sections in the document
        for section in doc.sections:
            # Remove header content safely
            if hasattr(section, 'header') and section.header:
                try:
                    # Clear header content by removing all paragraphs
                    for paragraph in section.header.paragraphs:
                        paragraph.clear()
                    # Clear any tables in header
                    for table in section.header.tables:
                        table._element.getparent().remove(table._element)
                except Exception as e:
                    logger.warning(f"Could not clear header content: {e}")
            
            # Remove footer content safely
            if hasattr(section, 'footer') and section.footer:
                try:
                    # Clear footer content by removing all paragraphs
                    for paragraph in section.footer.paragraphs:
                        paragraph.clear()
                    # Clear any tables in footer
                    for table in section.footer.tables:
                        table._element.getparent().remove(table._element)
                except Exception as e:
                    logger.warning(f"Could not clear footer content: {e}")
            
            # Ensure no header/footer spacing
            try:
                section.header_distance = 0
                section.footer_distance = 0
            except Exception as e:
                logger.warning(f"Could not set header/footer distances: {e}")
        
        logger.info("Successfully removed headers and footers from document")
        return doc
        
        # Also check for any header/footer references in the document XML
        try:
            from docx.oxml.ns import qn
            # Remove any w:headerReference and w:footerReference elements
            for section in doc.sections:
                sectPr = section._element.find(qn('w:sectPr'))
                if sectPr is not None:
                    # Remove header references
                    header_refs = sectPr.findall(qn('w:headerReference'))
                    for header_ref in header_refs:
                        sectPr.remove(header_ref)
                    
                    # Remove footer references
                    footer_refs = sectPr.findall(qn('w:footerReference'))
                    for footer_ref in footer_refs:
                        sectPr.remove(footer_ref)
        except Exception as e:
            logger.warning(f"Could not remove header/footer references: {e}")
        
        logger.info("CRITICAL: All headers and footers removed from document")
        return doc
    except Exception as e:
        logger.error(f"Error removing headers and footers: {e}")
        return doc

def fix_page_margins_for_3x3_grid(doc):
    """Fix page margins to ensure 3x3 grid fits properly."""
    try:
        from docx.shared import Inches
        
        # Get the first section (or create one if none exists)
        if not doc.sections:
            section = doc.add_section()
        else:
            section = doc.sections[0]
        
        # Set minimal margins to maximize available space for the 3x3 grid
        # Standard letter paper is 8.5" x 11"
        # We need space for: 3 columns × 3.4" = 10.2" total width
        # So we need very small margins
        section.left_margin = Inches(0.25)   # 0.25" left margin
        section.right_margin = Inches(0.25)  # 0.25" right margin
        section.top_margin = Inches(0.25)    # 0.25" top margin
        section.bottom_margin = Inches(0.25) # 0.25" bottom margin
        
        # REMOVE ANY HEADERS AND FOOTERS that might be taking up space
        # This ensures the full page area is available for the 3x3 grid
        if hasattr(section, 'header') and section.header:
            # Clear header content by removing all child elements
            for child in list(section.header._element):
                section.header._element.remove(child)
        if hasattr(section, 'footer') and section.footer:
            # Clear footer content by removing all child elements
            for child in list(section.footer._element):
                section.footer._element.remove(child)
        # Ensure no header/footer spacing
        section.header_distance = 0
        section.footer_distance = 0
        
        # Calculate available space (convert twips to inches for logging)
        # 1 inch = 1440 twips
        available_width_inches = (section.page_width - section.left_margin - section.right_margin) / 1440
        available_height_inches = (section.page_height - section.top_margin - section.bottom_margin) / 1440
        
        logger.debug(f"Page dimensions: {section.page_width/1440:.2f}\" x {section.page_height/1440:.2f}\"")
        logger.debug(f"Margins: L={section.left_margin.inches:.2f}\", R={section.right_margin.inches:.2f}\", T={section.top_margin.inches:.2f}\", B={section.bottom_margin.inches:.2f}\"")
        logger.debug(f"Available space: {available_width_inches:.2f}\" x {available_height_inches:.2f}\"")
        
        return doc
    except Exception as e:
        logger.error(f"Error fixing page margins: {e}")
        return doc

def fix_page_margins_for_horizontal_3x3_grid(doc):
    """Fix page margins and enforce landscape orientation for horizontal 3x3 grid."""
    try:
        from docx.shared import Inches
        from docx.enum.section import WD_ORIENT
        
        # Get the first section (or create one if none exists)
        if not doc.sections:
            section = doc.add_section()
        else:
            section = doc.sections[0]
        
        # ENFORCE LANDSCAPE ORIENTATION for horizontal templates
        section.orientation = WD_ORIENT.LANDSCAPE
        
        # Set optimized margins for landscape 3x3 grid with exact cell dimensions
        # Landscape: 11" x 8.5" (rotated from 8.5" x 11")
        # Required: 3 columns × 3.4" = 10.2" width, 3 rows × 2.4" = 7.2" height
        # Available: 11" - 0.4" margins = 10.6" width, 8.5" - 0.65" margins = 7.85" height
        section.left_margin = Inches(0.2)    # 0.2" left margin
        section.right_margin = Inches(0.2)   # 0.2" right margin
        section.top_margin = Inches(0.325)   # 0.325" top margin
        section.bottom_margin = Inches(0.325) # 0.325" bottom margin
        
        # REMOVE ANY HEADERS AND FOOTERS that might be taking up space
        # This ensures the full page area is available for the 3x3 grid
        if hasattr(section, 'header') and section.header:
            # Clear header content by removing all child elements
            for child in list(section.header._element):
                section.header._element.remove(child)
        if hasattr(section, 'footer') and section.footer:
            # Clear footer content by removing all child elements
            for child in list(section.footer._element):
                section.footer._element.remove(child)
        # Ensure no header/footer spacing
        section.header_distance = 0
        section.footer_distance = 0
        
        # Calculate available space in landscape mode
        available_width_inches = (section.page_width - section.left_margin - section.right_margin) / 1440
        available_height_inches = (section.page_height - section.top_margin - section.bottom_margin) / 1440
        
        logger.debug(f"LANDSCAPE Page dimensions: {section.page_width/1440:.2f}\" x {section.page_height/1440:.2f}\"")
        logger.debug(f"LANDSCAPE Margins: L={section.left_margin.inches:.2f}\", R={section.right_margin.inches:.2f}\", T={section.top_margin.inches:.2f}\", B={section.bottom_margin.inches:.2f}\"")
        logger.debug(f"LANDSCAPE Available space: {available_width_inches:.2f}\" x {available_height_inches:.2f}\"")
        
        return doc
    except Exception as e:
        logger.error(f"Error fixing horizontal page margins: {e}")
        return doc 

def fix_page_margins_for_4x3_grid(doc):
    """Fix page margins to ensure 4x3 grid fits properly for double templates."""
    try:
        from docx.shared import Inches
        
        # Get the first section (or create one if none exists)
        if not doc.sections:
            section = doc.add_section()
        else:
            section = doc.sections[0]
        
        # Set minimal margins to maximize available space for the 4x3 grid
        # Standard letter paper is 8.5" x 11"
        # We need space for: 4 columns × 2.4" = 9.6" total width
        # So we need very small margins
        section.left_margin = Inches(0.25)   # 0.25" left margin
        section.right_margin = Inches(0.25)  # 0.25" right margin
        section.top_margin = Inches(0.25)    # 0.25" top margin
        section.bottom_margin = Inches(0.25) # 0.25" bottom margin
        
        # REMOVE ANY HEADERS AND FOOTERS that might be taking up space
        # This ensures the full page area is available for the 4x3 grid
        if hasattr(section, 'header') and section.header:
            # Clear header content by removing all child elements
            for child in list(section.header._element):
                section.header._element.remove(child)
        if hasattr(section, 'footer') and section.footer:
            # Clear footer content by removing all child elements
            for child in list(section.footer._element):
                section.footer._element.remove(child)
        # Ensure no header/footer spacing
        section.header_distance = 0
        section.footer_distance = 0
        
        # Calculate available space (convert twips to inches for logging)
        # 1 inch = 1440 twips
        available_width_inches = (section.page_width - section.left_margin - section.right_margin) / 1440
        available_height_inches = (section.page_height - section.top_margin - section.bottom_margin) / 1440
        
        logger.debug(f"4x3 GRID Page dimensions: {section.page_width/1440:.2f}\" x {section.page_height/1440:.2f}\"")
        logger.debug(f"4x3 GRID Margins: L={section.left_margin.inches:.2f}\", R={section.right_margin.inches:.2f}\", T={section.top_margin.inches:.2f}\", B={section.bottom_margin.inches:.2f}\"")
        logger.debug(f"4x3 GRID Available space: {available_width_inches:.2f}\" x {available_height_inches:.2f}\"")
        
        return doc
    except Exception as e:
        logger.error(f"Error fixing 4x3 grid page margins: {e}")
        return doc

def prevent_table_expansion_enhanced(doc, template_type=None):
    """
    Enhanced table expansion prevention with multiple layers of protection.
    This function applies the most aggressive table expansion prevention measures.
    """
    try:
        from docx.shared import Inches, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from src.core.constants import CELL_DIMENSIONS
        
        logger.info(f"Applying enhanced table expansion prevention for template: {template_type}")
        
        for table in doc.tables:
            try:
                # LAYER 1: Disable all auto-sizing features
                table.autofit = False
                if hasattr(table, 'allow_autofit'):
                    table.allow_autofit = False
                
                # LAYER 2: Force fixed table layout at XML level
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)
                
                # Remove any existing layout
                for existing_layout in tblPr.findall(qn('w:tblLayout')):
                    tblPr.remove(existing_layout)
                
                # Force fixed layout
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                
                # LAYER 3: Set absolute table width constraints
                if template_type and template_type in CELL_DIMENSIONS:
                    cell_dims = CELL_DIMENSIONS[template_type]
                    
                    # Calculate total table width based on grid layout
                    if template_type in ['horizontal', 'vertical']:
                        # 3x3 grid
                        total_width = cell_dims['width'] * 3
                        total_height = cell_dims['height'] * 3
                    elif template_type == 'mini':
                        # 4x5 grid (20 labels)
                        total_width = cell_dims['width'] * 4
                        total_height = cell_dims['height'] * 5
                    elif template_type == 'double':
                        # 4x3 grid
                        total_width = cell_dims['width'] * 4
                        total_height = cell_dims['height'] * 3
                    elif template_type == 'inventory':
                        # 2x2 grid
                        total_width = cell_dims['width'] * 2
                        total_height = cell_dims['height'] * 2
                    else:
                        # Default 3x3
                        total_width = cell_dims['width'] * 3
                        total_height = cell_dims['height'] * 3
                    
                    # Set absolute table width
                    tblW = tblPr.find(qn('w:tblW'))
                    if tblW is None:
                        tblW = OxmlElement('w:tblW')
                        tblPr.append(tblW)
                    tblW.set(qn('w:w'), str(int(total_width * 1440)))  # Convert to twips
                    tblW.set(qn('w:type'), 'dxa')  # Absolute width
                    
                    # LAYER 4: Create precise column grid
                    tblGrid = table._element.find(qn('w:tblGrid'))
                    if tblGrid is not None:
                        tblGrid.getparent().remove(tblGrid)
                    
                    tblGrid = OxmlElement('w:tblGrid')
                    num_cols = len(table.rows[0].cells) if table.rows else 3
                    
                    for _ in range(num_cols):
                        gridCol = OxmlElement('w:gridCol')
                        gridCol.set(qn('w:w'), str(int(cell_dims['width'] * 1440)))
                        tblGrid.append(gridCol)
                    table._element.insert(0, tblGrid)
                
                # LAYER 5: Process each row with absolute height constraints
                for row in table.rows:
                    # Set exact row height
                    if template_type and template_type in CELL_DIMENSIONS:
                        row.height = Inches(CELL_DIMENSIONS[template_type]['height'])
                    else:
                        row.height = Inches(2.4)  # Default
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    
                    # LAYER 6: Process each cell with absolute constraints
                    for cell in row.cells:
                        # Check if this cell contains a DOH image before setting TOP alignment
                        from src.core.utils.common import cell_contains_doh_image
                        has_doh_image = cell_contains_doh_image(cell)
                        
                        # Set vertical alignment to prevent content from expanding cell
                        # BUT preserve center alignment for cells with DOH images
                        if not has_doh_image:
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        else:
                            # AGGRESSIVELY preserve center alignment for DOH images
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Apply XML-level vertical alignment to be absolutely sure
                            try:
                                from docx.oxml import OxmlElement
                                from docx.oxml.ns import qn
                                tc_element = cell._tc
                                tcPr = tc_element.get_or_add_tcPr()
                                
                                # Force vertical alignment at cell level
                                vAlign = tcPr.find(qn('w:vAlign'))
                                if vAlign is None:
                                    vAlign = OxmlElement('w:vAlign')
                                    tcPr.append(vAlign)
                                vAlign.set(qn('w:val'), 'center')
                                
                                # Also ensure paragraph centering
                                from docx.enum.text import WD_ALIGN_PARAGRAPH
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                            except Exception as e:
                                logger.warning(f"Error enforcing DOH center alignment in prevent_table_expansion_enhanced: {e}")
                            logger.debug("Preserved center alignment for cell with DOH image in prevent_table_expansion_enhanced")
                        
                        # Clear all margins and padding
                        clear_cell_margins(cell)
                        clear_table_cell_padding(cell)
                        
                        # Force absolute cell dimensions at XML level
                        tcPr = cell._tc.get_or_add_tcPr()
                        
                        # Set absolute cell width
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is None:
                            tcW = OxmlElement('w:tcW')
                            tcPr.append(tcW)
                        if template_type and template_type in CELL_DIMENSIONS:
                            tcW.set(qn('w:w'), str(int(CELL_DIMENSIONS[template_type]['width'] * 1440)))
                        else:
                            tcW.set(qn('w:w'), '3456')  # Default 2.4 inches in twips
                        tcW.set(qn('w:type'), 'dxa')
                        
                        # Set absolute cell height
                        tcH = tcPr.find(qn('w:tcH'))
                        if tcH is None:
                            tcH = OxmlElement('w:tcH')
                            tcPr.append(tcH)
                        if template_type and template_type in CELL_DIMENSIONS:
                            tcH.set(qn('w:w'), str(int(CELL_DIMENSIONS[template_type]['height'] * 1440)))
                        else:
                            tcH.set(qn('w:w'), '3456')  # Default 2.4 inches in twips
                        tcH.set(qn('w:hRule'), 'exact')
                        
                        # Disable cell auto-sizing
                        tcFitText = tcPr.find(qn('w:tcFitText'))
                        if tcFitText is None:
                            tcFitText = OxmlElement('w:tcFitText')
                            tcPr.append(tcFitText)
                        tcFitText.set(qn('w:val'), '0')  # Disable fit text
                        
                        # LAYER 7: Process paragraphs to prevent text overflow
                        for paragraph in cell.paragraphs:
                            # Set minimal spacing
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Force left alignment to prevent expansion
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Add text wrapping controls at XML level
                            pPr = paragraph._element.get_or_add_pPr()
                            
                            # Enable word wrapping
                            wrap = pPr.find(qn('w:wordWrap'))
                            if wrap is None:
                                wrap = OxmlElement('w:wordWrap')
                                pPr.append(wrap)
                            wrap.set(qn('w:val'), '1')
                            
                            # Disable overflow
                            overflow = pPr.find(qn('w:overflowPunct'))
                            if overflow is None:
                                overflow = OxmlElement('w:overflowPunct')
                                pPr.append(overflow)
                            overflow.set(qn('w:val'), '0')
                            
                            # LAYER 8: Process runs to prevent font expansion
                            for run in paragraph.runs:
                                # Ensure text doesn't cause expansion
                                if len(run.text) > 100:
                                    # Truncate very long text
                                    run.text = run.text[:97] + "..."
                                
                                # Set font properties to prevent expansion
                                if not run.font.size:
                                    run.font.size = Pt(12)
                                
                                # Add text wrapping control to run level
                                rPr = run._element.get_or_add_rPr()
                                
                                wrap_run = rPr.find(qn('w:wordWrap'))
                                if wrap_run is None:
                                    wrap_run = OxmlElement('w:wordWrap')
                                    rPr.append(wrap_run)
                                wrap_run.set(qn('w:val'), '1')
                
                # LAYER 9: Final table-level constraints
                # Disable all table auto-sizing features
                tblLook = tblPr.find(qn('w:tblLook'))
                if tblLook is None:
                    tblLook = OxmlElement('w:tblLook')
                    tblPr.append(tblLook)
                tblLook.set(qn('w:val'), '0000')  # Disable all auto-sizing
                tblLook.set(qn('w:firstRow'), '0')
                tblLook.set(qn('w:lastRow'), '0')
                tblLook.set(qn('w:firstColumn'), '0')
                tblLook.set(qn('w:lastColumn'), '0')
                tblLook.set(qn('w:noHBand'), '0')
                tblLook.set(qn('w:noVBand'), '0')
                
                # Applied enhanced table expansion prevention
                
            except Exception as table_error:
                logger.warning(f"Error applying enhanced expansion prevention to table: {table_error}")
                continue
        
        # Enhanced table expansion prevention completed
        return doc
        
    except Exception as e:
        logger.error(f"Error in enhanced table expansion prevention: {e}")
        return doc 