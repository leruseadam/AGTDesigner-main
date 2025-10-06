#!/usr/bin/env python3
"""
Fast DOCX Generator
Optimized for web performance with minimal processing overhead
"""

import os
import time
import logging
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from typing import List, Dict, Any

# Performance constants
MAX_RECORDS_PER_DOCX = 100  # Limit records for web performance
CHUNK_SIZE = 20  # Process in small chunks
MAX_PROCESSING_TIME = 30  # 30 second timeout

logger = logging.getLogger(__name__)

class FastDocxGenerator:
    """Ultra-fast DOCX generation with minimal overhead"""
    
    def __init__(self):
        self.start_time = time.time()
        self.processed_count = 0
        
    def generate_labels_fast(self, records: List[Dict], template_type: str = 'vertical', 
                           scale_factor: float = 1.0) -> BytesIO:
        """Generate labels with optimized processing"""
        try:
            logger.info(f"[FAST-DOCX] Starting generation for {len(records)} records")
            
            # Limit records for web performance
            if len(records) > MAX_RECORDS_PER_DOCX:
                logger.warning(f"[FAST-DOCX] Limiting records from {len(records)} to {MAX_RECORDS_PER_DOCX}")
                records = records[:MAX_RECORDS_PER_DOCX]
            
            # Create document
            doc = Document()
            
            # Set up page margins for labels
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
            # Process records in chunks
            for i in range(0, len(records), CHUNK_SIZE):
                # Check timeout
                if time.time() - self.start_time > MAX_PROCESSING_TIME:
                    logger.warning(f"[FAST-DOCX] Timeout reached, stopping at {self.processed_count} records")
                    break
                
                chunk = records[i:i + CHUNK_SIZE]
                self._process_chunk(doc, chunk, template_type, scale_factor)
                self.processed_count += len(chunk)
            
            # Save to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            generation_time = time.time() - self.start_time
            logger.info(f"[FAST-DOCX] Generated {self.processed_count} labels in {generation_time:.2f}s")
            
            return output_buffer
            
        except Exception as e:
            logger.error(f"[FAST-DOCX] Error generating labels: {e}")
            raise
    
    def _process_chunk(self, doc: Document, chunk: List[Dict], template_type: str, scale_factor: float):
        """Process a chunk of records with minimal formatting"""
        try:
            for record in chunk:
                # Check timeout
                if time.time() - self.start_time > MAX_PROCESSING_TIME:
                    break
                
                # Create label table
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Set table properties for label size
                table.autofit = False
                table.allow_autofit = False
                
                # Get cell
                cell = table.cell(0, 0)
                cell.width = Inches(2.0 * scale_factor)
                
                # Add content with minimal formatting
                self._add_label_content(cell, record, template_type)
                
                # Add spacing between labels
                doc.add_paragraph()
                
        except Exception as e:
            logger.error(f"[FAST-DOCX] Error processing chunk: {e}")
            raise
    
    def _add_label_content(self, cell, record: Dict, template_type: str):
        """Add content to label cell with minimal processing"""
        try:
            # Clear cell content
            cell.text = ''
            
            # Get basic fields
            product_name = record.get('ProductName', record.get('Product Name*', 'N/A'))
            product_type = record.get('ProductType', record.get('Product Type*', 'N/A'))
            brand = record.get('ProductBrand', record.get('Product Brand', 'N/A'))
            
            # Improved weight mapping - check multiple weight fields
            weight = record.get('CombinedWeight')
            if not weight or weight == 'N/A' or weight == '':
                # Try WeightWithUnits field (from database)
                weight = record.get('WeightWithUnits')
                
            if not weight or weight == 'N/A' or weight == '':
                # Create combined weight from Weight* and Units fields
                weight_value = record.get('Weight', record.get('Weight*', ''))
                units = record.get('Units', '')
                
                if weight_value and units and str(units) != 'None' and str(units) != '':
                    # Format weight properly (remove .0 if it's a whole number)
                    try:
                        if float(weight_value) == int(float(weight_value)):
                            weight = f"{int(float(weight_value))}{units}"
                        else:
                            weight = f"{weight_value}{units}"
                    except (ValueError, TypeError):
                        weight = f"{weight_value}{units}"
                elif weight_value:
                    weight = str(weight_value)
                else:
                    weight = 'N/A'
            
            # Improved price mapping - use correct Excel field name
            price = record.get('Price* (Tier Name for Bulk)', record.get('Price', 'N/A'))
            lineage = record.get('Lineage', 'N/A')
            
            # Add DOH field processing - check correct column name
            doh = record.get('DOH Compliant (Yes/No)', record.get('DOH', record.get('DOH Compliant*', '')))
            
            # Clean lineage
            if 'LINEAGE_START' in lineage and 'LINEAGE_END' in lineage:
                start_idx = lineage.find('LINEAGE_START') + len('LINEAGE_START')
                end_idx = lineage.find('LINEAGE_END')
                if start_idx != -1 and end_idx != -1:
                    lineage = lineage[start_idx:end_idx].strip()
            
            # Add content based on template type
            if template_type == 'vertical':
                self._add_vertical_content(cell, product_name, product_type, brand, weight, price, lineage, doh)
            elif template_type == 'horizontal':
                self._add_horizontal_content(cell, product_name, product_type, brand, weight, price, lineage, doh)
            elif template_type == 'mini':
                self._add_mini_content(cell, product_name, product_type, brand, weight, price, lineage, doh)
            else:
                self._add_vertical_content(cell, product_name, product_type, brand, weight, price, lineage, doh)
                
        except Exception as e:
            logger.error(f"[FAST-DOCX] Error adding label content: {e}")
            # Add fallback content
            cell.text = f"Product: {product_name}\nType: {product_type}\nBrand: {brand}"
    
    def _add_vertical_content(self, cell, product_name, product_type, brand, weight, price, lineage, doh):
        """Add vertical label content"""
        # Product Name (bold)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(product_name)
        run1.bold = True
        run1.font.size = Pt(10)
        
        # Product Type
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(product_type)
        run2.font.size = Pt(8)
        
        # Brand
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(brand)
        run3.font.size = Pt(8)
        
        # Weight and Price
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(f"{weight} - ${price}")
        run4.font.size = Pt(8)
        
        # Lineage
        p5 = cell.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = p5.add_run(lineage)
        run5.font.size = Pt(8)
        
        # DOH (if applicable)
        if doh and str(doh).strip().upper() == 'YES':
            p6 = cell.add_paragraph()
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = p6.add_run('DOH')
            run6.bold = True
            run6.font.size = Pt(8)
    
    def _add_horizontal_content(self, cell, product_name, product_type, brand, weight, price, lineage, doh):
        """Add horizontal label content"""
        # Single line format with DOH
        doh_text = ' | DOH' if doh and str(doh).strip().upper() == 'YES' else ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(f"{product_name} | {product_type} | {brand} | {weight} | ${price} | {lineage}{doh_text}")
        run1.font.size = Pt(8)
    
    def _add_mini_content(self, cell, product_name, product_type, brand, weight, price, lineage, doh):
        """Add mini label content"""
        # Compact format with DOH
        doh_text = ' | DOH' if doh and str(doh).strip().upper() == 'YES' else ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(f"{product_name}\n{product_type} | {weight}{doh_text}")
        run1.font.size = Pt(7)

def generate_fast_docx(records: List[Dict], template_type: str = 'vertical', 
                      scale_factor: float = 1.0) -> BytesIO:
    """Generate DOCX with fast processing"""
    generator = FastDocxGenerator()
    return generator.generate_labels_fast(records, template_type, scale_factor)

def create_fast_docx_routes(app):
    """Create fast DOCX generation routes"""
    
    @app.route('/api/generate-fast', methods=['POST'])
    def generate_labels_fast():
        """Fast label generation endpoint"""
        try:
            start_time = time.time()
            logger.info("[FAST-DOCX] Fast generation request received")
            
            # Get request data
            data = request.get_json()
            template_type = data.get('template_type', 'vertical')
            scale_factor = float(data.get('scale_factor', 1.0))
            selected_tags = data.get('selected_tags', [])
            
            logger.info(f"[FAST-DOCX] Template: {template_type}, Scale: {scale_factor}, Tags: {len(selected_tags)}")
            
            # Get records from Excel processor
            from src.core.data.excel_processor import get_excel_processor
            excel_processor = get_excel_processor()
            
            if not excel_processor.df or excel_processor.df.empty:
                return jsonify({'error': 'No data available'}), 400
            
            # Filter records by selected tags
            if selected_tags:
                # Simple filtering by product name
                filtered_df = excel_processor.df[
                    excel_processor.df['ProductName'].isin(selected_tags) |
                    excel_processor.df['Product Name*'].isin(selected_tags)
                ]
            else:
                # Use first 50 records for fast generation
                filtered_df = excel_processor.df.head(50)
            
            if filtered_df.empty:
                return jsonify({'error': 'No matching records found'}), 400
            
            # Convert to records format
            records = []
            for _, row in filtered_df.iterrows():
                record = {
                    'ProductName': row.get('ProductName', row.get('Product Name*', '')),
                    'ProductType': row.get('ProductType', row.get('Product Type*', '')),
                    'ProductBrand': row.get('ProductBrand', row.get('Product Brand', '')),
                    'Weight': row.get('CombinedWeight', row.get('Weight', row.get('Weight*', ''))),
                    'Price': row.get('Price* (Tier Name for Bulk)', row.get('Price', '0.00')),
                    'Lineage': row.get('Lineage', 'HYBRID')
                }
                records.append(record)
            
            # Generate DOCX
            output_buffer = generate_fast_docx(records, template_type, scale_factor)
            
            # Create filename
            from datetime import datetime
            today_str = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"AGT_Fast_Labels_{template_type}_{len(records)}tags_{today_str}.docx"
            
            # Return file
            from flask import send_file
            response = send_file(
                output_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            generation_time = time.time() - start_time
            logger.info(f"[FAST-DOCX] Fast generation completed in {generation_time:.2f}s")
            
            return response
            
        except Exception as e:
            logger.error(f"[FAST-DOCX] Error in fast generation: {e}")
            return jsonify({'error': str(e)}), 500
    
    return app

if __name__ == "__main__":
    # Test the generator
    test_records = [
        {
            'ProductName': 'Test Product 1',
            'ProductType': 'Flower',
            'ProductBrand': 'Test Brand',
            'Weight': '3.5g',
            'Price': '25.00',
            'Lineage': 'HYBRID'
        },
        {
            'ProductName': 'Test Product 2',
            'ProductType': 'Concentrate',
            'ProductBrand': 'Test Brand',
            'Weight': '1g',
            'Price': '40.00',
            'Lineage': 'INDICA'
        }
    ]
    
    output = generate_fast_docx(test_records, 'vertical', 1.0)
    print(f"Generated DOCX with {len(test_records)} records")