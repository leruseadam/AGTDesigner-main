#!/usr/bin/env python3
"""
Debug script to test DocxTemplate rendering of ProductStrain in vertical template
"""
import os
import sys
import logging
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from io import BytesIO

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_docxtemplate_rendering():
    """Test DocxTemplate rendering with ProductStrain"""
    
    # Load the vertical template
    template_path = "src/templates/vertical.docx"
    
    try:
        # Create DocxTemplate
        doc = DocxTemplate(template_path)
        logger.info(f"Loaded template: {template_path}")
        
        # Create test context with ProductStrain
        context = {
            'Label1': {
                'ProductName': 'Test Product',
                'ProductBrand': 'TestBrand',
                'ProductStrain': 'CBD Blend',  # This is the value we're testing
                'ProductVendor': 'TestVendor',
                'Lineage': 'CBD',
                'QR': None  # No QR for this test
            },
            # Empty labels
            'Label2': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label3': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label4': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label5': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label6': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label7': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label8': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''},
            'Label9': {'ProductBrand': '', 'DescAndWeight': '', 'Price': '', 'DOH': '', 'Ratio_or_THC_CBD': ''}
        }
        
        logger.info(f"Rendering with context: Label1.ProductStrain = '{context['Label1']['ProductStrain']}'")
        
        # Render the template
        doc.render(context)
        logger.info("DocxTemplate render completed successfully")
        
        # Save to buffer and create Document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        rendered_doc = Document(buffer)
        
        # Save the rendered document
        output_path = "./debug_docxtemplate_output.docx"
        rendered_doc.save(output_path)
        logger.info(f"Saved rendered document: {output_path}")
        
        # Check the content
        logger.info("=== CHECKING RENDERED DOCUMENT CONTENT ===")
        
        # Check paragraphs
        for i, paragraph in enumerate(rendered_doc.paragraphs):
            if paragraph.text.strip():
                logger.info(f"Paragraph {i}: '{paragraph.text}'")
                if 'CBD Blend' in paragraph.text:
                    logger.info(f"✅ Found ProductStrain in paragraph {i}")
        
        # Check tables
        for table_i, table in enumerate(rendered_doc.tables):
            logger.info(f"Table {table_i}:")
            for row_i, row in enumerate(table.rows):
                row_texts = [cell.text for cell in row.cells]
                logger.info(f"  Row {row_i}: {row_texts}")
                for cell_i, cell_text in enumerate(row_texts):
                    if 'CBD Blend' in cell_text:
                        logger.info(f"✅ Found ProductStrain in table {table_i}, row {row_i}, cell {cell_i}")
        
        # Search for ProductStrain specifically
        found_strain = False
        for table in rendered_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'CBD Blend' in cell.text:
                        found_strain = True
                        logger.info(f"✅ ProductStrain 'CBD Blend' found in rendered document!")
                        break
        
        if not found_strain:
            logger.error("❌ ProductStrain 'CBD Blend' NOT found in rendered document")
            
        return found_strain
        
    except Exception as e:
        logger.error(f"Error during DocxTemplate rendering: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("=== Testing DocxTemplate Rendering of ProductStrain ===")
    success = test_docxtemplate_rendering()
    if success:
        print("✅ DocxTemplate rendering test PASSED - ProductStrain appears in output")
    else:
        print("❌ DocxTemplate rendering test FAILED - ProductStrain missing from output")