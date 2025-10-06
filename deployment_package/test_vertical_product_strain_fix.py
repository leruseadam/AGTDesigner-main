#!/usr/bin/env python3
"""
Test script to verify ProductStrain fix for vertical template
"""
import sys
import os
import logging
from pathlib import Path

# Add the project root to Python path
sys.path.insert(0, str(Path(__file__).parent))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def test_product_strain_fix():
    """Test that ProductStrain shows up in vertical template"""
    
    try:
        # Initialize Excel processor
        from src.core.data.excel_processor import ExcelProcessor
        processor = ExcelProcessor()
        
        # Load default file
        default_file = "/Users/adamcordova/Desktop/labelMaker_ QR copy SAFEST copy 4/uploads/A Greener Today - Bothell_inventory_09-26-2025  4_51 PM.xlsx"
        processor.load_file(default_file)
        logger.info(f"Loaded {len(processor.df)} records from Excel")
        
        # Find a record with a strain value (not empty or "Paraphernalia")
        strain_records = processor.df[
            (processor.df['Product Strain'].notna()) & 
            (processor.df['Product Strain'] != '') & 
            (processor.df['Product Strain'] != 'Paraphernalia')
        ]
        
        if strain_records.empty:
            logger.error("No records found with valid strain values")
            return
            
        test_record = strain_records.iloc[0]
        logger.info(f"Testing with product: {test_record['ProductName']}")
        logger.info(f"ProductStrain: '{test_record['Product Strain']}'")
        
        # Test vertical template processing
        from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
        font_scheme = get_font_scheme('vertical')
        template_processor = TemplateProcessor('vertical', font_scheme, 1.0)
        
        # Generate one label
        processor.selected_tags = [test_record.to_dict()]
        
        # Generate document
        logger.info("=== GENERATING VERTICAL TEMPLATE WITH PRODUCT STRAIN ===")
        output_path = template_processor.generate_labels_from_processor(processor, output_dir="./")
        logger.info(f"Generated document: {output_path}")
        
        # Read the generated document and check for ProductStrain content
        from docx import Document
        doc = Document(output_path)
        
        product_strain_found = False
        strain_value = test_record['Product Strain']
        
        logger.info(f"=== CHECKING GENERATED DOCUMENT FOR STRAIN VALUE: '{strain_value}' ===")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if strain_value in cell.text:
                        logger.info(f"✅ SUCCESS: Found ProductStrain value '{strain_value}' in generated document!")
                        logger.info(f"   Cell text: {cell.text.strip()}")
                        product_strain_found = True
        
        for para in doc.paragraphs:
            if strain_value in para.text:
                logger.info(f"✅ SUCCESS: Found ProductStrain value '{strain_value}' in paragraph!")
                logger.info(f"   Paragraph text: {para.text.strip()}")
                product_strain_found = True
        
        if product_strain_found:
            logger.info("🎉 FIX SUCCESSFUL: ProductStrain is now appearing in vertical template!")
        else:
            logger.error("❌ FIX FAILED: ProductStrain value not found in generated document")
            # Show what placeholders exist
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'ProductStrain' in cell.text:
                            logger.error(f"   Found unreplaced placeholder: {cell.text.strip()}")
                            
        # Cleanup
        if os.path.exists(output_path):
            os.remove(output_path)
            logger.info(f"Cleaned up test file: {output_path}")
            
    except Exception as e:
        logger.error(f"Test failed: {e}", exc_info=True)

if __name__ == "__main__":
    test_product_strain_fix()