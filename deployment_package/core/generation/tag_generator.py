import io
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from io import BytesIO
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import re
from docxtpl import DocxTemplate, InlineImage
from copy import deepcopy
import cProfile
from itertools import groupby
import pandas as pd

from src.core.generation.docx_formatting import (
    fix_table_row_heights,
    safe_fix_paragraph_spacing,
    apply_conditional_formatting,
    set_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
    enforce_fixed_cell_dimensions,
)
from src.core.generation.context_builders import (
    build_context,
)
from src.core.formatting.markers import (
    wrap_with_marker,
    FIELD_MARKERS
)
from src.core.utils.resource_utils import resource_path
from src.core.constants import (
    FONT_SCHEME_HORIZONTAL,
    FONT_SCHEME_VERTICAL,
    FONT_SCHEME_MINI,
    LINEAGE_COLOR_MAP
)

# Performance optimization: disable debug logging in production
DEBUG_ENABLED = False

logger = logging.getLogger(__name__)

def _find_most_likely_ounce_weight(product_name, product_type):
    """
    Find the most common ounce weight for similar nonclassic products.
    This helps maintain consistency with actual product packaging rather than mathematical conversion.
    """
    # Common ounce weights for nonclassic products based on typical packaging
    common_oz_weights = {
        # Edibles - common sizes
        'edible (liquid)': ['2.5oz', '3.53oz', '1.7oz'],
        'edible (solid)': ['1oz', '2.5oz', '3.5oz'],
        'gummy': ['1oz', '2oz', '3.5oz'],
        'chocolate': ['1oz', '2oz', '3.5oz'],
        'cookie': ['1oz', '2oz'],
        'brownie': ['1oz', '2oz'],
        'candy': ['1oz', '2oz', '3.5oz'],
        
        # Tinctures and oils
        'tincture': ['1oz', '2oz', '4oz'],
        'drops': ['1oz', '2oz'],
        'liquid': ['1oz', '2oz', '4oz'],
        
        # Topicals
        'topical': ['1oz', '2oz', '4oz'],
        'cream': ['1oz', '2oz', '4oz'],
        'lotion': ['1oz', '2oz', '4oz'],
        'salve': ['1oz', '2oz'],
        'balm': ['1oz', '2oz'],
        
        # Capsules
        'capsule': ['1oz', '2oz'],
        
        # Beverages
        'beverage': ['12oz', '16oz', '20oz'],
        'drink': ['12oz', '16oz', '20oz'],
        'soda': ['12oz', '16oz'],
        'juice': ['12oz', '16oz'],
        
        # Default fallback
        'default': ['1oz', '2oz', '3.5oz']
    }
    
    # Get the most common weight for this product type
    product_type_lower = product_type.lower().strip()
    
    # Try exact match first
    if product_type_lower in common_oz_weights:
        return common_oz_weights[product_type_lower][0]  # Return the first (most common) weight
    
    # Try partial matches for product types that might have variations
    for key, weights in common_oz_weights.items():
        if key != 'default' and key in product_type_lower:
            return weights[0]
    
    # Special handling for Moonshot products (they seem to be 2.5oz or 3.53oz based on the image)
    if 'moonshot' in product_name.lower():
        return '2.5oz'  # Most common Moonshot size
    
    # Default fallback
    return common_oz_weights['default'][0]  # Return '1oz' as default

PLACEHOLDER_MARKERS = {
    "Description": ("DESC_START", "DESC_END"),
    "WeightUnits": ("WEIGHTUNITS_START", "WEIGHTUNITS_END"),
    "ProductBrand": ("PRODUCTBRAND_START", "PRODUCTBRAND_END"),
    "ProductBrand_Center": ("PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END"),
    "Price": ("PRICE_START", "PRICE_END"),
    "Lineage": ("LINEAGE_START", "LINEAGE_END"),
    "DOH": ("{{Label1.DOH}}", ""),
    "Ratio_or_THC_CBD": ("RATIO_START", "RATIO_END"),
    "THC_CBD": ("THC_CBD_START", "THC_CBD_END"),
    "ProductName": ("PRODUCTNAME_START", "PRODUCTNAME_END"),
    "ProductStrain": ("PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END"),
    "ProductType": ("PRODUCTTYPE_START", "PRODUCTTYPE_END"),
    "ProductVendor": ("PRODUCTVENDOR_START", "PRODUCTVENDOR_END")
}

# Import colors from docx_formatting to avoid duplication
from .docx_formatting import COLORS as LINEAGE_COLORS

def get_template_path(template_type):
    """Return the absolute path for a given template type."""
    # Map template types to filenames
    template_files = {
        'horizontal': 'horizontal.docx',
        'vertical': 'vertical.docx',
        'mini': 'mini.docx',
        'double': 'double.docx',
        'inventory': 'inventory.docx'
    }

    # Validate template type
    if template_type not in template_files:
        raise ValueError(f"Invalid template type: {template_type}")

    # Get template directory path
    base_dir = Path(__file__).parent
    template_dir = base_dir / "templates"
    
    # Debug logging
    print(f"Debug - Template lookup: type={template_type}, file={template_files[template_type]}")
    print(f"Debug - Paths: base_dir={base_dir}, template_dir={template_dir}")

    # Ensure template directory exists
    if not template_dir.exists():
        raise ValueError(f"Template directory not found: {template_dir}")

    # Build full path (expected filename)
    expected_filename = template_files[template_type]
    template_path = template_dir / expected_filename
    print(f"Debug - Final template path: {template_path}")

    # Verify template exists, with robust fallbacks for casing/hidden prefixes on some hosts
    if not template_path.exists():
        # Fallback: case-insensitive match ONLY for non-hidden files (ignore . and ~$ temp files)
        expected_lower = expected_filename.lower()
        fallback = None
        for p in template_dir.iterdir():
            if not p.is_file():
                continue
            name = p.name
            # Ignore hidden/office temp files
            if name.startswith('.') or name.startswith('~$'):
                continue
            if name.lower() == expected_lower:
                fallback = p
                break
        if fallback and fallback.exists():
            print(f"Warning - Using fallback template path due to case-only mismatch: {fallback}")
            return str(fallback)
        raise ValueError(f"Template file not found: {template_path}")

    return str(template_path)

def chunk_records(records, chunk_size=9):
    """Split the list of records into chunks of a given size."""
    # CRITICAL FIX: Only deduplicate if there are significantly more duplicates than expected
    # This prevents removing legitimate products that happen to have similar names
    if len(records) > 0:
        # Check if we have excessive duplicates (more than 50% duplicates)
        seen_products = set()
        unique_records = []
        duplicate_count = 0
        
        for record in records:
            product_name = record.get('ProductName', 'Unknown')
            if product_name not in seen_products:
                seen_products.add(product_name)
                unique_records.append(record)
            else:
                duplicate_count += 1
                # CRITICAL FIX: Keep duplicates but log them for transparency
                unique_records.append(record)
                logger.info(f"Keeping duplicate product in chunking: {product_name} (duplicate #{duplicate_count})")
        
        # Only deduplicate if we have excessive duplicates (more than 50% of records)
        duplicate_percentage = (duplicate_count / len(records)) * 100
        if duplicate_percentage > 50:
            logger.warning(f"Excessive duplicates detected ({duplicate_percentage:.1f}%), deduplicating")
            # Remove duplicates in this case
            seen_products = set()
            unique_records = []
            for record in records:
                product_name = record.get('ProductName', 'Unknown')
                if product_name not in seen_products:
                    seen_products.add(product_name)
                    unique_records.append(record)
                else:
                    logger.warning(f"Skipping duplicate product in chunking: {product_name}")
            records = unique_records
        else:
            logger.info(f"Keeping all {len(records)} records (duplicate rate: {duplicate_percentage:.1f}%)")
    
    return [records[i:i+chunk_size] for i in range(0, len(records), chunk_size)]

def flatten_tags(records):
    """Extract Description values from records for tag generation."""
    flat_tags = []
    for record in records:
        description = record.get("Description", "")
        if description and isinstance(description, str):
            flat_tags.append(description.strip())
    return flat_tags

def expand_template_to_4x5_fixed_scaled(template_path, scale_factor=1.0):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    num_cols, num_rows = 4, 5
    col_width_twips = str(int(1.5 * 1440))   # 1.5 inches per column for equal width
    row_height_pts  = Pt(1.5 * 72)           # 1.5 inches per row for equal height
    cut_line_twips  = int(0.001 * 1440)

    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')
    tblPr.insert(0, shd)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    borders = OxmlElement('w:tblBorders')
    for side in ('insideH','insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single"); b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3"); b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)
    cnt = 1
    for r in range(num_rows):
        for c in range(num_cols):
            cell = tbl.cell(r,c)
            cell._tc.clear_content()
            tc = deepcopy(src_tc)
            for t in tc.iter(qn('w:t')):
                if t.text and 'Label1' in t.text:
                    t.text = t.text.replace('Label1', f'Label{cnt}')
            for el in tc.xpath('./*'):
                cell._tc.append(deepcopy(el))
            cnt += 1
    from docx.oxml.shared import OxmlElement as OE
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing'); spacing.set(qn('w:w'), str(cut_line_twips)); spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def process_chunk(args):
    """Process a chunk of records to generate labels."""
    chunk, base_template, font_scheme, orientation, scale_factor = args
    # Mini template expands to 4x5 grid
    if orientation == "mini":
        local_template_buffer = expand_template_to_4x5_fixed_scaled(base_template, scale_factor=scale_factor)
        num_labels = 20  # Fixed: 4x5 grid = 20 labels per page
    else:
        local_template_buffer = base_template
        num_labels = 9
    tpl = DocxTemplate(local_template_buffer)
    context = {}
    image_width = Mm(8) if orientation == "mini" else Mm(9 if orientation == 'vertical' else 12)
    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
    if DEBUG_ENABLED:
        logger.debug(f"DOH image path: {doh_image_path}")
    
    for i in range(num_labels):
        label_data = {}
        if i < len(chunk):
            row = chunk[i]
            doh_value = str(row.get("DOH", "")).strip().upper()
            if DEBUG_ENABLED:
                logger.debug(f"Processing DOH value: {doh_value}")
            product_type = str(row.get("Product Type*", "")).strip().lower()
            if DEBUG_ENABLED:
                logger.debug(f"Product type: {product_type}")
            
            if doh_value == "YES":
                # Use HighCBD.png if product_type starts with 'high cbd'
                if product_type.startswith('high cbd'):
                    high_cbd_image_path = resource_path(os.path.join("templates", "HighCBD.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using HighCBD image: {high_cbd_image_path}")
                    label_data["DOH"] = InlineImage(tpl, high_cbd_image_path, width=image_width)
                else:
                    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using DOH image: {doh_image_path}")
                    label_data["DOH"] = InlineImage(tpl, doh_image_path, width=image_width)
                if DEBUG_ENABLED:
                    logger.debug(f"Created DOH image with width: {image_width}")
            else:
                label_data["DOH"] = ""
                if DEBUG_ENABLED:
                    logger.debug("Skipping DOH image - value is not 'YES'")
                
            # --- Wrap all fields with markers ---
            # Updated price mapping to use correct Excel field name
            price_val = f"{row.get('Price* (Tier Name for Bulk)', row.get('Price', ''))}"
            label_data["Price"] = wrap_with_marker(price_val, "PRICE")  # Fixed: Use "PRICE" marker to match markers.py definition
            
            lineage_text   = str(row.get("Lineage", "")).strip()
            product_brand  = str(row.get("Product Brand", "")).strip()
            product_type   = str(row.get("Product Type*", "")).strip().lower()
            product_strain = str(row.get("Product Strain", "")).strip()
            
            # Fix brand name for paraphernalia products
            if product_brand == "Paraphernalia" and product_type == "paraphernalia":
                product_name = str(row.get("ProductName", ""))
                if " by " in product_name:
                    product_brand = product_name.split(" by ")[-1].strip()
                else:
                    # Fallback to vendor if no "by" in name
                    product_brand = str(row.get("Vendor", "")).strip()
            
            # Only add brand markers for non-classic types
            # Classic types should show lineage instead of brand
            from src.core.constants import CLASSIC_TYPES
            is_classic_type = product_type in [ct.lower() for ct in CLASSIC_TYPES]
            
            if is_classic_type:
                # For classic types, don't add brand markers initially
                # They will be set to lineage later if lineage is available
                pass
            else:
                # For non-classic types, add brand with centering markers
                # Include Product Strain info for color determination
                product_strain = str(row.get("Product Strain", "")).strip()
                if product_strain:
                    brand_content = f"{product_strain} {product_brand.upper()}"
                else:
                    brand_content = product_brand.upper()
                
                label_data["ProductBrand"] = f"PRODUCTBRAND_CENTER_START{brand_content}PRODUCTBRAND_CENTER_END"
                label_data["ProductBrand_Center"] = f"PRODUCTBRAND_CENTER_START{brand_content}PRODUCTBRAND_CENTER_END"
            
            # Add other fields to label_data
            # Get product name
            product_name = str(row.get("ProductName", ""))
            
            # Use Description as the primary field (never force ProductName here)
            description_raw = row.get("Description", "")
            # Normalize None/NaN to empty string to avoid literal 'None'
            try:
                import math  # local import to avoid module overhead at top level
                if description_raw is None:
                    description = ""
                elif isinstance(description_raw, float) and math.isnan(description_raw):
                    description = ""
                else:
                    description = str(description_raw)
                    if description.lower() == "nan":
                        description = ""
            except Exception:
                description = str(description_raw or "")
            
            # Construct WeightUnits from Weight* and Units, but use JointRatio for pre-roll products
            product_type = str(row.get("Product Type*", "")).lower().strip()
            
            # For pre-roll and infused pre-roll products, use JointRatio instead of Weight* + Units
            if product_type in ["pre-roll", "infused pre-roll"]:
                joint_ratio = str(row.get("JointRatio", "")).strip()
                if joint_ratio and joint_ratio not in ['', 'NULL', 'null', '0', '0.0', 'None', 'nan']:
                    weight_units = joint_ratio  # Use JointRatio directly (e.g., "0.5g x 2 Pack")
                else:
                    weight_units = "0.5g x 2 Pack"  # Default for pre-rolls
            else:
                # Regular products: construct from Weight* + Units
                weight = str(row.get("Weight*", "")).strip()
                units = str(row.get("Units", "")).strip()
                
                # Apply weight conversion for nonclassic products (copy most likely ounce weight)
                if weight and units and units.lower() in ['g', 'grams', 'gram']:
                    # Define CLASSIC_TYPES locally to avoid import issues
                    CLASSIC_TYPES = {'flower', 'pre-roll', 'concentrate', 'infused pre-roll', 'solventless concentrate', 'vape cartridge', 'rso/co2 tankers'}
                    is_nonclassic = product_type not in [ct.lower() for ct in CLASSIC_TYPES]
                    
                    if is_nonclassic:
                        # Find the most common ounce weight for this product type
                        most_likely_oz_weight = _find_most_likely_ounce_weight(product_name, product_type)
                        if most_likely_oz_weight:
                            weight, units = most_likely_oz_weight.split(' ', 1) if ' ' in most_likely_oz_weight else (most_likely_oz_weight.replace('oz', ''), 'oz')
                            # Ensure units is 'oz' even if not in the found weight
                            if 'oz' not in units.lower():
                                units = 'oz'
                
                if weight and units:
                    weight_units = f"{weight}{units}"
                elif weight:
                    weight_units = weight
                else:
                    weight_units = ""
            
            # Preserve original ProductName; keep Description as the clean field
            label_data["ProductName"] = product_name  # Do not repurpose ProductName
            label_data["Description"] = description  # Primary clean display field
            label_data["WeightUnits"] = weight_units  # Don't wrap with markers for template rendering
            
            # For edibles, use brand instead of lineage in the label
            edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}
            is_edible = product_type in edible_types
            is_horizontal_or_double_or_vertical = orientation in {"horizontal", "double", "vertical"}
            
            # For classic types, try to get the strain's canonical lineage from the database
            if is_classic_type and product_strain:
                # DEBUG: Processing classic type '{product_type}' with strain '{product_strain}'
                try:
                    from src.core.data.product_database import get_product_database
                    product_db = get_product_database()
                    strain_info = product_db.get_strain_info(product_strain)
                    # DEBUG: Strain info: {strain_info}
                    if strain_info and strain_info.get('canonical_lineage'):
                        lineage_val = strain_info['canonical_lineage'].upper()
                        # DEBUG: Using database lineage: '{lineage_val}'
                    else:
                        # Fallback to Excel lineage if no database lineage found
                        lineage_val = lineage_text.upper() if lineage_text else ""
                        # DEBUG: Using Excel lineage fallback: '{lineage_val}'
                except Exception as e:
                    # Fallback to Excel lineage if database lookup fails
                    lineage_val = lineage_text.upper() if lineage_text else ""
                    # DEBUG: Using Excel lineage due to error: '{lineage_val}' (error: {e})
            elif is_edible:
                lineage_val = product_brand.upper() if product_brand else lineage_text
            else:
                lineage_val = lineage_text.upper() if lineage_text else ""
                
            # No extra space before Lineage in the output
            label_data["Lineage"] = lineage_val  # Don't wrap with markers for template rendering
            
            # For classic types, set ProductBrand and ProductBrand_Center to lineage
            if is_classic_type:
                if lineage_val:
                    # Use the lineage value from database or Excel
                    label_data["ProductBrand"] = lineage_val.strip()  # Don't wrap with markers for template rendering
                    label_data["ProductBrand_Center"] = lineage_val.strip()  # Don't wrap with markers for template rendering
                else:
                    # Fallback to Excel lineage if no database lineage found
                    fallback_lineage = lineage_text.upper() if lineage_text else ""
                    if fallback_lineage:
                        label_data["ProductBrand"] = fallback_lineage.strip()  # Don't wrap with markers for template rendering
                        label_data["ProductBrand_Center"] = fallback_lineage.strip()  # Don't wrap with markers for template rendering
                    else:
                        # No lineage available, set to empty
                        label_data["ProductBrand"] = ""
                        label_data["ProductBrand_Center"] = ""
            label_data["Ratio_or_THC_CBD"] = str(row.get("Ratio", ""))  # Don't wrap with markers for template rendering
            # ProductStrain is now handled by the template processor to ensure proper conversion
            # of non-classic types to "Mixed" or "CBD Blend"
            # Fix: Handle NaN values in JointRatio
            joint_ratio_value = row.get("JointRatio", "")
            if pd.isna(joint_ratio_value) or str(joint_ratio_value).lower() == 'nan':
                joint_ratio_value = ""
            label_data["JointRatio"] = wrap_with_marker(str(joint_ratio_value), "JOINT_RATIO")
            
            # Add THC and CBD from AI and AK columns
            ai_value = str(row.get("AI", "")).strip()
            ak_value = str(row.get("AK", "")).strip()
            
            # Clean up the values (remove 'nan', empty strings, etc.)
            if ai_value in ['nan', 'NaN', '']:
                ai_value = ""
            if ak_value in ['nan', 'NaN', '']:
                ak_value = ""
            
            # Note: Individual THC/CBD values removed - QR codes now provide this information
            
            # DescAndWeight should contain only the description text (mapped to DESC marker in template)
            # Use the processed Description field from above
            desc = description  # Use the processed description from above
            
            if desc:
                desc = re.sub(r'[-\s]+$', '', desc)
            product_type = str(row.get("Product Type*", "")).strip().lower()
            
            # For edibles in double template, use Product Brand instead of Description
            edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}
            if product_type in edible_types and orientation == "double":
                # Use Product Brand instead of Description for edibles in double template
                desc = product_brand if product_brand else desc
            
            # DescAndWeight should only contain the description text, not description + weight
            # This field is mapped to the DESC marker in templates
            combined = desc  # Just use the description, no weight combination
            
            label_data["DescAndWeight"] = wrap_with_marker(combined, "DESC")
            
            context[f"Label{i+1}"] = label_data
            if DEBUG_ENABLED:
                logger.debug(f"Created label data for Label{i+1}")
        else:
            # Empty label data for unused slots
            context[f"Label{i+1}"] = {
                "Description": "",
                "WeightUnits": "",
                "ProductBrand": "",
                "Price": "",
                "Lineage": "",
                "DOH": "",
                "Ratio_or_THC_CBD": "",
                # ProductStrain handled by template processor
                "DescAndWeight": "",
                "JointRatio": ""
            }
            if DEBUG_ENABLED:
                logger.debug(f"Created empty label data for Label{i+1}")

    # Render template
    if DEBUG_ENABLED:
        logger.debug("Rendering template...")
    tpl.render(context)
    if DEBUG_ENABLED:
        logger.debug("Template rendered successfully")
    
    # Save to buffer
    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    if DEBUG_ENABLED:
        logger.debug("Template saved to buffer")
    
    return buffer.getvalue()

def combine_documents(docs):
    """Combine multiple documents into one using a safer method."""
    if not docs:
        return None
        
    try:
        # Use the first document as the base
        master_doc = Document(BytesIO(docs[0]))
        
        # Add spacing between documents
        master_doc.add_paragraph()
        
        # Append each subsequent document
        for doc_bytes in docs[1:]:
            try:
                src_doc = Document(BytesIO(doc_bytes))
                
                # Copy all content from the source document
                for element in src_doc.element.body:
                    # Create a deep copy to avoid reference issues
                    new_element = deepcopy(element)
                    master_doc.element.body.append(new_element)
                    
                # Add spacing between documents
                master_doc.add_paragraph()
                
            except Exception as e:
                logger.error(f"Error combining document: {e}")
                # Continue with other documents instead of failing completely
                continue

        # Save final document to bytes
        final_buffer = BytesIO()
        master_doc.save(final_buffer)
        final_buffer.seek(0)
        
        # Validate the combined document
        try:
            test_doc = Document(final_buffer)
            # Center all tables in the document
            for table in test_doc.tables:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            final_buffer.seek(0)
            return final_buffer.getvalue()
        except Exception as validation_error:
            logger.error(f"Combined document validation failed: {validation_error}")
            raise ValueError(f"Combined document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in combine_documents: {e}")
        raise

def run_full_process_by_group(records, template_path, font_scheme):
    """Process all records using the template, grouped by strain color (lineage)."""
    if not records:
        return None
    # Define canonical lineage order
    lineage_order = list(LINEAGE_COLOR_MAP.keys())
    def get_lineage(rec):
        # Check multiple possible field names for lineage
        possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
        lin = ''
        for field in possible_fields:
            if field in rec and rec[field]:
                lin = str(rec[field]).strip()
                break
        
        # Normalize the lineage value
        lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
        
        # Debug logging
        if DEBUG_ENABLED:
            logger.debug(f"Record: {rec.get('ProductName', 'Unknown')}, Raw lineage: {lin}, Normalized: {lin if lin in lineage_order else 'MIXED'}")
        
        return lin if lin in lineage_order else 'MIXED'
    # Sort records by lineage order, then by ProductName
    records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
    # Group by lineage and chunk within each group
    tag_chunks = []
    for lineage, group in groupby(records_sorted, key=get_lineage):
        group_list = list(group)
        for i in range(0, len(group_list), 9):
            tag_chunks.append(group_list[i:i+9])
    def _inner():
        template_type = Path(template_path).stem
        docs = []
        for tag_chunk in tag_chunks:
            doc = process_chunk(
                tag_chunk,
                template_path,
                font_scheme
            )
            if doc:
                docs.append(doc)
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_group.prof')
    return result

def run_full_process_by_mini(records, template_type, font_scheme, scale_factor=1.0):
    """Process records with the mini template."""
    if not records:
        return None
    def _inner():
        template_path = get_template_path(template_type)
        if hasattr(template_path, "seek"):
            template_path.seek(0)
            try:
                Document(template_path)
                template_path.seek(0)
            except Exception as e:
                raise ValueError(f"Template buffer is not a valid DOCX: {e}")
        chunks = chunk_records(records, chunk_size=20)  # Fixed: 4x5 grid = 20 labels per page
        docs = []
        for chunk in chunks:
            args = (chunk, template_path, font_scheme, template_type, scale_factor)
            docs.append(process_chunk(args))
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_mini.prof')
    return result

def generate_multiple_label_tables(records, template_path):
    """
    For each record, render the template and append the resulting table to a new document.
    Returns a BytesIO buffer with the final DOCX.
    """
    try:
        # Sort records by canonical lineage order, then by name
        lineage_order = list(LINEAGE_COLOR_MAP.keys())
        def get_lineage(rec):
            possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
            lin = ''
            for field in possible_fields:
                if field in rec and rec[field]:
                    lin = str(rec[field]).strip()
                    break
            lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
            return lin if lin in lineage_order else 'MIXED'
        records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
        final_doc = Document()
        # Remove default paragraph if it exists
        if final_doc.paragraphs:
            p = final_doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        # Group by lineage and chunk within each group
        for lineage, group in groupby(records_sorted, key=get_lineage):
            group_list = list(group)
            for i in range(0, len(group_list), 9):
                chunk = group_list[i:i+9]
                # Create a single 3x3 table for this chunk
                table = final_doc.add_table(rows=3, cols=3)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False
                tblPr = table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                table._element.insert(0, tblPr)
                col_width = Inches(1.0)
                tblGrid = OxmlElement('w:tblGrid')
                for _ in range(3):
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), str(int(col_width.inches * 1440)))
                    tblGrid.append(gridCol)
                table._element.insert(0, tblGrid)
                # Fill the table cells with tag data
                for idx in range(9):
                    r, c = divmod(idx, 3)
                    cell = table.cell(r, c)
                    if idx < len(chunk):
                        record = chunk[idx]
                        try:
                            doc = DocxTemplate(template_path)
                            context = build_context(record, doc)
                            doc.render(context)
                            tmp_stream = BytesIO()
                            doc.save(tmp_stream)
                            tmp_stream.seek(0)
                            rendered_doc = Document(tmp_stream)
                            if rendered_doc.tables:
                                src_table = rendered_doc.tables[0]
                                src_cell = src_table.cell(0, 0)
                                for para in cell.paragraphs:
                                    para._element.getparent().remove(para._element)
                                for para in src_cell.paragraphs:
                                    new_para = cell.add_paragraph()
                                    new_para.alignment = para.alignment
                                    for run in para.runs:
                                        try:
                                            new_run = new_para.add_run(run.text)
                                            # Always force Arial Bold for consistency across platforms
                                            new_run.font.name = "Arial"
                                            new_run.font.bold = True
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            if run.font.size:
                                                new_run.font.size = run.font.size
                                            if run.font.color and run.font.color.rgb:
                                                new_run.font.color.rgb = run.font.color.rgb
                                        except Exception as run_error:
                                            logger.warning(f"Error copying run: {run_error}")
                                            new_para.add_run(run.text)
                                if not cell.paragraphs:
                                    cell.add_paragraph()
                        except Exception as record_error:
                            logger.error(f"Error processing record: {record_error}")
                            cell.text = ''
                    else:
                        cell.text = ''
                    cell.width = col_width
                for row in table.rows:
                    row.height = Inches(1.0)
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                # Enforce fixed cell dimensions to prevent any growth
                try:
                    # Safety check: ensure table has valid structure
                    if table and table.rows and len(table.rows) > 0:
                        first_row = table.rows[0]
                        if hasattr(first_row, '_element') and hasattr(first_row._element, 'tc_lst'):
                            enforce_fixed_cell_dimensions(table)
                        else:
                            logger.warning(f"Skipping table with invalid XML structure in tag generator")
                    else:
                        logger.warning(f"Skipping empty or invalid table in tag generator")
                except Exception as e:
                    logger.warning(f"Error enforcing fixed cell dimensions in tag generator: {e}")
                    continue
                
                final_doc.add_paragraph()
        # Center all tables in the final document
        for table in final_doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Ensure all fonts are Arial Bold for consistency across platforms
        from src.core.generation.docx_formatting import enforce_arial_bold_all_text
        enforce_arial_bold_all_text(final_doc)
        
        # Save final document
        output = BytesIO()
        final_doc.save(output)
        output.seek(0)
        
        # Validate the final document
        try:
            test_doc = Document(output)
            output.seek(0)
            return output
        except Exception as validation_error:
            logger.error(f"Final document validation failed: {validation_error}")
            raise ValueError(f"Generated document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in generate_multiple_label_tables: {e}")
        raise

def set_table_borders(table):
    """Apply consistent border formatting matching main application style."""
    tblPr = table._element.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
        
    tblBorders = OxmlElement('w:tblBorders')
    
    # Remove outer borders
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
        
    # Add light gray interior lines
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
        
    tblPr.append(tblBorders)

def debug_markers(text):
    """Debug helper to identify marker issues."""
    markers = ['DESC', 'WEIGHTUNITS', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 'THC_CBD']
    found_markers = []
    
    for marker in markers:
        start = f"{marker}_START"
        end = f"{marker}_END"
        if start in text:
            pos = text.find(start)
            found_markers.append(f"{marker} start at {pos}")
        if end in text:
            pos = text.find(end)
            found_markers.append(f"{marker} end at {pos}")
            
    if found_markers:
        if DEBUG_ENABLED:
            logger.debug(f"Found markers in text: {text}")
        for marker in found_markers:
            if DEBUG_ENABLED:
                logger.debug(f"  {marker}")
    return found_markers

def validate_and_repair_document(doc_bytes):
    """Validate a document and attempt to repair common issues."""
    try:
        # Try to load the document
        doc = Document(BytesIO(doc_bytes))
        
        # Check for common corruption issues
        issues_found = []
        
        # Check if document has content
        if not doc.paragraphs and not doc.tables:
            issues_found.append("Document has no content")
            
        # Check for malformed tables
        for table in doc.tables:
            try:
                # Try to access table properties
                _ = table.rows
                _ = table.columns
            except Exception as e:
                issues_found.append(f"Malformed table: {e}")
                
        # Check for malformed paragraphs
        for para in doc.paragraphs:
            try:
                # Try to access paragraph properties
                _ = para.runs
            except Exception as e:
                issues_found.append(f"Malformed paragraph: {e}")
        
        if issues_found:
            logger.warning(f"Document validation issues found: {issues_found}")
            return False, issues_found
            
        return True, []
        
    except Exception as e:
        logger.error(f"Document validation failed: {e}")
        return False, [f"Document cannot be loaded: {e}"]

def create_safe_document():
    """Create a minimal safe document as a fallback."""
    try:
        doc = Document()
        # Add a simple paragraph to ensure the document is valid
        doc.add_paragraph("Document generated successfully.")
        return doc
    except Exception as e:
        logger.error(f"Error creating safe document: {e}")
        raise



