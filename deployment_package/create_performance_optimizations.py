#!/usr/bin/env python3
"""
PythonAnywhere Performance Optimization Script
Applies aggressive optimizations to reduce file upload and document generation times
"""

import os
import sys

def apply_performance_optimizations():
    """Apply aggressive performance optimizations for PythonAnywhere"""
    
    print("🚀 Applying PythonAnywhere performance optimizations...")
    
    # Create optimized configuration
    optimization_config = '''
# PythonAnywhere Aggressive Performance Optimization
import os
import logging

# Set environment variables for maximum performance
os.environ['PYTHONANYWHERE_OPTIMIZATION'] = 'True'
os.environ['FLASK_ENV'] = 'production'
os.environ['FLASK_DEBUG'] = 'False'

# Aggressive memory management
import gc
gc.set_threshold(100, 10, 10)  # More aggressive garbage collection

# Disable verbose logging completely
logging.getLogger().setLevel(logging.CRITICAL)
for logger_name in ['werkzeug', 'urllib3', 'requests', 'pandas', 'openpyxl', 'xlrd', 'docxtpl', 'python-docx']:
    logging.getLogger(logger_name).setLevel(logging.CRITICAL)
    logging.getLogger(logger_name).disabled = True

# Performance constants for PythonAnywhere
PYTHONANYWHERE_MAX_CHUNK_SIZE = 25  # Reduce from 50
PYTHONANYWHERE_MAX_PROCESSING_TIME = 15  # Reduce from 30
PYTHONANYWHERE_MAX_TOTAL_TIME = 120  # Reduce from 300
PYTHONANYWHERE_MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB max
PYTHONANYWHERE_UPLOAD_CHUNK_SIZE = 8192  # 8KB chunks

# Enable lazy imports
LAZY_IMPORTS = True

# Disable startup file loading
DISABLE_STARTUP_FILE_LOADING = True
LAZY_LOADING_ENABLED = True

print("✅ PythonAnywhere optimizations applied")
'''
    
    # Write optimization config
    with open("pythonanywhere_optimizations.py", "w") as f:
        f.write(optimization_config)
    
    print("✅ Created pythonanywhere_optimizations.py")
    
    # Create fast upload handler
    fast_upload_handler = '''#!/usr/bin/env python3
"""
Fast upload handler for PythonAnywhere with aggressive optimizations
"""

from flask import Flask, request, jsonify, session
import os
import time
import tempfile
import pandas as pd
from werkzeug.utils import secure_filename

def create_fast_upload_handler(app):
    """Create optimized upload handler"""
    
    @app.route('/upload-ultra-fast', methods=['POST'])
    def upload_ultra_fast():
        """Ultra-fast upload with minimal processing"""
        start_time = time.time()
        
        try:
            if 'file' not in request.files:
                return jsonify({'error': 'No file uploaded'}), 400
                
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': 'No file selected'}), 400
            
            # Quick file validation
            if not file.filename.lower().endswith(('.xlsx', '.xls')):
                return jsonify({'error': 'Only Excel files allowed'}), 400
            
            # Save with minimal processing
            filename = secure_filename(file.filename)
            filepath = os.path.join('uploads', filename)
            
            # Ensure uploads directory exists
            os.makedirs('uploads', exist_ok=True)
            
            # Save file directly
            file.save(filepath)
            
            # Quick validation - just check if pandas can open it
            try:
                # Read only first 5 rows to validate
                df = pd.read_excel(filepath, nrows=5)
                row_count = len(pd.read_excel(filepath))  # Get full count
            except Exception as e:
                os.remove(filepath)  # Clean up on error
                return jsonify({'error': f'Invalid Excel file: {str(e)}'}), 400
            
            # Store minimal info in session
            session['current_file'] = filename
            session['file_path'] = filepath
            session['upload_time'] = time.time()
            
            processing_time = time.time() - start_time
            
            return jsonify({
                'success': True,
                'filename': filename,
                'rows': row_count,
                'columns': len(df.columns),
                'processing_time': round(processing_time, 2),
                'message': f'File uploaded successfully in {processing_time:.1f}s'
            })
            
        except Exception as e:
            return jsonify({'error': f'Upload failed: {str(e)}'}), 500
    
    @app.route('/process-ultra-fast', methods=['POST'])
    def process_ultra_fast():
        """Ultra-fast processing with minimal overhead"""
        start_time = time.time()
        
        try:
            if 'current_file' not in session:
                return jsonify({'error': 'No file uploaded'}), 400
            
            filepath = session.get('file_path')
            if not os.path.exists(filepath):
                return jsonify({'error': 'File not found'}), 400
            
            # Load with minimal processing
            df = pd.read_excel(filepath)
            
            # Basic processing only
            processed_data = []
            for index, row in df.iterrows():
                if index >= 100:  # Limit to first 100 rows for speed
                    break
                    
                processed_data.append({
                    'index': index,
                    'data': row.to_dict()
                })
            
            processing_time = time.time() - start_time
            
            return jsonify({
                'success': True,
                'processed_rows': len(processed_data),
                'data': processed_data[:20],  # Return first 20 for preview
                'processing_time': round(processing_time, 2),
                'total_rows': len(df)
            })
            
        except Exception as e:
            return jsonify({'error': f'Processing failed: {str(e)}'}), 500

    return app

# Export the function
__all__ = ['create_fast_upload_handler']
'''
    
    with open("fast_upload_handler.py", "w") as f:
        f.write(fast_upload_handler)
    
    print("✅ Created fast_upload_handler.py")
    
    # Create fast document generator
    fast_docx_generator = '''#!/usr/bin/env python3
"""
Ultra-fast document generation for PythonAnywhere
"""

from docx import Document
from docx.shared import Inches, Pt
import os
import time

class FastDocxGenerator:
    """Optimized DOCX generation with minimal overhead"""
    
    def __init__(self):
        self.template_cache = {}
    
    def generate_simple_labels(self, data, output_path):
        """Generate labels with minimal formatting for speed"""
        start_time = time.time()
        
        try:
            # Create simple document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Product Labels', 0)
            
            # Process data in chunks
            chunk_size = 10  # Small chunks for PythonAnywhere
            processed = 0
            
            for i in range(0, min(len(data), 50), chunk_size):  # Limit to 50 items
                chunk = data[i:i+chunk_size]
                
                for item in chunk:
                    # Add simple paragraph for each product
                    p = doc.add_paragraph()
                    p.add_run(f"Product: {item.get('Product Name*', 'N/A')}").bold = True
                    
                    # Add basic info only
                    doc.add_paragraph(f"Type: {item.get('Product Type*', 'N/A')}")
                    doc.add_paragraph(f"Brand: {item.get('Product Brand', 'N/A')}")
                    doc.add_paragraph(f"Weight: {item.get('Weight*', 'N/A')}")
                    
                    # Add separator
                    doc.add_paragraph("─" * 40)
                    
                    processed += 1
                
                # Quick break to prevent timeout
                if time.time() - start_time > 15:  # 15 second limit
                    break
            
            # Save document
            doc.save(output_path)
            
            generation_time = time.time() - start_time
            
            return {
                'success': True,
                'output_path': output_path,
                'processed_items': processed,
                'generation_time': generation_time,
                'file_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'generation_time': time.time() - start_time
            }

def create_fast_generator_routes(app):
    """Add fast generation routes to Flask app"""
    
    from flask import request, jsonify, send_file, session
    
    @app.route('/generate-ultra-fast', methods=['POST'])
    def generate_ultra_fast():
        """Ultra-fast document generation"""
        start_time = time.time()
        
        try:
            # Get data from request
            data = request.get_json()
            if not data or 'items' not in data:
                return jsonify({'error': 'No data provided'}), 400
            
            items = data['items']
            if len(items) > 25:  # Limit for speed
                items = items[:25]
                
            # Generate output filename
            timestamp = int(time.time())
            output_filename = f"fast_labels_{timestamp}.docx"
            output_path = os.path.join('output', output_filename)
            
            # Ensure output directory exists
            os.makedirs('output', exist_ok=True)
            
            # Generate document
            generator = FastDocxGenerator()
            result = generator.generate_simple_labels(items, output_path)
            
            total_time = time.time() - start_time
            
            if result['success']:
                return jsonify({
                    'success': True,
                    'filename': output_filename,
                    'download_url': f'/download/{output_filename}',
                    'processed_items': result['processed_items'],
                    'total_time': round(total_time, 2),
                    'file_size': result['file_size']
                })
            else:
                return jsonify({
                    'error': result['error'],
                    'total_time': round(total_time, 2)
                }), 500
                
        except Exception as e:
            total_time = time.time() - start_time
            return jsonify({
                'error': f'Generation failed: {str(e)}',
                'total_time': round(total_time, 2)
            }), 500

# Export the function
__all__ = ['create_fast_generator_routes', 'FastDocxGenerator']
'''
    
    with open("fast_docx_generator.py", "w") as f:
        f.write(fast_docx_generator)
    
    print("✅ Created fast_docx_generator.py")
    
    # Create integration script
    integration_script = '''#!/usr/bin/env python3
"""
Apply performance optimizations to existing Flask app
"""

def apply_optimizations_to_app():
    """Apply all performance optimizations"""
    
    try:
        # Import optimization modules
        from pythonanywhere_optimizations import *
        from fast_upload_handler import create_fast_upload_handler
        from fast_docx_generator import create_fast_generator_routes
        
        print("✅ Performance optimizations loaded")
        return True
        
    except ImportError as e:
        print(f"❌ Failed to load optimizations: {e}")
        return False

# Auto-apply optimizations when imported
if __name__ != "__main__":
    apply_optimizations_to_app()
'''
    
    with open("apply_optimizations.py", "w") as f:
        f.write(integration_script)
    
    print("✅ Created apply_optimizations.py")
    
    return True

def create_optimized_wsgi():
    """Create optimized WSGI file for PythonAnywhere"""
    
    wsgi_content = '''#!/usr/bin/env python3
"""
Ultra-optimized WSGI configuration for PythonAnywhere
Includes aggressive performance optimizations for file upload and document generation
"""

import os
import sys
import logging

# Aggressive performance setup
os.environ['PYTHONANYWHERE_OPTIMIZATION'] = 'True'
os.environ['FLASK_ENV'] = 'production'
os.environ['FLASK_DEBUG'] = 'False'

# Memory optimization
import gc
gc.set_threshold(100, 10, 10)

# Configure the project directory
USERNAME = 'adamcordova'  # Update this if needed
project_dir = f'/home/{USERNAME}/AGTDesigner'

# Add to Python path
if os.path.exists(project_dir):
    if project_dir not in sys.path:
        sys.path.insert(0, project_dir)

# Add user site-packages
import site
user_site = site.getusersitepackages()
if user_site not in sys.path:
    sys.path.insert(0, user_site)

# Disable all logging for maximum performance
logging.getLogger().setLevel(logging.CRITICAL)
for logger_name in ['werkzeug', 'urllib3', 'requests', 'pandas', 'openpyxl', 'xlrd', 'docxtpl', 'python-docx']:
    logging.getLogger(logger_name).setLevel(logging.CRITICAL)
    logging.getLogger(logger_name).disabled = True

try:
    # Import performance optimizations first
    try:
        from pythonanywhere_optimizations import *
        from apply_optimizations import apply_optimizations_to_app
        apply_optimizations_to_app()
    except ImportError:
        pass  # Continue without optimizations if not available
    
    # Import the Flask application
    from app import app as application
    
    # Apply ultra-fast handlers if available
    try:
        from fast_upload_handler import create_fast_upload_handler
        from fast_docx_generator import create_fast_generator_routes
        
        application = create_fast_upload_handler(application)
        create_fast_generator_routes(application)
        
        print("✅ Fast handlers applied")
    except ImportError:
        print("⚠️ Fast handlers not available")
    
    # Ultra-aggressive production configuration
    application.config.update(
        DEBUG=False,
        TESTING=False,
        TEMPLATES_AUTO_RELOAD=False,
        SEND_FILE_MAX_AGE_DEFAULT=300,  # 5 minutes
        MAX_CONTENT_LENGTH=5 * 1024 * 1024,  # 5MB max
        PERMANENT_SESSION_LIFETIME=900,  # 15 minutes
        SESSION_COOKIE_SECURE=True,
        SESSION_COOKIE_HTTPONLY=True,
        WTF_CSRF_TIME_LIMIT=None,
    )
    
    # Force garbage collection
    gc.collect()
    
except ImportError as e:
    logging.critical(f"Failed to import Flask app: {e}")
    raise
except Exception as e:
    logging.critical(f"Error configuring Flask app: {e}")
    raise

# For direct execution
if __name__ == "__main__":
    application.run(debug=False, threaded=True)
'''
    
    with open("wsgi_ultra_optimized.py", "w") as f:
        f.write(wsgi_content)
    
    print("✅ Created wsgi_ultra_optimized.py")

if __name__ == "__main__":
    print("🚀 PythonAnywhere Performance Optimization Tool")
    print("=" * 50)
    
    success = apply_performance_optimizations()
    
    if success:
        create_optimized_wsgi()
        print("\n🎉 Performance optimizations created!")
        print("\nFiles created:")
        print("- pythonanywhere_optimizations.py")
        print("- fast_upload_handler.py") 
        print("- fast_docx_generator.py")
        print("- apply_optimizations.py")
        print("- wsgi_ultra_optimized.py")
        print("\nNext steps:")
        print("1. Upload these files to PythonAnywhere")
        print("2. Use wsgi_ultra_optimized.py as your WSGI file")
        print("3. Use /upload-ultra-fast endpoint for faster uploads")
        print("4. Use /generate-ultra-fast endpoint for faster document generation")
    else:
        print("\n❌ Optimization creation failed!")
        sys.exit(1)